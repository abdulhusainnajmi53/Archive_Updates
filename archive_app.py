import sys
import os
import json
import sqlite3
try:
    import pandas as pd
except ImportError:
    pd = None

import subprocess
import urllib.parse
import shutil
import re
import time
import contextlib
import traceback
import tempfile
import numpy as np
from datetime import datetime, date, timedelta
from updater import UpdateChecker, UpdateDownloader, install_update
from version import APP_VERSION

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QTableWidget, QTableWidgetItem, QPushButton, QLineEdit, 
    QLabel, QFileDialog, QHeaderView, QStackedWidget, QFrame,
    QMessageBox, QAbstractItemView, QStyledItemDelegate, QTextEdit,
    QComboBox, QCompleter, QListWidget, QStyle, QListView, QScrollArea,
    QSizePolicy, QTabWidget, QTabBar, QDialog, QFormLayout, QStyleOptionViewItem,
    QGroupBox, QCheckBox, QSlider, QGridLayout, QSplashScreen, QGraphicsDropShadowEffect,
    QMenu, QInputDialog, QProgressDialog, QButtonGroup, QRadioButton, QListWidgetItem, QDialogButtonBox,
    QProgressBar
)
from PyQt6.QtCore import Qt, QSize, QEvent, QPropertyAnimation, QEasingCurve, QRect, QTimer, QThread, pyqtSignal, QRectF
from PyQt6.QtGui import (
    QFont, QFontDatabase, QColor, QKeyEvent, QShortcut, 
    QKeySequence, QFontMetrics, QPalette, QTextOption, QPainter,
    QTextDocument, QAbstractTextDocumentLayout, QPixmap, QImage, QPen, QPainterPath, 
    QLinearGradient, QBrush, QStandardItemModel, QStandardItem
)
import openpyxl
from openpyxl.styles import Font, Alignment


# Word processing library check
try:
    from docx import Document
except ImportError:
    Document = None
class ImportSelectionDialog(QDialog):
    def __init__(self, sheet_names, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Select Sheets to Import")
        self.resize(400, 500)
        self.sheet_names = sheet_names
        self.selected_sheets = []
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        lbl = QLabel("Select the sheets you want to import.\nUnchecked sheets will NOT be modified in the database.")
        lbl.setWordWrap(True)
        layout.addWidget(lbl)
        
        # Checkboxes List
        self.list_widget = QListWidget()
        for sheet in self.sheet_names:
            item = QListWidgetItem(sheet)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            item.setCheckState(Qt.CheckState.Checked)  # Default all checked
            self.list_widget.addItem(item)
            
        layout.addWidget(self.list_widget)
        
        # Select All / None
        btn_layout = QHBoxLayout()
        btn_all = QPushButton("Select All")
        btn_all.clicked.connect(lambda: self.toggle_all(True))
        btn_none = QPushButton("Select None")
        btn_none.clicked.connect(lambda: self.toggle_all(False))
        btn_layout.addWidget(btn_all)
        btn_layout.addWidget(btn_none)
        layout.addLayout(btn_layout)
        
        # Buttons
        btn_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btn_box.accepted.connect(self.accept_selection)
        btn_box.rejected.connect(self.reject)
        layout.addWidget(btn_box)
        
    def toggle_all(self, state):
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            item.setCheckState(Qt.CheckState.Checked if state else Qt.CheckState.Unchecked)
            
    def accept_selection(self):
        self.selected_sheets = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.checkState() == Qt.CheckState.Checked:
                self.selected_sheets.append(item.text())
        self.accept()


class AdvancedFilterDialog(QDialog):
    def __init__(self, parent=None, current_filters=None, history=None, db_path=None):
        super().__init__(parent)
        self.setWindowTitle("Advanced Filters")
        self.resize(1000, 700)
        self.parent_app = parent
        self.db_path = db_path
        
        # Current Filters: { col_index: {set of values} }
        self.current_filters: dict[int, set[str]] = current_filters if current_filters else {}
        # Make a deep copy for local editing
        self.local_filters: dict[int, set[str]] = {k: set(v) for k, v in self.current_filters.items()}
        
        self.history = history if history else []
        
        # Map Tab Index to DB Column Index
        # Tabs: Person, Occasion, Category, Place, Country, Year, A/V, Incomplete
        self.tab_map = [
            (2, "Person", "person"),
            (3, "Occasion", "occasion"),
            (4, "Category", "category"),
            (5, "Place", "place"),
            (6, "Country", "country"),
            (9, "Year", "year"),
            (13, "AV", "AV"),
            (16, "Incomplete", "incomplete")
        ]
        
        # Load Data into Pandas DataFrame for Smart Cascading
        self.df = pd.DataFrame()
        self.load_data()
        
        self.init_ui()
        self.load_history()
        self.show()

    def load_data(self):
        """Loads ALL data (Master + Other Sheets) into a single DataFrame."""
        if not self.db_path: return

        try:
            conn = sqlite3.connect(self.db_path)
            
            # 1. Load Master Sheet
            cols = ["person", "occasion", "category", "place", "country", "year", "AV", "incomplete"]
            # Note: "AV" column in DB might be "a_v" or "av". We'll select by index/alias if needed, 
            # but usually it's best to select * and map.
            # Safest is to read all and filter, or select specific.
            # The schema has: audio_no, person, occasion, category, place, country, hijri, esavi, year, outof, remarks, summary, av, ...
            
            query = "SELECT person, occasion, category, place, country, year, av, incomplete FROM events"
            df_master = pd.read_sql_query(query, conn)
            # Rename 'av' to 'AV' to match our internal key if needed, or just normalize keys.
            df_master.columns = [c.lower() for c in df_master.columns] # person, occasion... av...
            
            # 2. Load Other Sheets
            data_frames = [df_master]
            
            sheets = conn.execute("SELECT id, name FROM old_sheets_meta").fetchall()
            allowed_sheets = {"SAF", "Mumbai Cass", "Surat Cass", "Karachi Cass", "Karachi Video"}
            
            for sid, sname in sheets:
                if sname not in allowed_sheets: continue
                
                # Fetch headers to map columns
                headers = conn.execute("SELECT col_index, header_name FROM old_sheet_headers WHERE sheet_id = ?", (sid,)).fetchall()
                
                # Map standard keys to col_X
                # We need: person, occasion, category, place, country, year, av, incomplete
                col_mapping = {}
                target_keys = ["person", "occasion", "category", "place", "country", "year", "av", "incomplete"]
                
                for h_idx, h_name in headers:
                    h_lower = h_name.lower().strip()
                    # Normalize common variations
                    if h_lower == "a/v": h_lower = "av"
                    
                    if h_lower in target_keys:
                        col_mapping[f"col_{h_idx}"] = h_lower
                
                if col_mapping:
                    # Select only mapped columns
                    sel_cols = ", ".join(col_mapping.keys())
                    sheet_query = f"SELECT {sel_cols} FROM old_sheet_data_{sid}"
                    try:
                        df_sheet = pd.read_sql_query(sheet_query, conn)
                        df_sheet.rename(columns=col_mapping, inplace=True)
                        data_frames.append(df_sheet)
                    except:
                        pass
                        
            conn.close()
            
            # 3. Combine
            self.df = pd.concat(data_frames, ignore_index=True)
            
            # 4. Clean Data
            # Clean Strings
            for col in self.df.columns:
                self.df[col] = self.df[col].astype(str).replace(['None', 'nan', ''], np.nan)
            
            # Normalize Year: Remove decimals (1392.0 -> 1392)
            if 'year' in self.df.columns:
                 self.df['year'] = pd.to_numeric(self.df['year'], errors='coerce').fillna(0).astype(int).astype(str)
                 self.df['year'] = self.df['year'].replace('0', '') # Revert 0 to empty
                 
            # Fill NaNs with empty string for UI
            self.df.fillna("", inplace=True)
            
        except Exception as e:
            print(f"Error loading filter data: {e}")

    @contextlib.contextmanager
    def busy_cursor(self, button=None):
        """Standard indicator for long operations."""
        if button and hasattr(button, 'setEnabled'):
            button.setEnabled(False)
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            yield
        finally:
            QApplication.restoreOverrideCursor()

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 1. TABS
        self.tabs = QTabWidget()
        self.tab_widgets = [] # List of (list_widget, search_box) tuples
        
        # Use custom font
        font = QFont("Arial", 12)
        if hasattr(self.parent_app, 'lisan_font_family'):
             font = QFont(self.parent_app.lisan_font_family, 14)
        
        for i, (col_idx, label, db_col) in enumerate(self.tab_map):
            tab = QWidget()
            t_layout = QVBoxLayout(tab)
            
            # Search
            search = QLineEdit()
            search.setPlaceholderText(f"Search {label}...")
            search.textChanged.connect(lambda text, idx=i: self.filter_list_items(idx, text))
            t_layout.addWidget(search)
            
            # List
            lst = QListWidget()
            lst.setFont(font)
            lst.setSelectionMode(QAbstractItemView.SelectionMode.NoSelection) # We use checkboxes
            
            # Populate initially
            self.populate_list(lst, col_idx, db_col)
            
            # Connect Item Changed
            lst.itemChanged.connect(lambda item, idx=i: self.on_item_changed(item, idx))
            
            t_layout.addWidget(lst)
            
            # Select All / None
            btn_layout = QHBoxLayout()
            btn_all = QPushButton("Select All")
            btn_all.clicked.connect(lambda _, l=lst, idx=i: self.select_all(l, idx, True))
            btn_none = QPushButton("Select None")
            btn_none.clicked.connect(lambda _, l=lst, idx=i: self.select_all(l, idx, False))
            btn_layout.addWidget(btn_all)
            btn_layout.addWidget(btn_none)
            t_layout.addLayout(btn_layout)
            
            self.tab_widgets.append((lst, search))
            self.tabs.addTab(tab, label)
            
        self.tabs.currentChanged.connect(self.on_tab_changed)
        layout.addWidget(self.tabs)
        
        # 2. HISTORY
        grp_history = QGroupBox("Filter History (Last 10)")
        h_layout = QVBoxLayout(grp_history)
        self.list_history = QListWidget()
        self.list_history.setFixedHeight(100)
        self.list_history.itemClicked.connect(self.restore_history)
        h_layout.addWidget(self.list_history)
        layout.addWidget(grp_history)
        
        # 3. DIALOG BUTTONS
        btn_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btn_box.accepted.connect(self.accept)
        btn_box.rejected.connect(self.reject)
        layout.addWidget(btn_box)
        
    def populate_list(self, list_widget, col_idx, db_col):
        """Filters the DataFrame based on other selections and populates valid options."""
        list_widget.blockSignals(True)
        list_widget.clear()
        
        if self.df.empty:
            list_widget.blockSignals(False)
            return

        # Normalized column name (lowercase)
        target_col = db_col.lower()
        
        # Start with full dataframe
        df_filtered = self.df.copy()
        
        # Apply filters from OTHER tabs
        for c_idx, set_vals in self.local_filters.items():
            if c_idx == col_idx: continue # Don't filter self by self
            if not set_vals: continue
            
            # Find column name for c_idx
            c_name_map = next((x[2] for x in self.tab_map if x[0] == c_idx), None)
            if not c_name_map: continue
            c_name = c_name_map.lower()
            
            if c_name not in df_filtered.columns: continue
            
            # Handle "(Blanks)"
            filter_vals = list(set_vals)
            if "(Blanks)" in filter_vals:
                # Include empty strings or blanks
                # We normalized empty to ""
                if "(Blanks)" in filter_vals: filter_vals.remove("(Blanks)")
                filter_vals.append("")
            
            df_filtered = df_filtered[df_filtered[c_name].isin(filter_vals)]
            
        # Get unique values for this column
        if target_col in df_filtered.columns:
            unique_vals = df_filtered[target_col].unique()
            # Sort: empty first, then alpha
            # Filter out None/Nan just in case
            valid_vals = [str(x) for x in unique_vals if x is not None]
            valid_vals.sort(key=lambda x: (x != "", x.lower()))
            
            # Replace "" with "(Blanks)" for UI
            display_vals = []
            for v in valid_vals:
                if v == "":
                   display_vals.append("(Blanks)")
                else:
                   display_vals.append(v)
                   
            # Current selection for this tab
            current_selected = self.local_filters.get(col_idx, set())
            
            for v in display_vals:
                item = QListWidgetItem(v)
                item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                if v in current_selected:
                    item.setCheckState(Qt.CheckState.Checked)
                else:
                    item.setCheckState(Qt.CheckState.Unchecked)
                list_widget.addItem(item)
                
        list_widget.blockSignals(False)

    def filter_list_items(self, tab_idx, text):
        lst = self.tab_widgets[tab_idx][0]
        text = text.lower()
        for i in range(lst.count()):
            item = lst.item(i)
            item.setHidden(text not in item.text().lower())

    def on_item_changed(self, item, tab_index):
        """Update local state on checkbox toggle and Refresh Cascading."""
        col_idx = self.tab_map[tab_index][0]
        val = item.text()
        
        if col_idx not in self.local_filters:
            self.local_filters[col_idx] = set()
            
        if item.checkState() == Qt.CheckState.Checked:
            self.local_filters[col_idx].add(val)
        else:
            self.local_filters[col_idx].discard(val)
            if not self.local_filters[col_idx]:
                del self.local_filters[col_idx]
                
        # CASCADE: Refresh OTHER tabs to reflect new constraints
        # We iterate all tabs and re-populate if they are NOT the current tab
        for i, (l_widget, _) in enumerate(self.tab_widgets):
            if i == tab_index: continue # Don't refresh self (would lose scroll/focus)
            
            c_idx, _, db_col = self.tab_map[i]
            # Refresh this list
            self.populate_list(l_widget, c_idx, db_col)
            if i == tab_index: continue # Don't refresh self (would lose scroll/focus)
            
            c_idx, _, db_col = self.tab_map[i]
            # Refresh this list
            self.populate_list(l_widget, c_idx, db_col)

    def on_tab_changed(self, index):
        """Optional: Refresh the tab we just switched to, just in case."""
        if index < 0 or index >= len(self.tab_widgets): return
        
        l_widget, _ = self.tab_widgets[index]
        c_idx, _, db_col = self.tab_map[index]
        self.populate_list(l_widget, c_idx, db_col)
                
    def select_all(self, lst, tab_idx, select=True):
        lst.blockSignals(True)
        for i in range(lst.count()):
            item = lst.item(i)
            if not item.isHidden():
                item.setCheckState(Qt.CheckState.Checked if select else Qt.CheckState.Unchecked)
        lst.blockSignals(False)
        
        # Manually trigger update
        self.on_item_changed(None, tab_idx)

    def on_tab_changed(self, index):
        # Refresh the list in the NEW tab based on filters in OTHER tabs
        col_idx = self.tab_map[index][0]
        db_col = self.tab_map[index][2]
        lst = self.tab_widgets[index][0]
        
        # Save current selections (already in local_filters via on_item_changed)
        # But we need to make sure populate_list respects them
        self.populate_list(lst, col_idx, db_col)

    def load_history(self):
        self.list_history.clear()
        # Parse history list of dicts -> Strings
        for i, h_state in enumerate(reversed(self.history)):
            
            # Skip empty or value-less states
            if not h_state or not any(h_state.values()):
                continue

            # Create a summary string
            parts = []
            for k, v in h_state.items():
                if not v: continue # Skip empty sets
                
                # Find label
                label = next((x[1] for x in self.tab_map if x[0] == k), f"Col {k}")
                val_str = ", ".join(list(v)[:3])
                if len(v) > 3: val_str += "..."
                parts.append(f"{label}:[{val_str}]")
            
            if not parts: continue # Should not happen if we checked any(values)
            
            summary = " | ".join(parts)
            item = QListWidgetItem(summary)
            # Store the actual state dict in UserRole
            item.setData(Qt.ItemDataRole.UserRole, h_state)
            self.list_history.addItem(item)
            
    def restore_history(self, item):
        state = item.data(Qt.ItemDataRole.UserRole)
        print(f"DEBUG: Restoring history state: {state}")
        self.local_filters = {k: set(v) for k, v in state.items()}
        
        # Refresh current tab
        idx = self.tabs.currentIndex()
        col_idx = self.tab_map[idx][0]
        db_col = self.tab_map[idx][2]
        lst = self.tab_widgets[idx][0]
        
        self.populate_list(lst, col_idx, db_col)

# --- 1. GLOBAL PATHS ---
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))

    return os.path.join(base_path, relative_path)


def resolve_paths():
    """
    Determines the Database and Backup paths.
    Prioritizes External Volume: '/Volumes/MAHAD 1/Documents/Microsoft Acess'
    Falls back to Local Folder if Volume is missing.
    """
    VOL_NAME = "MAHAD 1"
    VOL_PATH = "/Volumes/MAHAD 1"
    
    TARGET_BASE = os.path.join(VOL_PATH, "Documents", "Microsoft Acess", "Dh Archive")
    TARGET_DB = os.path.join(TARGET_BASE, "archive_database.db")
    TARGET_BACKUP = os.path.join(TARGET_BASE, "Bakcup", "DH Archive")
    
    # Check External Volume DB
    if os.path.exists(TARGET_DB):
        return TARGET_BASE, TARGET_DB, TARGET_BACKUP

    # Fallback to Local (Script Directory)
    LOCAL_BASE = os.path.dirname(os.path.abspath(__file__))
    LOCAL_DB = os.path.join(LOCAL_BASE, "archive_database.db")
    LOCAL_BACKUP = os.path.join(LOCAL_BASE, "backups")

    # If Local DB exists, use it
    if os.path.exists(LOCAL_DB):
        return LOCAL_BASE, LOCAL_DB, LOCAL_BACKUP

    # Both Missing - Prompt User
    app = QApplication.instance()
    if not app:
        app = QApplication(sys.argv)

    msg = QMessageBox()
    msg.setIcon(QMessageBox.Icon.Question)
    msg.setWindowTitle("Database Not Found")
    msg.setText("No database file was found.")
    msg.setInformativeText(f"checked:\n1. {TARGET_DB}\n2. {LOCAL_DB}\n\nWould you like to Create a new database or Import an existing one?")
    
    btn_import = msg.addButton("📁 Import Database", QMessageBox.ButtonRole.ActionRole)
    btn_create = msg.addButton("🆕 Create New (Local)", QMessageBox.ButtonRole.ActionRole)
    btn_exit = msg.addButton("❌ Exit App", QMessageBox.ButtonRole.RejectRole)
    
    msg.exec()
    clicked = msg.clickedButton()

    if clicked == btn_exit:
        sys.exit(0)

    # Ensure directories exist (Local by default for safety if creating new)
    FINAL_BASE = LOCAL_BASE
    FINAL_DB = LOCAL_DB
    FINAL_BACKUP = LOCAL_BACKUP

    # If volume exists, try to use it for creation? 
    # User requested robustness, suggesting local is safer if volume missing.
    if os.path.exists(VOL_PATH):
        FINAL_BASE = TARGET_BASE
        FINAL_DB = TARGET_DB
        FINAL_BACKUP = TARGET_BACKUP
        os.makedirs(FINAL_BASE, exist_ok=True)
        os.makedirs(FINAL_BACKUP, exist_ok=True)
    else:
        os.makedirs(FINAL_BACKUP, exist_ok=True)

    if clicked == btn_import:
        f_path, _ = QFileDialog.getOpenFileName(None, "Select Database to Import", "", "SQLite DB (*.db);;All Files (*)")
        if f_path:
            try:
                shutil.copy2(f_path, FINAL_DB)
                QMessageBox.information(None, "Success", f"Database imported successfully to:\n{FINAL_DB}")
                return FINAL_BASE, FINAL_DB, FINAL_BACKUP
            except Exception as e:
                QMessageBox.critical(None, "Import Error", f"Failed to copy database:\n{e}")
                sys.exit(1)
        else:
            sys.exit(0)
            
    elif clicked == btn_create:
        try:
            with open(FINAL_DB, 'w') as f:
                pass
            QMessageBox.information(None, "Success", f"New database created at:\n{FINAL_DB}")
            return FINAL_BASE, FINAL_DB, FINAL_BACKUP
        except Exception as e:
            QMessageBox.critical(None, "Error", f"Failed to create database:\n{e}")
            sys.exit(1)

    return FINAL_BASE, FINAL_DB, FINAL_BACKUP

# --- INITIALIZE PATHS ---
BASE_ARCHIVE_PATH, DB_FILE, DEFAULT_BACKUP_DIR = resolve_paths()

# ENABLE WAL MODE (Robustness)
try:
    _conn = sqlite3.connect(DB_FILE)
    _conn.execute("PRAGMA journal_mode=WAL;")
    _conn.close()
except:
    pass

LEGACY_ARCHIVE_PATH = os.path.expanduser("~/Desktop/Archive_Data")
DROPDOWN_FILE = os.path.join(BASE_ARCHIVE_PATH, "DropdownData.xlsx")

FONT_FILE_NAME = "AlFatemi14241.ttf"
SPLASH_FILE_NAME = "splash_background.png"

# Resolve Font Path (Check Bundle -> Documents -> Desktop Legacy)
# We use BASE_ARCHIVE_PATH which is now dynamically resolved
_font_candidates = [
    resource_path(FONT_FILE_NAME),
    os.path.join(BASE_ARCHIVE_PATH, FONT_FILE_NAME),
    os.path.join(LEGACY_ARCHIVE_PATH, FONT_FILE_NAME)
]
FONT_PATH = _font_candidates[0]
for p in _font_candidates:
    if os.path.exists(p):
        FONT_PATH = p
        break

# Resolve Splash Path (Check Bundle -> Documents -> Desktop Legacy)
_splash_candidates = [
    resource_path(SPLASH_FILE_NAME),
    os.path.join(BASE_ARCHIVE_PATH, SPLASH_FILE_NAME),
    os.path.join(LEGACY_ARCHIVE_PATH, SPLASH_FILE_NAME)
]
SPLASH_PATH = _splash_candidates[0]
for p in _splash_candidates:
    if os.path.exists(p):
        SPLASH_PATH = p
        break

# Custom Role for tracking original values in Master Lists
OriginalValueRole = Qt.ItemDataRole.UserRole + 1


# --- THEMES ---
STYLESHEET_LIGHT = """
    QMainWindow { background-color: #f7fafc; color: #2d3748; }
    
    /* HYBRID THEME: Sidebar is Dark Navy to match Splash */
    QFrame#Sidebar { background-color: #0F172A; border-right: 1px solid #1e293b; }
    
    QLabel { color: #2d3748; }
    QLabel#SidebarTitle { color: #D69E2E; } /* Golden Title */
    
    /* Standard Buttons */
    QPushButton { background-color: transparent; border: none; padding: 6px; border-radius: 4px; color: #4a5568; }
    QPushButton:hover { background-color: #f0f4f8; }
    
    /* PRIMARY ACTION: Cyan (from Splash Logo) */
    QPushButton#PrimaryBtn { background-color: #38B2AC; color: white; font-weight: bold; border: none; }
    QPushButton#PrimaryBtn:hover { background-color: #319795; }
    
    QPushButton#SecondaryBtn { background-color: #edf2f7; color: #2d3748; border: 1px solid #dcdfe3; font-weight: bold; padding: 5px 15px; }
    QPushButton#SecondaryBtn:hover { background-color: #e2e8f0; }
    
    /* Toggle Colors - Locked is Cyan, Unlocked is Orange */
    QPushButton#LockBtn { background-color: #fbd38d; color: #744210; border: 1px solid #d69e2e; font-weight: bold; padding: 4px 12px; }
    QPushButton#LockBtn:checked { background-color: #38B2AC; color: white; border-color: #319795; }

    /* SIDEBAR MENU BUTTONS (Dark Background logic) */
    QPushButton#MenuBtn { text-align: left; padding-left: 15px; font-size: 14px; color: #cbd5e0; }
    QPushButton#MenuBtn:hover { background-color: #1e293b; color: white; }
    QPushButton#MenuBtn:checked { background-color: #38B2AC; color: white; font-weight: bold; }
    
    QPushButton#CollapseBtn { font-size: 20px; color: #a0aec0; background: transparent; }
    QPushButton#CollapseBtn:hover { color: white; }

    /* MAIN CONTENT AREA (Light) */
    QTableWidget { background-color: white; color: #2d3748; gridline-color: #edf2f7; border: 1px solid #dcdfe3; outline: none; }
    
    QTableWidget::item:selected { 
        background-color: #e6fffa !important; /* Very light cyan tint */
        color: #2d3748 !important; 
    }
    
    QTableWidget::item:focus { border: none; outline: none; }
    
    QHeaderView::section { background-color: #f7fafc; padding: 8px; border: 1px solid #dcdfe3; font-weight: bold; color: #4a5568; text-align: center; }
    
    QTableCornerButton::section {
        background-color: #f7fafc;
        border: 1px solid #dcdfe3;
    }

    QLineEdit, QComboBox, QTextEdit { background-color: white !important; border: 1px solid #dcdfe3; color: #2d3748 !important; border-radius: 4px; padding: 4px; }
    
    QComboBox QAbstractItemView { 
        background-color: white !important; 
        color: #2d3748 !important; 
        selection-background-color: #e6fffa; 
        selection-color: #2d3748 !important; 
        border: 1px solid #dcdfe3;
    }
    
    QTabWidget::pane { border: 1px solid #dcdfe3; background: white; top: -1px; }
    QTabBar::tab { background: #edf2f7; border: 1px solid #dcdfe3; padding: 12px 30px; margin-right: 2px; color: #4a5568; border-top-left-radius: 4px; border-top-right-radius: 4px; }
    QTabBar::tab:selected { background: white; border-bottom-color: white; font-weight: bold; color: #38B2AC; }
    
    QGroupBox { 
        font-weight: bold; 
        border: 1px solid #dcdfe3; 
        border-radius: 6px; 
        margin-top: 10px; 
        padding-top: 25px; 
        background-color: white; 
        color: #2d3748; 
    }
    QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; left: 10px; padding: 0 5px; }
    
    QTextEdit#DebugConsole { font-size: 10px; background: #fdfdfd; border: 1px solid #dcdfe3; color: #2d3748; }

    /* --- Slider Styling for Light Mode --- */
    QSlider::groove:horizontal {
        border: 1px solid #cbd5e0;
        height: 8px;
        background: #e2e8f0;
        margin: 2px 0;
        border-radius: 4px;
    }
    QSlider::handle:horizontal {
        background: #38B2AC;
        border: 1px solid #38B2AC;
        width: 18px;
        height: 18px;
        margin: -5px 0;
        border-radius: 9px;
    }
    
    QScrollArea { border: none; background-color: transparent; }
    QScrollArea > QWidget > QWidget { background-color: transparent; }

    /* Dialogs & MessageBoxes (Fix for macOS Light Mode) */
    QDialog, QMessageBox { background-color: white; color: #2d3748; }
    QMessageBox QLabel { color: #2d3748; background-color: transparent; }
"""

STYLESHEET_DARK = """
    QMainWindow { background-color: #0F172A; color: #e2e8f0; } /* Dark Navy Main BG */
    QFrame#Sidebar { background-color: #0F172A; border-right: 1px solid #1e293b; }
    QLabel { color: #e2e8f0; }
    QLabel#SidebarTitle { color: #D69E2E; } /* Golden Title */
    
    QPushButton { background-color: transparent; border: none; padding: 6px; border-radius: 4px; color: #a0aec0; }
    QPushButton:hover { background-color: #1e293b; color: #e2e8f0; }
    
    /* Primary Action: Cyan */
    QPushButton#PrimaryBtn { background-color: #38B2AC; color: white; font-weight: bold; border: none; }
    QPushButton#PrimaryBtn:hover { background-color: #319795; }
    
    QPushButton#SecondaryBtn { background-color: #1e293b; color: #e2e8f0; border: 1px solid #334155; font-weight: bold; padding: 5px 15px; }
    QPushButton#SecondaryBtn:hover { background-color: #334155; }
    
    QPushButton#LockBtn { background-color: #744210; color: #fefcbf; border: 1px solid #ecc94b; font-weight: bold; padding: 4px 12px; }
    QPushButton#LockBtn:checked { background-color: #319795; color: white; border-color: #285e61; }
    
    /* Primary Action: Cyan */
    QPushButton#PrimaryBtn { background-color: #38B2AC; color: white; font-weight: bold; border: none; }
    QPushButton#PrimaryBtn:hover { background-color: #319795; }
    
    QPushButton#SecondaryBtn { background-color: #2d3748; color: #e2e8f0; border: 1px solid #4a5568; font-weight: bold; padding: 5px 15px; }
    QPushButton#SecondaryBtn:hover { background-color: #4a5568; }
    
    QPushButton#LockBtn { background-color: #744210; color: #fbd38d; border: 1px solid #d69e2e; font-weight: bold; padding: 4px 12px; }
    QPushButton#LockBtn:checked { background-color: #38B2AC; color: white; border-color: #319795; }

    /* Sidebar Menu */
    QPushButton#MenuBtn { text-align: left; padding-left: 15px; font-size: 14px; color: #cbd5e0; }
    QPushButton#MenuBtn:hover { background-color: #1e293b; color: white; }
    QPushButton#MenuBtn:checked { background-color: #38B2AC; color: white; font-weight: bold; }
    
    QPushButton#CollapseBtn { font-size: 20px; color: #a0aec0; background: transparent; }
    QPushButton#CollapseBtn:hover { color: white; }

    QTableWidget { 
        background-color: #1e293b; 
        color: #e2e8f0; 
        gridline-color: #2d3748; 
        border: 1px solid #4a5568; 
        outline: none;
        selection-color: white; /* FORCE WHITE SELECTION TEXT */
        selection-background-color: #2d3748; /* FORCE DARK SELECTION BG */
    }
    
    QTableWidget::item:selected { 
        background-color: #2d3748 !important; 
        color: white !important; /* REDUNDANT FORCE */
    }
    
    QTableWidget::item:focus { border: none; outline: none; }
    
    QHeaderView::section { background-color: #0F172A; padding: 8px; border: 1px solid #4a5568; font-weight: bold; color: #cbd5e0; text-align: center; }
    
    QTableCornerButton::section {
        background-color: #0F172A;
        border: 1px solid #4a5568;
    }

    QLineEdit, QComboBox, QTextEdit { background-color: #1e293b !important; border: 1px solid #4a5568; color: #e2e8f0 !important; border-radius: 4px; padding: 4px; }
    
    QComboBox QAbstractItemView { 
        background-color: #1e293b !important; 
        color: #e2e8f0 !important; 
        selection-background-color: #2d3748; 
        selection-color: #e2e8f0 !important; 
        border: 1px solid #4a5568;
    }
    
    QTabWidget::pane { border: 1px solid #4a5568; background: #1e293b; top: -1px; }
    QTabBar::tab { background: #0F172A; border: 1px solid #4a5568; padding: 12px 30px; margin-right: 2px; color: #a0aec0; border-top-left-radius: 4px; border-top-right-radius: 4px; }
    QTabBar::tab:selected { background: #1e293b; border-bottom-color: #1e293b; font-weight: bold; color: #38B2AC; }
    
    QGroupBox { 
        font-weight: bold; 
        border: 1px solid #4a5568; 
        border-radius: 6px; 
        margin-top: 10px; 
        padding-top: 25px; 
        background-color: #1e293b; 
        color: #e2e8f0; 
    }
    QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; left: 10px; padding: 0 5px; }

    QTextEdit#DebugConsole { font-size: 10px; background: #0F172A; border: 1px solid #4a5568; color: #a0aec0; }

    QSlider::groove:horizontal {
        border: 1px solid #4a5568;
        height: 8px;
        background: #2d3748;
        margin: 2px 0;
        border-radius: 4px;
    }
    QSlider::handle:horizontal {
        background: #38B2AC;
        border: 1px solid #38B2AC;
        width: 18px;
        height: 18px;
        margin: -5px 0;
        border-radius: 9px;
    }
    
    QScrollArea { border: none; background-color: transparent; }
    QScrollArea > QWidget > QWidget { background-color: transparent; }
    
    QDialog, QMessageBox { background-color: #1e293b; color: #e2e8f0; }
    QMessageBox QLabel { color: #e2e8f0; background-color: transparent; }
"""

# --- 2. CALENDAR CALCULATOR LOGIC ---

def misri_to_julian(hd, hm, hy):
    """Converts Misri Hijri date to Julian Day number."""
    hy_val = hy - 1
    jd = 1948439 + (hy_val // 30) * 10631
    hy_rem = hy_val % 30
    leap_years = [2, 5, 7, 10, 13, 16, 18, 21, 24, 26, 29]
    for i in range(1, hy_rem + 1):
        if i in leap_years:
            jd = jd + 355
        else:
            jd = jd + 354
    for m in range(1, hm):
        if m % 2 != 0:
            jd = jd + 30
        else:
            jd = jd + 29
    jd = jd + hd - 1
    return jd

def julian_to_gregorian(jd):
    """Converts Julian Day number to a Gregorian date."""
    L = jd + 68569
    N = (4 * L) // 146097
    L = L - (146097 * N + 3) // 4
    I = (4000 * (L + 1)) // 1461001
    L = L - (1461 * I) // 4 + 31
    J = (80 * L) // 2447
    day = L - (2447 * J) // 80
    L = J // 11
    month = J + 2 - (12 * L)
    year = 100 * (N - 49) + I + L
    return date(year, month, day)

def gregorian_to_julian(year, month, day):
    """Converts Gregorian date to Julian Day number."""
    if month <= 2:
        year = year - 1
        month = month + 12
    A = year // 100
    B = 2 - A + (A // 4)
    jd = int(365.25 * (year + 4716)) + int(30.6001 * (month + 1)) + day + B - 1524.5
    return int(jd + 0.5)

def julian_to_misri(jd):
    """Converts Julian Day number back to Misri Hijri."""
    jd_rel = jd - 1948439
    cycle = jd_rel // 10631
    jd_rel = jd_rel % 10631
    hy = cycle * 30 + 1
    leap_years = [2, 5, 7, 10, 13, 16, 18, 21, 24, 26, 29]
    while True:
        year_len = 354
        if (hy % 30) in leap_years:
            year_len = 355
        if jd_rel < year_len:
            break
        jd_rel = jd_rel - year_len
        hy = hy + 1
    hm = 1
    while True:
        month_len = 29
        if hm % 2 != 0:
            month_len = 30
        if hm == 12 and (hy % 30) in leap_years:
            month_len = 30
        if jd_rel < month_len:
            break
        jd_rel = jd_rel - month_len
        hm = hm + 1
    hd = jd_rel + 1
    return hd, hm, hy

def calculate_esavi(hijri_str):
    try:
        parts = hijri_str.split('/')
        day = int(parts[0])
        month = int(parts[1])
        year = int(parts[2])
        jd = misri_to_julian(day, month, year)
        dt = julian_to_gregorian(jd)
        formatted_date = dt.strftime("%d/%m/%Y")
        return formatted_date
    except:
        return ""

def calculate_hijri(esavi_str):
    try:
        parts = esavi_str.split('/')
        day = int(parts[0])
        month = int(parts[1])
        year = int(parts[2])
        jd = gregorian_to_julian(year, month, day)
        hd, hm, hy = julian_to_misri(jd)
        formatted_hijri = f"{hd:02d}/{hm:02d}/{hy}"
        return formatted_hijri
    except:
        return ""

# --- 3. BACKGROUND WORKER FOR BACKUPS ---

class BackupWorker(QThread):
    finished = pyqtSignal(str) # Emits status message

    def __init__(self, target_dir, prefix="archive"):
        super().__init__()
        self.target_dir = target_dir
        self.prefix = prefix

    def run(self):
        try:
            if not os.path.exists(DB_FILE):
                return

            if not os.path.exists(self.target_dir):
                os.makedirs(self.target_dir)

            now = datetime.now()
            year_folder = os.path.join(self.target_dir, str(now.year))
            if not os.path.exists(year_folder):
                os.makedirs(year_folder)

            # 1. Create Backup
            filename = f"{self.prefix}_{now.strftime('%Y-%m-%d_%H-%M-%S')}.db"
            dest_path = os.path.join(year_folder, filename)
            shutil.copy2(DB_FILE, dest_path)
            
            # 2. Smart Consolidation Logic
            self.consolidate_backups(year_folder)
            
            self.finished.emit(f"Backup created: {filename}")
        except Exception as e:
            self.finished.emit(f"Backup Failed: {str(e)}")

    def consolidate_backups(self, folder_path):
        """Keeps only the last backup of past days, and last backup of past weeks."""
        files = []
        for f in os.listdir(folder_path):
            if f.endswith(".db") and f.startswith("archive_"):
                full_path = os.path.join(folder_path, f)
                files.append(full_path)
        
        files.sort() # Sort by name (which includes timestamp) => oldest first
        
        if not files:
            return

        # Parse file info
        file_info = []
        for fp in files:
            try:
                # Extract date from filename: archive_2024-01-20_14-30-00.db
                basename = os.path.basename(fp)
                date_part = basename.split('_')[1] # 2024-01-20
                file_date = datetime.strptime(date_part, "%Y-%m-%d").date()
                file_info.append({'path': fp, 'date': file_date})
            except:
                continue

        today = date.today()
        
        # Group by date
        files_by_date = {}
        for item in file_info:
            d = item['date']
            if d not in files_by_date:
                files_by_date[d] = []
            files_by_date[d].append(item['path'])

        # RULE 1: Daily Consolidation (Past Days)
        # For every date EXCEPT today, keep only the last file
        for d, paths in files_by_date.items():
            if d < today:
                # This is a past day. If multiple files, keep only the last one.
                # paths are sorted because original list was sorted
                while len(paths) > 1:
                    file_to_remove = paths.pop(0) # Remove first (oldest)
                    try:
                        os.remove(file_to_remove)
                    except:
                        pass

        # Re-scan remaining files for Weekly Rule
        remaining_files = []
        for f in os.listdir(folder_path):
            if f.endswith(".db"):
                full_path = os.path.join(folder_path, f)
                remaining_files.append(full_path)
        remaining_files.sort()

        # Parse again
        weekly_groups = {}
        for fp in remaining_files:
            try:
                basename = os.path.basename(fp)
                date_part = basename.split('_')[1]
                file_date = datetime.strptime(date_part, "%Y-%m-%d").date()
                
                # Check age > 7 days
                age = (today - file_date).days
                if age > 7:
                    # Get Year-Week key (e.g., (2024, 3))
                    # isocalendar returns (year, week, weekday)
                    year_week = file_date.isocalendar()[:2]
                    if year_week not in weekly_groups:
                        weekly_groups[year_week] = []
                    weekly_groups[year_week].append(fp)
            except:
                pass

        # RULE 2: Weekly Consolidation (Older than 7 days)
        # For older weeks, keep only the very last file of that week
        for yw, paths in weekly_groups.items():
            while len(paths) > 1:
                file_to_remove = paths.pop(0) # Remove oldest
                try:
                    os.remove(file_to_remove)
                except:
                    pass

# --- 4. CUSTOM UI COMPONENTS ---

class LoadWorker(QThread):
    """
    Background worker for loading data pages without freezing UI.
    Returns: (rows, total_count, column_widths_if_saved)
    """
    finished = pyqtSignal(list, int, float) # rows, total_count, time_taken
    
    def __init__(self, db_path, query_sql, query_params, limit, offset, count_sql=None, count_params=None):
        super().__init__()
        self.db_path = db_path
        self.sql = query_sql
        self.params = query_params
        self.limit = limit
        self.offset = offset
        self.count_sql = count_sql
        self.count_params = count_params
        
    def run(self):
        t0 = time.time()
        conn = None
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # 1. Get Total Count (if needed for pagination)
            total = 0
            if self.count_sql:
                cursor.execute(self.count_sql, self.count_params)
                res = cursor.fetchone()
                if res: total = res[0]
            
            # 2. Get Page Data
            # Inject Limit/Offset:
            # We assume the SQL passed has placeholders or is robust. 
            # Actually, safe way is to append LIMIT OFFSET to the string passed in, 
            # assuming the caller prepared the base query correctly without them.
            if "LIMIT" not in self.sql.upper():
                final_sql = f"{self.sql} LIMIT {self.limit} OFFSET {self.offset}"
            else:
                final_sql = self.sql # Calling code handled it?
                
            cursor.execute(final_sql, self.params)
            rows = cursor.fetchall()
            
            t1 = time.time()
            self.finished.emit(rows, total, t1-t0)
            
        except Exception as e:
            traceback.print_exc()
            self.finished.emit([], 0, 0.0)
        finally:
            if conn:
                conn.close()

class ImportWorker(QThread):
    progress_update = pyqtSignal(str, int) # Message, %
    finished_success = pyqtSignal()
    finished_error = pyqtSignal(str)

    def __init__(self, excel_path, db_file, selected_sheets):
        super().__init__()
        self.excel_path = excel_path
        self.db_file = db_file
        self.selected_sheets = [s.strip() for s in selected_sheets]
    
    def run(self):
        conn = None
        try:
            print("DEBUG: Starting ImportWorker.run")
            self.progress_update.emit("Reading Excel File...", 5)
            xls = pd.ExcelFile(self.excel_path)
            print(f"DEBUG: All sheets in file: {xls.sheet_names}")
            print(f"DEBUG: Selected sheets: {self.selected_sheets}")

            # Normalize sheet names from file
            file_sheet_map = {s.strip().lower(): s.strip() for s in xls.sheet_names}
            
            conn = sqlite3.connect(self.db_file)
            cursor = conn.cursor()
            cursor.execute("BEGIN TRANSACTION")
            
            # Track handled sheets
            handled_sheets = set()
            
            # Helper to check if a sheet was selected (case-insensitive)
            def is_selected(target_name):
                for s in self.selected_sheets:
                    if s.lower() == target_name.lower():
                        return True
                return False

            total_steps = len(self.selected_sheets)
            step_progress: float = 90.0 / max(total_steps, 1) if total_steps > 0 else 0.0
            current_progress: float = 10.0

            # --- 1. EVENTS (Master) ---
            if is_selected("Master"):
                target_sheet = file_sheet_map.get("master")
                if target_sheet:
                    print(f"DEBUG: Processing Master -> {target_sheet}")
                    handled_sheets.add(target_sheet.lower())
                    self.progress_update.emit("Importing Master Sheet...", int(current_progress))
                    cursor.execute("DELETE FROM events")
                    
                    df = pd.read_excel(xls, target_sheet)
                    df.columns = [c.strip() for c in df.columns]
                    
                    col_map = {
                        "Audio No": "audio_no", "Person": "person", "Occasion": "occasion", 
                        "Category": "category", "Place": "place", "Country": "country",
                        "Hijri Date": "hijri_date", "Esavi Date": "esavi_date", "Year": "year",
                        "Out of": "out_of", "Remarks": "remarks",
                        "Tracks": "Tracks", "Track": "Tracks", "Track List": "Tracks",
                        "A/V": "AV", "AV": "AV", "Cass No": "cass_no", "Came From": "came_from", "Incomplete": "incomplete"
                    }
                    
                    # Prepare data
                    total_rows = len(df)
                    db_data = []
                    
                    for idx, row in df.iterrows():
                        if idx % 500 == 0:
                            pct = int(current_progress + (idx / total_rows) * (step_progress * 0.8))
                            self.progress_update.emit(f"Importing Master Sheet ({idx}/{total_rows})...", pct)

                        vals = {}
                        for head, db_col in col_map.items():
                            if head in df.columns:
                                val = row[head]
                                if pd.isna(val):
                                    final_val = ""
                                elif isinstance(val, (pd.Timestamp, datetime, date)):
                                    final_val = val.strftime("%d/%m/%Y")
                                elif isinstance(val, float):
                                    if val.is_integer():
                                        final_val = str(int(val))
                                    else:
                                        final_val = str(val)
                                else:
                                    final_val = str(val).strip()
                                vals[db_col] = final_val
                        
                        raw_no = str(vals.get("audio_no", "")).strip()
                        if not raw_no: continue
                        try:
                            if raw_no.endswith(".0"): raw_no = raw_no[:-2]
                            an = int(raw_no)
                        except: continue

                        tup = (
                            an,
                            vals.get("person", ""), vals.get("occasion", ""), vals.get("category", ""),
                            vals.get("place", ""), vals.get("country", ""), vals.get("hijri_date", ""),
                            vals.get("esavi_date", ""), vals.get("year", ""), vals.get("out_of", ""),
                            vals.get("remarks", ""), vals.get("Tracks", ""), vals.get("AV", ""),
                            vals.get("cass_no", ""), vals.get("came_from", ""), vals.get("incomplete", "")
                        )
                        db_data.append(tup)
                    
                    if db_data:
                        cursor.executemany("""
                            INSERT OR REPLACE INTO events (
                                audio_no, person, occasion, category, place, country, 
                                hijri_date, esavi_date, year, out_of, remarks, 
                                Tracks, AV, cass_no, came_from, incomplete
                            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                        """, db_data)
                
                current_progress += step_progress

            # --- 2. LOCATIONS ---
            if is_selected("Locations") or is_selected("Location"):
                target_sheet = file_sheet_map.get("locations")
                if not target_sheet: target_sheet = file_sheet_map.get("location")
                if target_sheet:
                    print(f"DEBUG: Processing Locations -> {target_sheet}")
                    handled_sheets.add(target_sheet.lower())
                    self.progress_update.emit("Importing Locations...", int(current_progress))
                    cursor.execute("DELETE FROM locations")
                    df = pd.read_excel(xls, target_sheet)
                    df.columns = [c.strip() for c in df.columns]
                    
                    print(f"DEBUG: Locations Columns: {list(df.columns)}")
                    
                    # Try to find the place column (prefer exact match or contains)
                    p_col = None
                    for c in df.columns:
                        if c.strip().lower() in ["place", "location", "city"]:
                            p_col = c
                            break
                    if not p_col:
                        p_col = next((c for c in df.columns if "Place" in c or "City" in c), None)
                    
                    if not p_col and len(df.columns) > 0: p_col = df.columns[0]
                    
                    # Try to find the country column
                    c_col = None
                    for c in df.columns:
                        if c.strip().lower() == "country":
                            c_col = c
                            break
                    if not c_col:
                        c_col = next((c for c in df.columns if "Country" in c), None)
                        
                    if not c_col and len(df.columns) > 1:
                        # If p_col is df.columns[0], try df.columns[1] for country
                        if p_col == df.columns[0]:
                             c_col = df.columns[1]
                    
                    print(f"DEBUG: Selected Place Column: '{p_col}', Selected Country Column: '{c_col}'")
                    if len(df) > 0:
                         print(f"DEBUG: First row data: {df.iloc[0].to_dict()}")

                    if p_col:
                        db_locs = []
                        valid_places = set()
                        for _, row in df.iterrows():
                            p_val = str(row[p_col]).strip()
                            if not p_val or p_val.lower() == "nan": continue
                            if p_val in valid_places: continue 
                            
                            c_val = ""
                            if c_col:
                                raw_c = str(row[c_col]).strip()
                                if raw_c and raw_c.lower() != "nan": c_val = raw_c
                                    
                            db_locs.append((p_val, c_val))
                            valid_places.add(p_val)
                            
                        if db_locs:
                            cursor.executemany("INSERT INTO locations (place, country) VALUES (?,?)", db_locs)
                
                current_progress += step_progress

            # --- 3. TRACKS ---
            if is_selected("Tracks"):
                target_sheet = file_sheet_map.get("tracks")
                if target_sheet:
                    print(f"DEBUG: Processing Tracks -> {target_sheet}")
                    handled_sheets.add(target_sheet.lower())
                    self.progress_update.emit("Importing Tracks...", int(current_progress))
                    cursor.execute("DELETE FROM tracks")
                    df = pd.read_excel(xls, target_sheet)
                    df.columns = [str(c).strip() for c in df.columns]
                    
                    print(f"DEBUG: Tracks Columns: {list(df.columns)}")
                    
                    eid_col = next((c for c in df.columns if "Audio" in c and "No" in c), None)
                    if not eid_col:
                        eid_col = next((c for c in df.columns if "Id" in c), None)
                    
                    # Be specific about Track No - look for "Track" AND "No" together
                    tn_col = next((c for c in df.columns if "Track" in c and "No" in c), None)
                    if not tn_col:
                        tn_col = next((c for c in df.columns if c.strip().lower() == "track"), None)
                    
                    nm_col = next((c for c in df.columns if "Tracks" in c or "Name" in c or "Title" in c or "Description" in c), None)
                    
                    print(f"DEBUG: Selected Tracks EntryID Col: '{eid_col}', TrackNo Col: '{tn_col}', Name Col: '{nm_col}'")
                    if len(df) > 0:
                        print(f"DEBUG: Tracks First row data: {df.iloc[0].to_dict()}")

                    if eid_col:
                        db_tracks = []
                        for _, row in df.iterrows():
                            raw_eid = str(row[eid_col]).strip()
                            if not raw_eid or raw_eid.lower() == "nan": continue
                            try:
                                if raw_eid.endswith(".0"): raw_eid = raw_eid[:-2]
                                eid = int(raw_eid)
                            except: continue

                            raw_tn = ""
                            if tn_col:
                                raw_tn = str(row[tn_col]).strip()
                                if raw_tn.lower() == "nan": raw_tn = ""
                                if raw_tn.endswith(".0"): raw_tn = raw_tn[:-2]
                            
                            raw_nm = ""
                            if nm_col:
                                raw_nm = str(row[nm_col]).strip()
                                if raw_nm.lower() == "nan": raw_nm = ""
                                
                            db_tracks.append((eid, raw_tn, raw_nm))
                            
                        if db_tracks:
                            cursor.executemany("INSERT INTO tracks (event_id, track_no, track_name) VALUES (?,?,?)", db_tracks)

                current_progress += step_progress

            # --- 4. DROPDOWNS (Monolithic & Individual) ---
            s_drop_mono = is_selected("Dropdown Options")
            if not s_drop_mono: s_drop_mono = is_selected("Dropdowns")
            
            # List of individual sheets that should be mapped to dropdowns
            # "Location" is handled above. 
            individual_lists = ["Person", "Occasion", "Category", "AV", "AV", "Incomplete", "Created", "Came From"]
            
            # Check if ANY dropdown related sheet is selected
            s_drop_any = s_drop_mono
            if not s_drop_any:
                for ind in individual_lists:
                    if is_selected(ind):
                        s_drop_any = True
                        break

            if s_drop_any:
                self.progress_update.emit("Importing Master Lists...", int(current_progress))
                cursor.execute("DELETE FROM dropdown_options")
                
                # 4a. Monolithic Sheet
                if s_drop_mono:
                    target_sheet = file_sheet_map.get("dropdown options")
                    if not target_sheet: target_sheet = file_sheet_map.get("dropdowns")
                    if target_sheet:
                        print(f"DEBUG: Processing Dropdowns (Mono) -> {target_sheet}")
                        handled_sheets.add(target_sheet.lower())
                        try:
                            df = pd.read_excel(xls, target_sheet)
                            df.columns = [str(c).strip() for c in df.columns]
                            
                            cat_col = next((c for c in df.columns if "Cat" in c), None)
                            val_col = next((c for c in df.columns if "Val" in c or "Name" in c), None)
                            if not cat_col and len(df.columns) > 0: cat_col = df.columns[0]
                            if not val_col and len(df.columns) > 1: val_col = df.columns[1]
                            
                            if cat_col and val_col:
                                db_drops = []
                                for _, row in df.iterrows():
                                    cat = str(row[cat_col]).strip()
                                    val = str(row[val_col]).strip()
                                    if not cat or not val or cat.lower()=="nan" or val.lower()=="nan": continue
                                    db_drops.append((cat, val))
                                if db_drops:
                                    cursor.executemany("INSERT INTO dropdown_options (category, value) VALUES (?,?)", db_drops)
                        except Exception as e:
                            print(f"Error importing monolithic dropdowns: {e}")

                # 4b. Individual Sheets
                for ind_cat in individual_lists:
                    if is_selected(ind_cat):
                        target_sheet = file_sheet_map.get(ind_cat.lower())
                        if target_sheet:
                            print(f"DEBUG: Processing Individual List -> {target_sheet}")
                            handled_sheets.add(target_sheet.lower())
                            try:
                                df = pd.read_excel(xls, target_sheet)
                                df.columns = [str(c).strip() for c in df.columns]
                                
                                # Assume value is in first column unless "Name" or "Value" exists
                                val_col = next((c for c in df.columns if "Val" in c or "Name" in c), None)
                                if not val_col and len(df.columns) > 0: val_col = df.columns[0]
                                
                                print(f"DEBUG: Sheet '{target_sheet}' - Columns: {df.columns.tolist()}")
                                print(f"DEBUG: Selected Value Column: '{val_col}'")
                                print(f"DEBUG: Row Count: {len(df)}")
                                if len(df) > 0:
                                    print(f"DEBUG: First row data: {df.iloc[0].to_dict()}")

                                
                                if val_col:
                                    db_ind = []
                                    # Normalized Category Name (e.g. "AV" -> "AV")
                                    norm_cat = ind_cat
                                    if norm_cat.upper() in ["AV", "A/V", "AV"]: norm_cat = "AV"
                                    
                                    for _, row in df.iterrows():
                                        val = str(row[val_col]).strip()
                                        if not val or val.lower()=="nan": continue
                                        db_ind.append((norm_cat, val))
                                    
                                    if db_ind:
                                        cursor.executemany("INSERT INTO dropdown_options (category, value) VALUES (?,?)", db_ind)
                            except Exception as e:
                                print(f"Error importing individual list {ind_cat}: {e}")

                current_progress += step_progress

            # --- 5. FOLDER MAPPINGS (Info Sheets) ---
            # Support multiple possible names - more robust detection
            def is_folder_mapping_sheet(sheet_name):
                name_lower = sheet_name.lower()
                # Check for various patterns:
                # 1. Contains "folder" or "info" AND ("mapping" or "sheet" or "mp3")
                # 2. Contains "mp3" AND "mapping"
                if ("folder" in name_lower or "info" in name_lower):
                    if ("mapping" in name_lower or "sheet" in name_lower or "mp3" in name_lower):
                        return True
                # Also match "Mp3 Mappings" pattern
                if "mp3" in name_lower and "mapping" in name_lower:
                    return True
                return False
            
            folder_sheets = [s for s in self.selected_sheets if is_folder_mapping_sheet(s)]
            if folder_sheets:
                # Prioritize: "Folder Mappings" > "Info Sheet" > "Info Sheet MP3" or process ALL?
                # User might have multiple. Let's process valid ones.
                for s_name in folder_sheets:
                    target_sheet = file_sheet_map.get(s_name.lower())
                    if not target_sheet: continue
                    
                    print(f"DEBUG: Processing Folder Mapping -> {target_sheet}")
                    handled_sheets.add(target_sheet.lower())
                    self.progress_update.emit(f"Importing {target_sheet}...", int(current_progress))
                    
                    # Determine mapping type
                    m_type = "folder" # Default folder mapping
                    if "mp3" in s_name.lower():
                        m_type = "mp3"
                    
                    # Wipe existing mappings of this type??
                    # User request: "selective wipe".
                    # If importing mapping sheets, we should probably clear old ones?
                    # The previous logic did `cursor.execute("DELETE FROM folder_mappings")`.
                    # But now we might support multiple.
                    # Safety: Clear ONLY if it's the first time processing mappings in this run?
                    # Or specific types?
                    # Let's pivot to: DELETE based on type.
                    cursor.execute("DELETE FROM folder_mappings WHERE mapping_type=?", (m_type,))
                    
                    try:
                        df = pd.read_excel(xls, target_sheet)
                        df.columns = [str(c).strip() for c in df.columns]
                        
                        start_col = next((c for c in df.columns if "Start" in c), None)
                        lnk_col = next((c for c in df.columns if "Link" in c or "Location" in c or "Path" in c), None)

                        # Fallback: Assume Col 0 = Start, Col 1 = Link if 2 cols
                        if not start_col and len(df.columns) > 0: start_col = df.columns[0]
                        if not lnk_col and len(df.columns) > 1: lnk_col = df.columns[1]

                        if start_col and lnk_col:
                            db_maps = []
                            for _, row in df.iterrows():
                                try:
                                    s_raw = str(row[start_col]).strip()
                                    if not s_raw or s_raw.lower()=="nan": continue
                                    if s_raw.endswith(".0"): s_raw = s_raw[:-2]
                                    start_no = int(s_raw)
                                except: continue
                                
                                f_link = str(row[lnk_col]).strip()
                                if f_link.lower()=="nan": f_link = ""
                                if not f_link: continue
                                    
                                # No longer sorting drive/name in DB
                                db_maps.append((start_no, f_link, m_type))
                            
                            if db_maps:
                                cursor.executemany("INSERT OR REPLACE INTO folder_mappings (start_no, full_link, mapping_type) VALUES (?,?,?)", db_maps)
                    except Exception as e:
                        print(f"Error importing folder mapping {s_name}: {e}")

                current_progress += step_progress

            # --- 6. GENERIC / OLD SHEETS (Fallback) ---
            for s_sel in self.selected_sheets:
                real_name = file_sheet_map.get(s_sel.strip().lower())
                if real_name and real_name.lower() not in handled_sheets:
                    print(f"DEBUG: Importing Generic Sheet -> {real_name}")
                    self.progress_update.emit(f"Importing {real_name}...", int(current_progress))
                    
                    try:
                        df = pd.read_excel(xls, real_name)
                        cursor.execute("INSERT INTO old_sheets_meta (name) VALUES (?)", (real_name,))
                        sheet_id = cursor.lastrowid
                        
                        headers = [str(c) for c in df.columns]
                        cols_def = []
                        for idx, h in enumerate(headers):
                            cursor.execute("INSERT INTO old_sheet_headers (sheet_id, col_index, header_name) VALUES (?,?,?)", (sheet_id, idx, h))
                            cols_def.append(f"col_{idx} TEXT")
                        
                        t_name = f"old_sheet_data_{sheet_id}"
                        col_str = ", ".join(cols_def)
                        cursor.execute(f"CREATE TABLE IF NOT EXISTS {t_name} ({col_str})")
                        
                        q_marks = ",".join(["?"] * len(headers))
                        data_rows = []
                        for _, row in df.iterrows():
                            vals = [str(val) if pd.notna(val) else "" for val in row]
                            data_rows.append(vals)
                        if data_rows:
                            cursor.executemany(f"INSERT INTO {t_name} VALUES ({q_marks})", data_rows)
                    except Exception as e_gen:
                        print(f"DEBUG: Error importing generic sheet {real_name}: {e_gen}")
                        traceback.print_exc()
                    
                    current_progress += step_progress

            conn.commit()
            conn.close()
            self.finished_success.emit()
            
        except Exception as e:
            traceback.print_exc()
            if conn: conn.close()
            self.finished_error.emit(str(e))

# --- 5. CUSTOM UI COMPONENTS ---

class ToggleSwitch(QCheckBox):
    """Custom toggle switch widget that looks like a modern slider."""
    def __init__(self, width=50, height=26, parent=None):
        super().__init__(parent)
        self.setFixedSize(width, height)
        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def hitButton(self, pos):
        """Ensure the entire widget area is clickable."""
        return self.rect().contains(pos)

    def paintEvent(self, event):
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        rect = self.rect()
        # Track
        if self.isChecked():
            # Dark Mode Active (Gemini Blueish)
            brush = QColor("#a8c7fa") 
            thumb_color = QColor("#04080f")
        else:
            # Light Mode Active (Grayish)
            brush = QColor("#cbd5e0")
            thumb_color = QColor("white")

        p.setBrush(brush)
        p.setPen(Qt.PenStyle.NoPen)
        p.drawRoundedRect(0, 0, rect.width(), rect.height(), rect.height() / 2, rect.height() / 2)
        
        # Thumb
        thumb_x = rect.width() - rect.height() + 2 if self.isChecked() else 2
        thumb_rect = QRect(thumb_x, 2, rect.height() - 4, rect.height() - 4)
        p.setBrush(thumb_color)
        p.drawEllipse(thumb_rect)

class LisanTableItem(QTableWidgetItem):
    """Custom table item used to force RTL and Fatemi font."""
    def __init__(self, text, lisan_font_family="Arial", font_size=14, is_valid=True):
        super().__init__(str(text))
        self.is_valid_item = is_valid
        
        font = QFont(lisan_font_family, font_size)
        if not is_valid:
            self.setForeground(QColor("#e53e3e")) # Bright Red
            font.setBold(True)
        # We don't set default foreground here, EditorDelegate.paint handles theme-specific colors
            
        self.setFont(font)
        self.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self.setData(Qt.ItemDataRole.TextAlignmentRole, Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

class FolderMappingDialog(QDialog):
    """Dialog for adding folder mappings (2-column)."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add New Folder Mapping")
        self.setFixedWidth(500)
        
        self.main_layout = QVBoxLayout(self)
        self.form = QFormLayout()
        
        self.input_start_no = QLineEdit()
        self.input_start_no.setPlaceholderText("Starting Audio ID (e.g. 17501)")
        self.form.addRow("Start Number:", self.input_start_no)
        
        self.link_layout = QHBoxLayout()
        self.input_link = QLineEdit()
        self.input_link.setReadOnly(True)
        self.input_link.setPlaceholderText("Select folder using the button...")
        
        self.btn_browse = QPushButton("📁 Browse...")
        self.btn_browse.clicked.connect(self.browse_folder)
        
        self.link_layout.addWidget(self.input_link)
        self.link_layout.addWidget(self.btn_browse)
        self.form.addRow("File Location:", self.link_layout)
        
        self.main_layout.addLayout(self.form)
        
        self.button_box = QHBoxLayout()
        self.btn_save = QPushButton("Save Mapping")
        self.btn_save.clicked.connect(self.accept)
        self.btn_cancel = QPushButton("Cancel")
        self.btn_cancel.clicked.connect(self.reject)
        
        self.button_box.addStretch()
        self.button_box.addWidget(self.btn_cancel)
        self.button_box.addWidget(self.btn_save)
        self.main_layout.addLayout(self.button_box)

    def browse_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder", "/Volumes")
        if folder_path:
            file_url = f"file://{folder_path}/"
            self.input_link.setText(file_url)

    def get_data(self):
        return {
            "start_no": self.input_start_no.text().strip(),
            "link": self.input_link.text().strip()
        }

class CheckableComboBox(QComboBox):
    """
    A robust ComboBox allowing multiple selection via checkboxes.
    Fixed: Font support, Performance, Selection Reliability.
    """
    def __init__(self, placeholder_text="Select...", parent=None, font_family="Arial"):
        super().__init__(parent)
        self.placeholder_text = placeholder_text
        self.font_family = font_family
        
        self.setEditable(True)
        self.lineEdit().setReadOnly(True)
        self.lineEdit().setAlignment(Qt.AlignmentFlag.AlignLeft)
        
        # Performance: Use Standard Item Model expressly
        self.model_ = QStandardItemModel(self)
        self.setModel(self.model_)
        
        # Event Filter to keep popup open
        self.view().viewport().installEventFilter(self)
        
        # Handle interaction via Pressed signal (fires on MouseDown)
        # We manually toggle and block standard behavior if needed.
        self.view().pressed.connect(self.on_item_pressed)
        
        # Font
        font = QFont(self.font_family, 13)
        self.setFont(font)
        self.view().setFont(font)
        
        # Initial State
        self.update_display_text()

    def eventFilter(self, widget, event):
        # Prevent popup closing when clicking inside view on Release
        if widget == self.view().viewport():
            if event.type() == QEvent.Type.MouseButtonRelease:
                index = self.view().indexAt(event.pos())
                if index.isValid():
                    return True # Consume event, don't close
        return super().eventFilter(widget, event)
        
    def addItems(self, items):
        for text in items:
            item = QStandardItem(str(text))
            # Enable Checkable, Enabled. DISABLE Selectable to prevent double-visuals/issues.
            # Using only UserCheckable + Enabled.
            item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
            item.setCheckState(Qt.CheckState.Unchecked)
            self.model_.appendRow(item)
        self.update_display_text()

    def on_item_pressed(self, index):
        # Manually toggle logic.
        item = self.model_.itemFromIndex(index)
        if item.checkState() == Qt.CheckState.Checked:
            item.setCheckState(Qt.CheckState.Unchecked)
        else:
            item.setCheckState(Qt.CheckState.Checked)
        self.update_display_text()

        # NOTE: Since we consume MouseRelease, we rely on this Press to do the work.
        
    def update_display_text(self):
        selected_items = self.get_checked_items()
        if not selected_items:
            self.lineEdit().setText(self.placeholder_text)
            self.lineEdit().setStyleSheet("color: gray;")
        else:
            if len(selected_items) == 1:
                self.lineEdit().setText(selected_items[0])
            else:
                self.lineEdit().setText(f"{len(selected_items)} Selected")
            self.lineEdit().setStyleSheet("color: black;")

    def get_checked_items(self):
        checked_items = []
        for i in range(self.model_.rowCount()):
            item = self.model_.item(i, 0)
            if item.checkState() == Qt.CheckState.Checked:
                checked_items.append(item.text())
        return checked_items

    def clear_selection(self):
        for i in range(self.model_.rowCount()):
            item = self.model_.item(i, 0)
            item.setCheckState(Qt.CheckState.Unchecked)
        self.update_display_text()
        
    def showPopup(self):
        super().showPopup()
        
    def hidePopup(self):
        super().hidePopup()
        self.update_display_text()

class FilterComboBox(QComboBox):

    def __init__(self, label_text, parent=None, font_family="Arial"):
        super().__init__(parent)
        self.label_text = label_text
        self.setEditable(True)
        self.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        
        self._is_validating = False
        
        # Set Fatemi Font
        self.fatemi_font = QFont(font_family, 13)
        self.setFont(self.fatemi_font)
        self.lineEdit().setFont(self.fatemi_font)
        
        self.setStyleSheet("border: 1px solid #dcdfe3; border-radius: 4px;")
        
        # Configure search-as-you-type engine
        self.completer().setFilterMode(Qt.MatchFlag.MatchContains)
        self.completer().setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.completer().setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
        
        popup = self.completer().popup()
        if popup:
            popup.setFont(self.fatemi_font)
        
        self.lineEdit().textChanged.connect(self.trigger_suggest)
        self.lineEdit().editingFinished.connect(self.validate_content)

    def trigger_suggest(self, text):
        """Triggers the completer popup quietly without stealing cursor focus."""
        if text:
            if self.hasFocus():
                self.completer().complete()

    def showPopup(self):
        """Smart Width Logic: Expands dropdown list based on longest content."""
        fm = QFontMetrics(self.fatemi_font)
        max_content_w = 0
        for i in range(self.count()):
            item_text = self.itemText(i)
            w = fm.horizontalAdvance(item_text) + 60
            if w > max_content_w:
                max_content_w = w
        
        final_w = max(max_content_w, self.width())
        self.view().setMinimumWidth(final_w)
        super().showPopup()

    def validate_content(self):
        if self._is_validating:
            return
            
        text = self.currentText().strip()
        if not text:
            return
            
        found = False
        for i in range(self.count()):
            if self.itemText(i) == text:
                found = True
                break
        
        if not found:
            self._is_validating = True
            self.blockSignals(True)
            QMessageBox.warning(self, "Invalid Selection", f"'{text}' is not a valid {self.label_text}.")
            self.setCurrentIndex(-1)
            self.setEditText("")
            self.blockSignals(False)
            self._is_validating = False

class EditorDelegate(QStyledItemDelegate):
    def initStyleOption(self, option, index):
        super().initStyleOption(option, index)
        # Force Wrap for Summary/Remarks
        if index.column() == 11: 
            option.features |= QStyleOptionViewItem.ViewItemFeature.WrapText

    def __init__(self, font_family, mode="text", parent=None, master_lists=None):
        super().__init__(parent)
        self.font_family = font_family
        self.mode = mode
        self.parent_table = parent
        self.master_lists = master_lists or {}

    def sizeHint(self, option, index):
        """Calculate exact height needed for wrapped text."""
        # Check for Tracks (Col 3 in Tracks Sheet) or Description/Remarks (Col 11 in Master)
        should_wrap = False
        if index.column() == 11: should_wrap = True
        
        # Check for Tracks Sheet (Col 3) or Track Search (Col 2 has track name)
        if hasattr(self, 'parent_table'):
            tbl_name = self.parent_table.property("table_name")
            if tbl_name == "tracks" and index.column() == 3: should_wrap = True
            
        if not should_wrap:
            return super().sizeHint(option, index)

        text = str(index.data(Qt.ItemDataRole.DisplayRole) or "")
        # Remove extra newlines for calculation to avoid huge rows
        text = text.strip() 
        
        if not text:
            return super().sizeHint(option, index)
            
        # Get width from parent table
        width = option.rect.width()
        if width <= 0 and hasattr(self, 'parent_table') and self.parent_table:
            width = self.parent_table.columnWidth(index.column())
            
        if width <= 0: width = 100
        
        text_width = width - 8 
        if text_width < 10: text_width = 10
        
        # Default font from option
        font = option.font
        
        if hasattr(self, 'parent_table') and hasattr(self.parent_table, 'parent_app'):
             try:
                 app = self.parent_table.parent_app
                 font = QFont(self.font_family, app.current_font_size)
             except:
                 pass
                 
        try:
            # OPTIMIZATION: Use QFontMetrics
            fm = QFontMetrics(font)
            rect = fm.boundingRect(QRect(0, 0, text_width, 0), Qt.TextFlag.TextWordWrap, text)
            h = rect.height() + 10 # Buffer
            return QSize(width, h)
        except:
            return super().sizeHint(option, index)

    def paint(self, painter, option, index):
        """Hardened delegate to force selection colors, correct rendering widgets, and professional RTL behavior."""
        opt = QStyleOptionViewItem(option)
        self.initStyleOption(opt, index)
        
        # Determine colors based on app mode
        app = getattr(self.parent_table, 'parent_app', None)
        is_dark = app.is_dark_mode if app else False
        
        # Check for search match highlight
        search_term = ""
        if app and app.highlight_search:
            search_term = app.current_search_term

        text = index.data(Qt.ItemDataRole.DisplayRole) or ""
        
        # Should we highlight search term?
        highlight_bg = False
        if search_term and text and str(search_term).lower() in str(text).lower():
            highlight_bg = True

        # Respect item-level foreground color if set
        fg_data = index.data(Qt.ItemDataRole.ForegroundRole)
        if fg_data:
            if isinstance(fg_data, QBrush):
                text_color = fg_data.color()
            elif isinstance(fg_data, QColor):
                text_color = fg_data
            else:
                text_color = QColor("white") if is_dark else QColor("#2d3748")
        else:
            text_color = QColor("white") if is_dark else QColor("#2d3748")

        painter.save()
        
        # 1. Background State
        if opt.state & QStyle.StateFlag.State_Selected:
            painter.fillRect(opt.rect, QColor("#e6fffa") if not is_dark else QColor("#2c5282"))
            # High-contrast selection text
            if not fg_data:
                text_color = QColor("#2d3748") if not is_dark else QColor("white")
        elif highlight_bg:
            painter.fillRect(opt.rect, QColor("#fefcbf") if not is_dark else QColor("#553c9a"))
        
        # 2. Text Drawing Logic
        painter.setPen(text_color)
        font_size = app.current_font_size if app else 14
        painter.setFont(QFont(self.font_family, font_size))
        
        # Alignment
        alignment = Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
        item = self.parent_table.item(index.row(), index.column())
        if item:
            alignment = item.textAlignment()
            
        text_flags = alignment
        # Summary & Remarks Wrap (Col 11 in Master) OR Track Data (Col 1 in Track Table)
        is_track_table = (hasattr(self.parent_table, "property") and self.parent_table.property("table_name") == "tracks") or (self.parent_table.objectName() == "TrackTable")
        
        if index.column() == 11 or (is_track_table and index.column() == 1):
             text_flags |= Qt.TextFlag.TextWordWrap
             
        text_rect = opt.rect.adjusted(8, 4, -8, -4)
        painter.drawText(text_rect, text_flags, str(text))
        
        painter.restore()
            
    def createEditor(self, parent, option, index):
        # 0. STRICT READ-ONLY for Audio No
        
        is_audio_no = False
        if hasattr(self.parent_table, "horizontalHeaderItem"):
            hi = self.parent_table.horizontalHeaderItem(index.column())
            if hi:
                h_text = hi.text().lower()
                # print(f"DEBUG: Delegate Editor Check: Col {index.column()}, Header: {h_text}")
                if "audio" in h_text and ("no" in h_text or "number" in h_text):
                    is_audio_no = True
                
        # Fallback for Master Table which definitely has Audio No at Col 1
        # IMPORTANT: Only apply Col 1 check if it's the Master Table (NOT TrackTable)
        is_master_table = (not hasattr(self.parent_table, "property") or not self.parent_table.property("sheet_id")) and (self.parent_table.objectName() != "TrackTable")
        
        if is_master_table and index.column() == 1:
             is_audio_no = True

        if is_audio_no:
            # print(f"DEBUG: Column {index.column()} LOCKED (Audio No)")
            return None

        # 1. Determine Font Size
        font_size = 14
        if hasattr(self.parent_table, 'parent_app'):
            font_size = self.parent_table.parent_app.current_font_size
        
        font = QFont(self.font_family, font_size)

        # 2. Check Alignment
        item = self.parent_table.item(index.row(), index.column())
        is_left_aligned = False
        if item and (item.textAlignment() & Qt.AlignmentFlag.AlignLeft):
            is_left_aligned = True

        # 3. Create Widget based on Mode
        
        # Dynamic Mode Detection for "text" mode (Other Sheets)
        current_mode = self.mode
        if current_mode == "text" and hasattr(self.parent_table, "horizontalHeaderItem"):
            header_item = self.parent_table.horizontalHeaderItem(index.column())
            if header_item:
                header_text = header_item.text().lower()
                if "person" in header_text: current_mode = "person"
                elif "occasion" in header_text: current_mode = "occasion"
                elif "category" in header_text: current_mode = "category"
                elif "place" in header_text: current_mode = "place"
                elif "country" in header_text: current_mode = "country"
                elif "incomplete" in header_text: current_mode = "incomplete"
                elif "a/v" in header_text: current_mode = "AV"
        
        if current_mode in ["person", "occasion", "category", "place", "AV", "country", "incomplete"]:
            # --- COMBOBOX LOGIC ---

            combo = QComboBox(parent)
            combo.setEditable(True)
            combo.setFont(font)
            combo.lineEdit().setFont(font)
            combo.installEventFilter(self)
            
            items = self.get_items_from_db()
            combo.addItems(items)
            
            # Popup Width Logic
            fm = QFontMetrics(font)
            max_item_w = 0
            for i in items:
                w = fm.horizontalAdvance(str(i)) + 70
                if w > max_item_w: max_item_w = w
            
            popup_min_w = max(max_item_w, self.parent_table.columnWidth(index.column()))
            combo.view().setMinimumWidth(popup_min_w)
            combo.view().setFont(font)
            
            completer = QCompleter(items, combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            if completer.popup():
                completer.popup().setFont(font)
                completer.popup().setMinimumWidth(popup_min_w)
            combo.setCompleter(completer)
            
            if is_left_aligned:
                combo.lineEdit().setAlignment(Qt.AlignmentFlag.AlignLeft)
            else:
                combo.lineEdit().setAlignment(Qt.AlignmentFlag.AlignRight)
            
            # CONNECT IMMEDIATE SAVE ON SELECTION
            # When user selects an item from the list, commit and close.
            combo.activated.connect(lambda: self.commit_and_close(combo))
            
            return combo

        elif index.column() in [7, 8, 9]: # Date Columns + Year
             date_editor = QLineEdit(parent)
             if index.column() in [7, 8]:
                 date_editor.setInputMask("99/99/9999")
             else:
                 # Year column, just numbers/text
                 pass
             date_editor.setFont(QFont("Arial", font_size)) # Arial for numbers
             date_editor.installEventFilter(self)
             return date_editor

        else:
            # --- TEXT EDIT LOGIC (Tracks + Standard Fields) ---
            editor = QTextEdit(parent)
            editor.setFont(font)
            editor.setAcceptRichText(False)
            editor.installEventFilter(self)
            
            # Force RTL / Alignment logic
            if is_left_aligned:
                editor.setLayoutDirection(Qt.LayoutDirection.LeftToRight)
                editor.setAlignment(Qt.AlignmentFlag.AlignLeft)
            else:
                # Force Right-to-Left for everything else (including Tracks)
                editor.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                text_option = QTextOption()
                text_option.setTextDirection(Qt.LayoutDirection.RightToLeft)
                text_option.setAlignment(Qt.AlignmentFlag.AlignRight)
                editor.document().setDefaultTextOption(text_option)
                editor.setAlignment(Qt.AlignmentFlag.AlignRight)
            
            # Style
            app = getattr(self.parent_table, 'parent_app', None)
            is_dark = app.is_dark_mode if app else False
            bg = "#1e293b" if is_dark else "white"
            color = "white" if is_dark else "#2d3748"
            border = "#4a5568" if is_dark else "#3182ce"
            
            editor.setStyleSheet(f"background-color: {bg}; color: {color}; border: 2px solid {border}; padding: 2px;")
            editor.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            
            # Connect auto-resize ONLY for QTextEdit (Tracks)
            # Disable for ArchiveTable (Master Sheet) to prevent layout crashes on large data
            if self.parent_table.__class__.__name__ != "ArchiveTable":
                editor.document().contentsChanged.connect(lambda: self.on_text_changed(editor, index))
            return editor

    def commit_and_close(self, editor):
        """Helper to delay commit slightly to allow QComboBox to finalize selection state."""
        QTimer.singleShot(0, lambda: self._do_commit(editor))

    def _do_commit(self, editor):
        try:
            self.commitData.emit(editor)
            self.closeEditor.emit(editor)
        except:
            pass

    def on_text_changed(self, editor, index):
        """Auto-expand row height as user types."""
        doc_h = editor.document().size().height()
        new_h = int(doc_h + 10)
        current_h = self.parent_table.rowHeight(index.row())
        
        # Enforce minimum height (e.g. 45) but allow growth
        if new_h < 45: new_h = 45
        
        # Only update if meaningful difference to prevent jitter/glitches
        if abs(new_h - current_h) > 3:
            self.parent_table.setRowHeight(index.row(), new_h)

    def setEditorData(self, editor, index):
        text = str(index.data(Qt.ItemDataRole.EditRole) or "")
        
        if isinstance(editor, QComboBox):
            editor.setCurrentText(text)
        elif isinstance(editor, QTextEdit):
            editor.setPlainText(text)
            
            # Re-assert Alignment
            opt = QTextOption()
            opt.setTextDirection(Qt.LayoutDirection.RightToLeft)
            opt.setAlignment(Qt.AlignmentFlag.AlignRight)
            editor.document().setDefaultTextOption(opt)
            editor.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
            editor.setAlignment(Qt.AlignmentFlag.AlignRight)

            cursor = editor.textCursor()
            cursor.movePosition(cursor.MoveOperation.End)
            editor.setTextCursor(cursor)
        elif isinstance(editor, QLineEdit):
            editor.setText(text)

    def setModelData(self, editor, model, index):
        """Override to capture OLD vs NEW value for Delta-Undo and block HTML."""
        old_val = str(index.model().data(index, Qt.ItemDataRole.EditRole) or "")
        
        # Determine current text based on editor type
        if isinstance(editor, QComboBox):
            text = editor.currentText().strip()
        elif isinstance(editor, QTextEdit):
            text = editor.toPlainText().strip()
        elif isinstance(editor, QLineEdit):
            text = editor.text().strip()
        else:
            text = ""
            
        # VALIDATION: If we accidentally got HTML rubbish, try to strip it (failsafe)
        if text.startswith("<!DOCTYPE HTML"):
            # This shouldn't happen with toPlainText(), but just in case
            pass

        # Manual Data Set to avoid super() possibly using HTML
        model.setData(index, text, Qt.ItemDataRole.EditRole)
        
        if text != old_val:
             # Push Undo Action via parent table's app reference
             if hasattr(self.parent_table, 'parent_app'):
                 app = self.parent_table.parent_app
                 pk = None
                 
                 if self.parent_table.__class__.__name__ == "ArchiveTable":
                     pk_item = self.parent_table.item(index.row(), 1) # Audio No
                     pk = pk_item.data(Qt.ItemDataRole.UserRole) if pk_item else None
                 
                 if pk:
                      app.push_undo_action({'type':'edit', 'pk':pk, 'col':index.column(), 'old_val':old_val})
                 else:
                      app.save_undo_state(self.parent_table)
                 
                 app.mark_unsaved()
        
        # Location Auto-Fill Country (Legacy Logic)
        if self.mode == "place":
            conn = sqlite3.connect(DB_FILE, timeout=30)
            conn.execute("PRAGMA journal_mode=WAL")
            res = conn.execute("SELECT country FROM locations WHERE place = ?", (text,)).fetchone()
            if res:
                col_idx = model.index(index.row(), 6) # Assuming Col 6 is country
                model.setData(col_idx, str(res[0]), Qt.ItemDataRole.EditRole)
            conn.close()

    def eventFilter(self, editor, event):
        """Intercepts Enter key inside cell editors."""
        if event.type() == QEvent.Type.KeyPress:
            if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                if isinstance(editor, QComboBox) and editor.view().isVisible():
                     # If popup is open, pressing Enter selects the item. 
                     # The 'activated' signal (connected above) will handle the save/close.
                     return False
                
                if isinstance(editor, QTextEdit):
                    # Allow Shift+Enter for new lines in TextEdit? 
                    # Or just standard Enter to commit? 
                    # User likely wants multi-line if row resizes.
                    # If modifiers & Shift, allow newline.
                    if event.modifiers() & Qt.KeyboardModifier.ShiftModifier:
                        return False # Let default handler insert newline
                
                # Commit data on Enter (without Shift)
                if editor is not None:
                    try:
                        self.commitData.emit(editor)
                        self.closeEditor.emit(editor)
                    except:
                        pass
                return True
        return super().eventFilter(editor, event)

    def get_items_from_db(self):
        conn = sqlite3.connect(DB_FILE, timeout=30)
        conn.execute("PRAGMA journal_mode=WAL")
        items = []
        try:
            if self.mode == "place":
                res = conn.execute("SELECT DISTINCT place FROM locations ORDER BY place ASC").fetchall()
            elif self.mode == "country":
                res = conn.execute("SELECT DISTINCT country FROM locations WHERE country != '' ORDER BY country ASC").fetchall()
            else:
                cat_search = ""
                if self.mode == "AV": cat_search = "AV"
                else: cat_search = self.mode.capitalize()
                res = conn.execute("SELECT value FROM dropdown_options WHERE category = ? ORDER BY value ASC", (cat_search,)).fetchall()
            for r in res:
                if r[0]: items.append(str(r[0]))
        except:
            pass
        finally:
            conn.close()
        return items

class FilterDialog(QDialog):
    """Excel-style filter popup with Search and Select All."""
    def __init__(self, title, items, selected_items=None, parent=None, font_family="Arial"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(300, 400)
        self.items = items
        self.selected_items = set(selected_items) if selected_items else set()
        self.font_family = font_family
        
        layout = QVBoxLayout(self)
        
        # Search
        self.txt_search = QLineEdit()
        self.txt_search.setPlaceholderText("Search items...")
        self.txt_search.textChanged.connect(self.filter_list)
        layout.addWidget(self.txt_search)
        
        # Select All / Clear All
        self.chk_all = QCheckBox("Select All")
        self.chk_all.setStyleSheet("color: black; font-weight: bold; margin-bottom: 5px;")
        self.chk_all.stateChanged.connect(self.toggle_all)
        layout.addWidget(self.chk_all)
        
        # List
        self.list_widget = QListWidget()
        self.list_widget.setFont(QFont(self.font_family, 13))
        # Ensure (Blanks) is top if present
        if "(Blanks)" in items:
             items.remove("(Blanks)")
             items.insert(0, "(Blanks)")
             
        layout.addWidget(self.list_widget)
        
        # Populate
        self.populate_list(self.items)
        
        # Buttons
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)
        
        # Initial Check for Select All
        self.update_select_all_state()

    def populate_list(self, all_items):
        self.list_widget.clear()
        for text in all_items:
            item = QListWidgetItem(str(text))
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            
            if text in self.selected_items:
                item.setCheckState(Qt.CheckState.Checked)
            else:
                item.setCheckState(Qt.CheckState.Unchecked)
            self.list_widget.addItem(item)

    def filter_list(self):
        term = self.txt_search.text().lower()
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if term in item.text().lower():
                item.setHidden(False)
            else:
                item.setHidden(True)

    def toggle_all(self, state):
        is_checked = (state == Qt.CheckState.Checked.value) # 2
        # Only toggle VISIBLE items
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if not item.isHidden():
                item.setCheckState(Qt.CheckState.Checked if is_checked else Qt.CheckState.Unchecked)

    def update_select_all_state(self):
        # We could check if all are selected to set tri-state, but keeping it simple.
        pass

    def get_selected(self):
        selected = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.checkState() == Qt.CheckState.Checked:
                selected.append(item.text())
        return selected


class TrackManagerWidget(QFrame):
    def __init__(self, audio_no, parent_app):
        super().__init__()
        self.audio_no = str(audio_no).strip()
        self.parent_app = parent_app
        self.setObjectName("TrackBox")
        
        # Dynamic Style
        if self.parent_app.is_dark_mode:
             self.setStyleSheet("QFrame#TrackBox { background-color: #1e1f20; border: 2px solid #444746; border-radius: 8px; }")
        else:
             self.setStyleSheet("QFrame#TrackBox { background-color: white; border: 2px solid #3182ce; border-radius: 8px; }")
        
        # Container to allow horizontal "stickiness" inside the spanned cell
        self.container = QFrame(self)
        self.container.setObjectName("TrackContainer")
        self.container_layout = QVBoxLayout(self.container)
        self.container_layout.setContentsMargins(10, 10, 10, 10)
        self.container_layout.setSpacing(5)
        
        header = QHBoxLayout()
        header_text = f"Tracks - Audio No: {self.audio_no}"
        self.lbl_title = QLabel(header_text)
        color_text = "#e3e3e3" if self.parent_app.is_dark_mode else "#2d3748"
        self.lbl_title.setStyleSheet(f"font-weight: bold; color: {color_text}; font-size: 14px;")
        header.addWidget(self.lbl_title)
        header.addStretch()
        
        self.btn_import_word = QPushButton("📂 Import Word")
        self.btn_import_word.clicked.connect(self.import_from_word)
        self.btn_import_word.setObjectName("SecondaryBtn") 
        header.addWidget(self.btn_import_word)
        
        self.btn_add_track = QPushButton("➕ Add Row")
        self.btn_add_track.clicked.connect(self.add_empty_track)
        header.addWidget(self.btn_add_track)
        self.container_layout.addLayout(header)
        
        self.track_table = QTableWidget()
        self.track_table.setObjectName("TrackTable")
        self.track_table.parent_app = self.parent_app 
        self.track_table.setAlternatingRowColors(self.parent_app.alternating_rows)
        self.track_table.setColumnCount(2)
        self.track_table.setHorizontalHeaderLabels(["No", "Track Data"])
        self.track_table.verticalHeader().setVisible(False)
        self.track_table.setWordWrap(True)
        self.track_table.setLayoutDirection(Qt.LayoutDirection.LeftToRight) 
        self.track_table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.track_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.track_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.track_table.setMinimumHeight(250)
        
        self.track_delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", self.track_table)
        for i in range(2):
            self.track_table.setItemDelegateForColumn(i, self.track_delegate)
            
        self.track_table.itemChanged.connect(self.save_track_data)
        self.container_layout.addWidget(self.track_table)
        
        # Sync row heights on column resize
        self.track_table.horizontalHeader().sectionResized.connect(lambda: QTimer.singleShot(50, self.track_table.resizeRowsToContents))
        
        # Initial Locked State
        self.set_locked(self.parent_app.master_locked)
        
        self.load_tracks()
        
        # Initial Position Sync
        QTimer.singleShot(10, self.sync_to_viewport)

    def set_locked(self, locked):
        """Disables editing and buttons based on global lock."""
        if locked:
            self.track_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.btn_import_word.setEnabled(False)
            self.btn_add_track.setEnabled(False)
            # Remove grey style as it confuses users, just lock editing
            self.track_table.setStyleSheet("") 
        else:
            self.track_table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.SelectedClicked | QAbstractItemView.EditTrigger.AnyKeyPressed)
            self.btn_import_word.setEnabled(True)
            self.btn_add_track.setEnabled(True)
            self.track_table.setStyleSheet("")

    def sync_to_viewport(self):
        """Moves the container so it stays within the visible viewport of the table."""
        master_table = self.parent_app.table
        viewport_w = master_table.viewport().width()
        scroll_x = master_table.horizontalScrollBar().value()
        
        # Sizing: Leave some margin
        desired_w = viewport_w - 20
        self.container.setFixedSize(desired_w, self.height() - 10)
        
        # Position: Offset by scroll_x to 'counter' the cell's movement
        self.container.move(scroll_x + 10, 5)

    def showEvent(self, event):
        super().showEvent(event)
        self.sync_to_viewport()
        if hasattr(self, 'track_table'):
            QTimer.singleShot(50, self.track_table.resizeRowsToContents)
            
    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.sync_to_viewport()

    def load_tracks(self):
        self.track_table.blockSignals(True)
        self.track_table.setRowCount(0)
        conn = sqlite3.connect(DB_FILE, timeout=30)
        conn.execute("PRAGMA journal_mode=WAL")
        res = conn.execute("SELECT track_no, track_name FROM tracks WHERE event_id = ? ORDER BY id ASC", (self.audio_no,)).fetchall()
        conn.close()
        if not res:
            self.add_empty_track()
        else:
            for i, row in enumerate(res):
                self.insert_track_row(i, row[0], row[1])
        self.track_table.blockSignals(False)

    def add_empty_track(self):
        self.track_table.blockSignals(True)
        row = self.track_table.rowCount()
        self.insert_track_row(row, str(row + 1), "")
        conn = sqlite3.connect(DB_FILE, timeout=30)
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("INSERT INTO tracks (event_id, track_no, track_name) VALUES (?,?,?)", (self.audio_no, str(row + 1), ""))
        conn.commit()
        conn.close()
        self.track_table.blockSignals(False)

    def insert_track_row(self, row, no, data):
        self.track_table.insertRow(row)
        
        item_no = LisanTableItem(str(no), self.parent_app.lisan_font_family, self.parent_app.current_font_size)
        item_no.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self.track_table.setItem(row, 0, item_no)
        
        item_data = LisanTableItem(str(data), self.parent_app.lisan_font_family, self.parent_app.current_font_size)
        item_data.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self.track_table.setItem(row, 1, item_data)
        
        if self.track_table.columnWidth(1) > 200:
            col_w = self.track_table.columnWidth(1)
        else:
            col_w = 400 # Default fallback width for calculation
            
        font = QFont(self.parent_app.lisan_font_family, self.parent_app.current_font_size)
        fm = QFontMetrics(font)
        rect = fm.boundingRect(QRect(0,0, col_w, 1000), Qt.TextFlag.TextWordWrap, str(data))
        est_height = rect.height() + 20 # Padding
        
        self.track_table.setRowHeight(row, max(45, est_height))

    def save_track_data(self, item):
        conn = sqlite3.connect(DB_FILE, timeout=30)
        conn.execute("PRAGMA journal_mode=WAL")
        tracks = []
        for r in range(self.track_table.rowCount()):
            it_n = self.track_table.item(r, 0)
            it_t = self.track_table.item(r, 1)
            n_text = it_n.text().strip() if it_n else ""
            t_text = it_t.text().strip() if it_t else ""
            
            # Skip empty rows (Auto-Cleanup)
            if not n_text and not t_text:
                continue
                
            tracks.append((n_text, t_text))
            
        conn.execute("DELETE FROM tracks WHERE event_id = ?", (self.audio_no,))
        for n, t in tracks:
            conn.execute("INSERT INTO tracks (event_id, track_no, track_name) VALUES (?,?,?)", (self.audio_no, n, t))
        conn.commit()
        conn.close()

    def import_from_word(self):
        if not Document:
            return
            
        path_tuple = QFileDialog.getOpenFileName(self, "Select Word File", "", "Word Documents (*.docx *.doc)")
        path = path_tuple[0]
        if not path:
            return
            
        temp_docx = None
        try:
            # Handle legacy .doc files by converting to .docx via macOS textutil
            if path.lower().endswith(".doc"):
                fd, temp_docx = tempfile.mkstemp(suffix=".docx")
                os.close(fd) 
                
                # Command: textutil -convert docx -output [temp_docx] [original_doc]
                cmd = ["textutil", "-convert", "docx", "-output", temp_docx, path]
                result = subprocess.run(cmd, capture_output=True, text=True)
                
                if result.returncode != 0:
                    QMessageBox.warning(self, "Import Error", f"Could not convert legacy .doc file: {result.stderr}")
                    return
                
                process_path = temp_docx
            else:
                process_path = path

            doc = Document(process_path)
            
            # Legacy logic: Check tables
            if len(doc.tables) < 2:
                # Some .doc converted files might put tracks in a different table index or structure
                # We'll try to find the track table by content if the 2nd one is missing
                target_table = None
                for t in doc.tables:
                    if len(t.rows) > 0 and len(t.columns) >= 2:
                        # Simple heuristic: first cell is often a number or small text
                        target_table = t
                        break
                if not target_table:
                    QMessageBox.warning(self, "Import Error", "No track table found in the document.")
                    return
                table = target_table
            else:
                table = doc.tables[1]
                
            self.track_table.blockSignals(True)
            conn = sqlite3.connect(DB_FILE, timeout=30)
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("DELETE FROM tracks WHERE event_id = ?", (self.audio_no,))
            self.track_table.setRowCount(0)
            
            for row in table.rows:
                # Ensure we have at least 2 cells
                if len(row.cells) < 2: continue
                
                t_num = row.cells[0].text.strip()
                t_content = " ".join([c.text.strip() for c in row.cells[1:]]).strip()
                
                if t_num or t_content:
                    idx = self.track_table.rowCount()
                    self.insert_track_row(idx, t_num, t_content)
                    conn.execute("INSERT INTO tracks (event_id, track_no, track_name) VALUES (?,?,?)", (self.audio_no, t_num, t_content))
            
            conn.commit()
            conn.close()
            self.track_table.blockSignals(False)
            
        except Exception as e:
            QMessageBox.warning(self, "Import Error", f"An error occurred during import: {e}")
            traceback.print_exc()
        finally:
            # Cleanup temp file if created
            if temp_docx and os.path.exists(temp_docx):
                try: os.remove(temp_docx)
                except: pass

class MasterListTable(QTableWidget):
    """Specialized table for Master Lists with protection and locking."""
    def __init__(self, parent_app, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.parent_app = parent_app
        self.is_locked = True 
        
        v_header = self.verticalHeader()
        v_header.setVisible(True)
        v_header.setFont(QFont("Arial", 12))
        v_header.setMinimumWidth(45)
        v_header.setDefaultSectionSize(40)
        
        self.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.horizontalHeader().setStretchLastSection(True)

    def set_lock_state(self, locked):
        self.is_locked = locked
        if locked:
            self.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        else:
            self.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.AnyKeyPressed)

    def keyPressEvent(self, event: QKeyEvent):
        if self.is_locked:
            return
            
        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            if self.state() != QAbstractItemView.State.EditingState:
                items = self.selectedItems()
                if len(items) > 1:
                    msg = "You are about to clear multiple items.\n\nAre you sure you want to proceed?"
                    reply = QMessageBox.question(self, "Confirm Clear", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                    if reply == QMessageBox.StandardButton.No:
                        return
                for item in items:
                    item.setText("")
                return
        super().keyPressEvent(event)

class ArchiveTable(QTableWidget):
    """Primary Grid component shared across main pages."""
    def __init__(self, parent_app, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.parent_app = parent_app
        self.horizontalScrollBar().valueChanged.connect(self.sync_all_tracks)

    def sync_all_tracks(self):
        # Sync all open track managers
        if hasattr(self.parent_app, 'open_tracks'):
            for manager in self.parent_app.open_tracks.values():
                if hasattr(manager, 'sync_to_viewport'):
                    manager.sync_to_viewport()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.sync_all_tracks()
            
    def keyPressEvent(self, event: QKeyEvent):
        # LOCK GUARD: Prevent clearing cells with Backspace/Delete when locked
        if hasattr(self.parent_app, 'master_locked') and self.parent_app.master_locked:
            if event.key() in (Qt.Key.Key_Backspace, Qt.Key.Key_Delete):
                return # Ignore
                
        if event.matches(QKeySequence.StandardKey.Copy):
            self.copy_selection()
            return
        if event.matches(QKeySequence.StandardKey.Paste):
            self.paste_selection()
            return
        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            if self.state() != QAbstractItemView.State.EditingState:
                items = self.selectedItems()
                if items:
                    self.parent_app.save_undo_state()
                    for item in items:
                        item.setText("")
                return 
        super().keyPressEvent(event)
        
    def copy_selection(self):
        ranges = self.selectedRanges()
        if not ranges:
            return
        r = ranges[0]
        rows = []
        for i in range(r.rowCount()):
            row_data = []
            for j in range(r.columnCount()):
                it = self.item(r.topRow() + i, r.leftColumn() + j)
                if it:
                    row_data.append(it.text())
                else:
                    row_data.append("")
            rows.append("\t".join(row_data))
        QApplication.clipboard().setText("\n".join(rows))
        
    def paste_selection(self):
        text = QApplication.clipboard().text()
        curr = self.currentIndex()
        if not text or not curr.isValid():
            return
        
        self.parent_app.save_undo_state()
        self.setUpdatesEnabled(False)
        self.blockSignals(True)
        
        try:
            import sqlite3, traceback
            conn = sqlite3.connect(DB_FILE)
            cursor = conn.cursor()
            
            # CRLF normalization
            raw_rows = text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
            if raw_rows and not raw_rows[-1]: raw_rows.pop()
            
            start_row = curr.row()
            start_col = curr.column()
            
            db_map = ["audio_no", "person", "occasion", "category", "place", "country", "hijri_date", "esavi_date", "year", "out_of", "remarks", "Tracks", "AV", "cass_no", "came_from", "incomplete"]

            for i, row_str in enumerate(raw_rows):
                # Calculate the target data row (skipping any expansion rows in between)
                target_row = -1
                found_count = 0
                for r in range(start_row, self.rowCount()):
                    # Expansion rows have span > 1
                    if self.columnSpan(r, 0) > 1:
                        continue
                    if found_count == i:
                        target_row = r
                        break
                    found_count += 1
                
                if target_row == -1 or target_row >= self.rowCount():
                    break
                
                cells = row_str.split('\t')
                
                # Get ID for DB update
                audio_no_item = self.item(target_row, 1)
                audio_no = audio_no_item.data(Qt.ItemDataRole.UserRole) if audio_no_item else None
                
                if not audio_no: continue

                for j, cell_text in enumerate(cells):
                    target_col = start_col + j
                    if target_col >= self.columnCount() or target_col < 1: # Col 0 is "+"
                        continue
                        
                    item = self.item(target_row, target_col)
                    if not item: continue
                    
                    val = cell_text.strip()
                    item.setText(val)
                    
                    # Update DB
                    db_col_idx = target_col - 1
                    if 0 <= db_col_idx < len(db_map):
                        db_col_name = db_map[db_col_idx]
                        cursor.execute(f"UPDATE events SET {db_col_name} = ? WHERE audio_no = ?", (val, audio_no))
            
            conn.commit()
            conn.close()
            self.parent_app.mark_unsaved()
            
        except Exception as e:
            print(f"Paste error: {e}")
            traceback.print_exc()
        finally:
            self.blockSignals(False)
            self.setUpdatesEnabled(True)
            self.viewport().update()
            # Final touch: Ensure rows fit the new text
            self.resizeRowsToContents()

# --- 4. MAIN APPLICATION WINDOW ---

class AddSheetDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add New Sheet")
        self.setFixedSize(400, 200)
        self.setStyleSheet("background-color: white; color: #2d3748;")
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        
        lbl = QLabel("Enter Sheet Name:")
        lbl.setStyleSheet("font-weight: bold; font-size: 14px;")
        layout.addWidget(lbl)
        
        self.inp_name = QLineEdit()
        self.inp_name.setPlaceholderText("e.g. Finance 2024")
        self.inp_name.setStyleSheet("padding: 8px; border: 1px solid #cbd5e0; border-radius: 4px;")
        layout.addWidget(self.inp_name)
        
        btn_layout = QHBoxLayout()
        self.btn_import = QPushButton("📂 Import Excel")
        self.btn_import.setStyleSheet("background-color: #38a169; color: white; padding: 8px; border-radius: 4px; font-weight: bold;")
        self.btn_import.clicked.connect(self.accept_import)
        
        self.btn_manual = QPushButton("📝 Create Manual")
        self.btn_manual.setStyleSheet("background-color: #3182ce; color: white; padding: 8px; border-radius: 4px; font-weight: bold;")
        self.btn_manual.clicked.connect(self.accept_manual)
        
        btn_layout.addWidget(self.btn_import)
        btn_layout.addWidget(self.btn_manual)
        layout.addLayout(btn_layout)
        
        self.mode = None # "import" or "manual"
        
    def accept_import(self):
        if not self.inp_name.text().strip():
             QMessageBox.warning(self, "Error", "Please enter a sheet name first.")
             return
        self.mode = "import"
        self.accept()
        
    def accept_manual(self):
        if not self.inp_name.text().strip():
             QMessageBox.warning(self, "Error", "Please enter a sheet name first.")
             return
        self.mode = "manual"
        self.accept()

class OtherSheetTable(QTableWidget):
    """Custom Table logic for Other Sheets (Paste Support)."""
    def __init__(self, parent_app, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.parent_app = parent_app
        
    def keyPressEvent(self, event):
        if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            # Consume Enter so it doesn't open the editor with HTML rubbish
            event.accept()
            # Move focus down
            curr = self.currentIndex()
            if curr.isValid() and curr.row() < self.rowCount() - 1:
                self.setCurrentIndex(self.model().index(curr.row()+1, curr.column()))
            return

        if hasattr(self.parent_app, 'master_locked') and self.parent_app.master_locked:
            # Block all keys that might trigger edit when master locked
            if event.text() and not event.modifiers():
                 return
            if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace, Qt.Key.Key_Return, Qt.Key.Key_Enter):
                 return

        if event.matches(QKeySequence.StandardKey.Copy):
            self.copy_selection()
            return
        if event.matches(QKeySequence.StandardKey.Paste):
            self.paste_selection()
            return

        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            if self.state() != QAbstractItemView.State.EditingState:
                items = self.selectedItems()
                if items:
                    self.parent_app.save_undo_state(self)
                    for item in items:
                        item.setText("")
                return
        super().keyPressEvent(event)

        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            if self.state() != QAbstractItemView.State.EditingState:
                items = self.selectedItems()
                if items:
                    self.parent_app.save_undo_state(self)
                    for item in items:
                        item.setText("")
                return
        super().keyPressEvent(event)

    def copy_selection(self):
        ranges = self.selectedRanges()
        if not ranges: return
        r = ranges[0]
        rows = []
        for i in range(r.rowCount()):
            row_data = []
            for j in range(r.columnCount()):
                it = self.item(r.topRow() + i, r.leftColumn() + j)
                row_data.append(it.text() if it else "")
            rows.append("\t".join(row_data))
        QApplication.clipboard().setText("\n".join(rows))

    def paste_selection(self):
        text = QApplication.clipboard().text()
        if not text: return
        
        curr = self.currentIndex()
        if not curr.isValid(): return
        
        self.setUpdatesEnabled(False)
        self.blockSignals(True)
        
        try:
            raw_rows = text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
            if raw_rows and not raw_rows[-1]: raw_rows.pop()
            
            start_row = curr.row()
            start_col = curr.column()
            
            self.parent_app.save_undo_state(self)
            
            for i, row_str in enumerate(raw_rows):
                r = start_row + i
                if r >= self.rowCount(): break
                p_cells = row_str.split('\t')
                for j, val in enumerate(p_cells):
                    c = start_col + j
                    if c >= self.columnCount(): break
                    item = self.item(r, c)
                    if not item:
                        item = QTableWidgetItem()
                        self.setItem(r, c, item)
                    item.setText(val.strip())

            self.parent_app.mark_unsaved()
            
        except Exception as e:
            print(f"OtherSheet Paste Error: {e}")
        finally:
            self.blockSignals(False)
            self.setUpdatesEnabled(True)
            self.viewport().update()
            self.resizeRowsToContents()

class OldSheetsManager(QWidget):
    def __init__(self, parent_app):
        super().__init__()
        self.parent_app = parent_app
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header = QHBoxLayout()
        title = QLabel("Other Sheets Archive")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2d3748;")
        header.addWidget(title)
        header.addStretch()
        
        btn_add = QPushButton("➕ Add Sheet")
        btn_add.setFixedSize(120, 40)
        btn_add.setStyleSheet("background-color: #3182ce; color: white; font-weight: bold; border-radius: 6px;")
        btn_add.clicked.connect(self.add_sheet_flow)
        header.addWidget(btn_add)

        btn_add_row = QPushButton("➕ Add Row")
        btn_add_row.setFixedSize(120, 40)
        btn_add_row.setStyleSheet("background-color: #48bb78; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_add_row.clicked.connect(self.add_row)
        header.addWidget(btn_add_row)

        btn_del_row = QPushButton("➖ Delete Row")
        btn_del_row.setFixedSize(120, 40)
        btn_del_row.setStyleSheet("background-color: #e53e3e; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_del_row.clicked.connect(self.delete_row)
        header.addWidget(btn_del_row)

        btn_del_sheet = QPushButton("🗑️ Delete Sheet")
        btn_del_sheet.setFixedSize(130, 40)
        btn_del_sheet.setStyleSheet("background-color: #718096; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_del_sheet.clicked.connect(self.delete_sheet)
        header.addWidget(btn_del_sheet)

        btn_add_col = QPushButton("➕ Add Col")
        btn_add_col.setFixedSize(120, 40)
        btn_add_col.setStyleSheet("background-color: #805ad5; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_add_col.clicked.connect(self.add_column)
        header.addWidget(btn_add_col)

        btn_del_col = QPushButton("➖ Del Col")
        btn_del_col.setFixedSize(120, 40)
        btn_del_col.setStyleSheet("background-color: #e53e3e; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_del_col.clicked.connect(self.delete_column)
        header.addWidget(btn_del_col)

        # Export Button
        btn_export = QPushButton("📊 Export")
        btn_export.setFixedSize(120, 40)
        btn_export.setStyleSheet("background-color: #2b6cb0; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        btn_export.clicked.connect(self.export_selection_current_sheet)
        header.addWidget(btn_export)

        # Global Lock Button
        self.btn_lock = QPushButton("🔒 Locked")
        self.btn_lock.setFixedSize(120, 40)
        self.btn_lock.setCheckable(True)
        self.btn_lock.setChecked(True)
        self.btn_lock.setStyleSheet("background-color: #c53030; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        self.btn_lock.clicked.connect(self.toggle_lock)
        header.addWidget(self.btn_lock)
        
        layout.addLayout(header)
        
        # Tabs
        self.tabs = QTabWidget()
        self.tabs.setStyleSheet("""
            QTabWidget::pane { border: 1px solid #cbd5e0; background: white; }
            QTabBar::tab { background: #edf2f7; color: #4a5568; padding: 10px 20px; margin-right: 2px; }
            QTabBar::tab:selected { background: white; color: #2d3748; font-weight: bold; border-bottom: 2px solid #3182ce; }
        """)
        layout.addWidget(self.tabs)
        
        self.load_sheets()
        
    def load_sheets(self):
        self.tabs.clear()
        conn = sqlite3.connect(DB_FILE)
        sheets = conn.execute("SELECT id, name FROM old_sheets_meta ORDER BY id ASC").fetchall()
        conn.close()
        
        for input_id, name in sheets:
            self.create_tab(input_id, name)
            
        # Add Live Tracks Sheet
        self.create_tracks_tab()

    def reset_lock(self):
        """Force resets to Locked state (called when navigating away/to)."""
        if not self.btn_lock.isChecked():
             self.btn_lock.setChecked(True)
             self.btn_lock.setText("🔒 Locked")
             self.btn_lock.setStyleSheet("background-color: #c53030; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
             self.toggle_lock() # Apply state

    def toggle_lock(self):
        """Toggles global lock state with confirmation."""
        # Check Master Lock
        if hasattr(self.parent_app, 'master_locked') and self.parent_app.master_locked:
            QMessageBox.warning(self, "Master Locked", "The Master Database is Locked.\nYou must unlock the Master Lock (Main Sheet) first.")
            self.btn_lock.setChecked(True)
            return

        locked = self.btn_lock.isChecked()
        
        if not locked: # User trying to UNLOCK
             # Ask confirmation
             reply = QMessageBox.question(self, "Unlock Editing", "Are you sure you want to UNLOCK editing for all sheets?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
             if reply != QMessageBox.StandardButton.Yes:
                 self.btn_lock.setChecked(True) # Revert
                 return
                 
             self.btn_lock.setText("🔓 Unlocked")
             self.btn_lock.setStyleSheet("background-color: #38a169; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
        else:
             self.btn_lock.setText("🔒 Locked")
             self.btn_lock.setStyleSheet("background-color: #c53030; color: white; font-weight: bold; border-radius: 6px; margin-left: 10px;")
             
        # Apply to all tabs
        for i in range(self.tabs.count()):
            page = self.tabs.widget(i)
            table = page.findChild(QTableWidget)
            if table:
                if locked:
                    table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
                else:
                    table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed)

    def create_tracks_tab(self):
        """Creates the Live Tracks sheet view."""
        tab = QWidget()
        t_layout = QVBoxLayout(tab)
        t_layout.setContentsMargins(0,0,0,0)
        
        table = OtherSheetTable(self.parent_app)
        table.setProperty("is_tracks", True) 
        table.setAlternatingRowColors(self.parent_app.alternating_rows)
        
        delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", table)
        table.setItemDelegate(delegate)
        table.setWordWrap(True)
        table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        # Debounced Resize
        self.resize_timer = QTimer()
        self.resize_timer.setSingleShot(True)
        self.resize_timer.setInterval(500)
        self.resize_timer.timeout.connect(table.resizeRowsToContents)
        table.horizontalHeader().sectionResized.connect(lambda: self.resize_timer.start())
        
        t_layout.addWidget(table)
        self.tabs.addTab(tab, "Tracks (Live)")
        
        # Load Data
        conn = sqlite3.connect(DB_FILE)
        try:
            data = conn.execute("SELECT id, event_id, track_no, track_name FROM tracks ORDER BY event_id ASC, CAST(track_no AS FLOAT) ASC").fetchall()

            # 1. SETUP STRUCTURE FIRST
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels(["ID", "Audio No", "Track No", "Track Name"])
            table.setRowCount(len(data))

            # 2. THEN CONFIGURE SIZING
            header = table.horizontalHeader()
            header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # ID
            header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents) # Audio No
            header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents) # Track No
            header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)          # Track Name (STRETCH)
            table.setWordWrap(True)
            table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
            
            # FONT
            font = QFont(self.parent_app.lisan_font_family, getattr(self.parent_app, 'current_font_size', 12))
            
            for r, row in enumerate(data):
                for c, val in enumerate(row):
                    cell_val = str(val) if val is not None else ""
                    it = QTableWidgetItem(cell_val)
                    # Store ID for ID column (0)
                    if c == 0: it.setData(Qt.ItemDataRole.UserRole, val) 
                    
                    if c == 0: it.setFlags(Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable) # ID read only
                    
                    # APPLY FONT
                    it.setFont(font)
                    
                    table.setItem(r, c, it)
            
            # HIDE ID COLUMN
            table.setColumnHidden(0, True)
            
            table.resizeColumnsToContents() 
            
            # Apply Lock State
            if self.btn_lock.isChecked():
                table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            else:
                table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed)

        except Exception as e:
            print("Tracks Load Error:", e)
        finally:
            conn.close()

    def create_tab(self, sheet_id, name):
        tab = QWidget()
        t_layout = QVBoxLayout(tab)
        t_layout.setContentsMargins(0, 0, 0, 0)
        
        # Table (Custom for Paste)
        table = OtherSheetTable(self.parent_app)
        table.setProperty("sheet_id", sheet_id)
        table.setAlternatingRowColors(self.parent_app.alternating_rows)
        
        # Apply EditorDelegate for Wrapping/Sizing
        delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", table)
        table.setItemDelegate(delegate)

        # Enable Text Wrapping & Auto-Resize
        table.setWordWrap(True)
        table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        # Debounced Resize Logic
        # Store timer on table to prevent gc
        table.resize_timer = QTimer(table)
        table.resize_timer.setSingleShot(True)
        table.resize_timer.setInterval(500)
        table.resize_timer.timeout.connect(table.resizeRowsToContents)
        
        table.horizontalHeader().sectionResized.connect(lambda: table.resize_timer.start())
        
        # Apply Lock State
        if self.btn_lock.isChecked():
            table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        else:
             table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed)
        
        # Load Headers
        conn = sqlite3.connect(DB_FILE)
        headers_res = conn.execute("SELECT col_index, header_name FROM old_sheet_headers WHERE sheet_id = ? ORDER BY col_index ASC", (sheet_id,)).fetchall()
        
        if headers_res:
             cols = len(headers_res)
             table.setColumnCount(cols)
             table.setHorizontalHeaderLabels([h[1] for h in headers_res])
        
        # Load Data
        try:
             # Fetch all columns dynamically, including ROWID
             data = conn.execute(f"SELECT rowid, * FROM old_sheet_data_{sheet_id}").fetchall()
             table.setRowCount(len(data))
             
             for r_idx, row_data in enumerate(data):
                 row_id = row_data[0] # ROWID is first
                 vals = row_data[1:] # Actual data
                 for c_idx, val in enumerate(vals):
                     # DATE CLEANUP: Fix Esavi bad format "1960_06_26 00:00:00" or with double underscores
                     val_str = str(val) if val is not None else ""
                     # Regex matches YYYY (digit 4) + underscores (1 or more) + MM + underscores + DD + anything else
                     # And replaces it with DD/MM/YYYY
                     if re.match(r"\d{4}_+\d{2}_+\d{2}", val_str):
                         val_str = re.sub(r"(\d{4})_+(\d{2})_+(\d{2}).*", r"\3/\2/\1", val_str)
                         
                     it = QTableWidgetItem(val_str)
                     it.setData(Qt.ItemDataRole.UserRole, row_id) # Store ROWID
                     table.setItem(r_idx, c_idx, it)
                     
             table.resizeColumnsToContents()
             
        except Exception:
             pass
             
        conn.close()
        
        # Set Font
        if hasattr(self.parent_app, 'lisan_font_family'):
            # Use app-wide font size setting if available, else default to 12
            # Use correct attribute 'current_font_size'
            font_size = getattr(self.parent_app, 'current_font_size', 12)
            table.setFont(QFont(self.parent_app.lisan_font_family, int(font_size)))
        
        # Data Save -> Deferred (Master like)
        table.itemChanged.connect(lambda item: self.parent_app.mark_unsaved())
        
        # Apply EditorDelegate for generic editing
        # We need to know column count.
        delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", table)
        for c in range(table.columnCount()):
            table.setItemDelegateForColumn(c, delegate)

        t_layout.addWidget(table)
        self.tabs.addTab(tab, name)

    def save_all_sheets(self):
        """Iterates through all tabs and saves their table data to secondary DB tables."""
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        try:
            cursor.execute("BEGIN TRANSACTION")
            for i in range(self.tabs.count()):
                t_widget = self.tabs.widget(i)
                tbl = t_widget.findChild(QTableWidget)
                if not tbl: continue
                sid = tbl.property("sheet_id")
                if not sid: continue
                
                # Get DB columns
                table_info = cursor.execute(f"PRAGMA table_info(old_sheet_data_{sid})").fetchall()
                real_cols = [info[1] for info in table_info] # col_0, col_1...
                
                for r in range(tbl.rowCount()):
                    # Get ROWID
                    row_item = tbl.item(r, 0)
                    row_id = row_item.data(Qt.ItemDataRole.UserRole) if row_item else None
                    if row_id is None: continue
                    
                    # Build Update
                    updates = []
                    params = []
                    for c in range(tbl.columnCount()):
                        it = tbl.item(r, c)
                        val = it.text().strip() if it else ""
                        
                        target_db_col = None
                        std_col = f"col_{c}"
                        if std_col in real_cols:
                            target_db_col = std_col
                        
                        if target_db_col:
                            updates.append(f"\"{target_db_col}\" = ?")
                            params.append(val)
                    
                    if updates:
                        params.append(row_id)
                        sql = f"UPDATE old_sheet_data_{sid} SET {', '.join(updates)} WHERE rowid = ?"
                        cursor.execute(sql, tuple(params))
            
            cursor.execute("COMMIT")
            self.parent_app.log_message("All Other Sheets saved successfully.")
        except Exception as e:
            print(f"Error saving Other Sheets: {e}")
            cursor.execute("ROLLBACK")
        finally:
            conn.close()
        
    def add_sheet_flow(self):
        dlg = AddSheetDialog(self)
        if dlg.exec():
            name = dlg.inp_name.text().strip()
            mode = dlg.mode
            if mode == "import":
                self.import_excel(name)
            elif mode == "manual":
                self.create_manual(name)
                
    def import_excel(self, name):
        path_tuple = QFileDialog.getOpenFileName(self, "Import Excel", "", "Excel Files (*.xlsx *.xls)")
        path = path_tuple[0]
        if not path: return
        
        try:
            df = pd.read_excel(path, dtype=str)
            df = df.fillna("")
            
            # DATE CLEANUP on Import
            # Use regex to fix "YYYY_MM_DD..." patterns globally in the dataframe (1 or more underscores)
            df = df.replace(r"(\d{4})_+(\d{2})_+(\d{2}).*", r"\3/\2/\1", regex=True)
            
            conn = sqlite3.connect(DB_FILE)
            cur = conn.cursor()
            cur.execute("INSERT INTO old_sheets_meta (name, created_at) VALUES (?, datetime('now'))", (name,))
            sheet_id = cur.lastrowid
            
            # Save Headers
            cols = list(df.columns)
            for i, col in enumerate(cols):
                cur.execute("INSERT INTO old_sheet_headers (sheet_id, col_index, header_name) VALUES (?,?,?)", (sheet_id, i, col))
                
            # Create Table
            safe_cols = [f'"{c}" TEXT' for c in cols] # quote columns to be safe
            col_def = ", ".join(safe_cols)
            cur.execute(f"CREATE TABLE old_sheet_data_{sheet_id} ({col_def})")
            
            # Insert Data
            placeholders = ",".join(["?"] * len(cols))
            data = df.values.tolist()
            cur.executemany(f"INSERT INTO old_sheet_data_{sheet_id} VALUES ({placeholders})", data)
            
            conn.commit()
            conn.close()
            
            self.create_tab(sheet_id, name)
            self.tabs.setCurrentIndex(self.tabs.count()-1)
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", str(e))
            
    def create_manual(self, name):
        conn = sqlite3.connect(DB_FILE)
        cur = conn.cursor()
        cur.execute("INSERT INTO old_sheets_meta (name, created_at) VALUES (?, datetime('now'))", (name,))
        sheet_id = cur.lastrowid
        
        # Default 10 columns
        cols = [f"Column {i+1}" for i in range(10)]
        for i, col in enumerate(cols):
            cur.execute("INSERT INTO old_sheet_headers (sheet_id, col_index, header_name) VALUES (?,?,?)", (sheet_id, i, col))
            
        # FIX: Use standardized col_0, col_1 naming
        safe_cols = [f"col_{i} TEXT" for i in range(10)]
        col_def = ", ".join(safe_cols)
        cur.execute(f"CREATE TABLE old_sheet_data_{sheet_id} ({col_def})")
        
        conn.commit()
        conn.close()
        
        self.create_tab(sheet_id, name)
        self.tabs.setCurrentIndex(self.tabs.count()-1)

    def save_cell(self, item, table, sheet_id):
        row_id = item.data(Qt.ItemDataRole.UserRole)
        col = item.column()
        val = item.text()
        
        if row_id is None:
            return

        conn = sqlite3.connect(DB_FILE)
        try:
             # FIX: Determine actual column name in DB
            headers = conn.execute("SELECT header_name FROM old_sheet_headers WHERE sheet_id = ? ORDER BY col_index ASC", (sheet_id,)).fetchall()
            
            if col < len(headers):
                # Check actual table schema
                table_info = conn.execute(f"PRAGMA table_info(old_sheet_data_{sheet_id})").fetchall()
                real_cols = [info[1] for info in table_info]
                
                target_col = None
                
                # 1. Try Standard Schema (col_0, col_1...)
                std_col = f"col_{col}"
                if std_col in real_cols:
                    target_col = std_col
                else:
                    # 2. Legacy Fallback (Use Header Name)
                    header_name = headers[col][0]
                    target_col = header_name

                if target_col:
                    conn.execute(f"UPDATE old_sheet_data_{sheet_id} SET \"{target_col}\" = ? WHERE rowid = ?", (val, row_id))
                    conn.commit()
        except Exception as e:
            print(f"Save Error: {e}")
        finally:
            conn.close()

    def rename_header_slot(self, idx, table, sheet_id):
        old_name = table.horizontalHeaderItem(idx).text()
        text, ok = QInputDialog.getText(self, "Rename Column", "New Name:", text=old_name)
        if ok and text:
            try:
                conn = sqlite3.connect(DB_FILE)
                conn.execute("UPDATE old_sheet_headers SET header_name = ? WHERE sheet_id = ? AND col_index = ?", (text, sheet_id, idx))
                # FIX: Do NOT alter table column names. Keep them as col_X or original.
                # conn.execute(f"ALTER TABLE old_sheet_data_{sheet_id} RENAME COLUMN \"{old_name}\" TO \"{text}\"")
                conn.commit()
                conn.close()
                table.horizontalHeaderItem(idx).setText(text)
            except Exception as e:
                print(e)
                QMessageBox.warning(self, "Error", f"Could not rename column: {e}")

    def add_row(self):
        current_widget = self.tabs.currentWidget()
        if not current_widget: return
        table = current_widget.findChild(QTableWidget)
        if not table: return
        
        self.parent_app.save_undo_state(table)
        
        # Find sheet_id - we need to look it up or store it. 
        # The tab text is the name. But better to store ID.
        # Current implementation of create_tab doesn't store ID on the widget.
        # Let's rely on the fact that tabs are created in order or stored in a way we can retrieve.
        # Actually create_tab connects signals with sheet_id.
        # We can store sheet_id property on the table.
        sheet_id = table.property("sheet_id")
        
        conn = sqlite3.connect(DB_FILE)
        try:
            cur = conn.cursor()
            cur.execute(f"INSERT INTO old_sheet_data_{sheet_id} DEFAULT VALUES")
            row_id = cur.lastrowid
            conn.commit()
            
            row_idx = table.rowCount()
            table.insertRow(row_idx)
            cols = table.columnCount()
            for c in range(cols):
                it = QTableWidgetItem("")
                it.setData(Qt.ItemDataRole.UserRole, row_id)
                table.setItem(row_idx, c, it)
                
            # Apply delegate
            delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", table)
            for c in range(cols):
                table.setItemDelegateForColumn(c, delegate)
                
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))
        finally:
            conn.close()

    def delete_row(self):
        current_widget = self.tabs.currentWidget()
        if not current_widget: return
        table = current_widget.findChild(QTableWidget)
        if not table: return
        
        self.parent_app.save_undo_state(table)
        
        selection = table.selectionModel().selectedRows()
        if not selection:
            QMessageBox.warning(self, "Selection", "Please select row(s) to delete.")
            return
            
        sheet_id = table.property("sheet_id")
        
        rows_to_delete = sorted([idx.row() for idx in selection], reverse=True)
        row_ids_to_delete = []
        
        for r in rows_to_delete:
            item = table.item(r, 0)
            if item:
                row_ids_to_delete.append(item.data(Qt.ItemDataRole.UserRole))
        
        reply = QMessageBox.question(self, "Confirm Delete", f"Are you sure you want to delete {len(rows_to_delete)} row(s)?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            conn = sqlite3.connect(DB_FILE)
            try:
                for rid in row_ids_to_delete:
                    if rid is not None:
                        conn.execute(f"DELETE FROM old_sheet_data_{sheet_id} WHERE rowid = ?", (rid,))
                conn.commit()
                
                # Remove rows from table
                for r in rows_to_delete:
                    table.removeRow(r)
                    
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))
            finally:
                conn.close()

    def import_database(self):
        """Overwrites the current database with an external .db file."""
        msg = "You are about to OVERWRITE the current database with an existing .db file.\n\n" \
              "Requirements:\n" \
              "- The file must be a valid SQLite database.\n" \
              "- All current data will be replaced.\n\n" \
              "A safety backup will be created before proceeding. Continue?"
        
        reply = QMessageBox.warning(self, "Import Database", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.No:
            return

        path = QFileDialog.getOpenFileName(self, "Select Database File", "", "SQLite Database (*.db)")[0]
        if not path:
            return

        # 1. Safety Backup
        try:
            self.log_message("Creating pre-import backup...")
            self.manual_backup(tag="pre_db_import")
            # Wait a moment for backup logic? manual_backup is async thread.
            # We want to be safe. Since manual_backup spawns a thread, we can't strict wait unless we join.
            # But the backup thread copies READ-ONLY from DB usually.
            # However, if we overwrite immediately, the backup might read the NEW file or corrupt partial.
            # CRITICAL: We need Synchronous Backup here.
            
            # Inline Synchronous Backup for Safety
            if not os.path.exists(self.backup_dir): os.makedirs(self.backup_dir)
            pre_name = f"backup_pre_db_import_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.db"
            shutil.copy2(DB_FILE, os.path.join(self.backup_dir, pre_name))
            self.log_message(f"Safety backup created: {pre_name}")
            
        except Exception as e:
            QMessageBox.critical(self, "Backup Status", f"Could not create safety backup. Operation aborted.\nError: {e}")
            return

        # 2. Overwrite Logic
        try:
            # Copy new file over existing DB_FILE
            shutil.copy2(path, DB_FILE)
            self.log_message(f"Database replaced with: {path}")
            
            # 3. Reload App State
            self.init_db_and_migrate()
            self.load_data()
            self.refresh_master_lists_action()
            
            QMessageBox.information(self, "Import Success", "Database imported successfully.\nThe application has been reloaded.")
            
        except Exception as e:
            self.log_message(f"Import Error: {e}")
            QMessageBox.critical(self, "Import Failed", f"Could not replace database file:\n{e}\n\nCheck if the file is in use.")

    def add_column(self):
        current_widget = self.tabs.currentWidget()
        if not current_widget: return
        table = current_widget.findChild(QTableWidget)
        if not table: return
        
        self.parent_app.save_undo_state(table)
        sheet_id = table.property("sheet_id")
        
        text, ok = QInputDialog.getText(self, "Add Column", "Column Name:")
        if ok and text:
            text = text.strip()
            if not text: return
            
            # Check for duplicate
            for i in range(table.columnCount()):
                if table.horizontalHeaderItem(i).text() == text:
                    QMessageBox.warning(self, "Error", "Column with this name already exists.")
                    return
            
            try:
                conn = sqlite3.connect(DB_FILE)
                # Get next col index
                cur = conn.cursor()
                cur.execute("SELECT MAX(col_index) FROM old_sheet_headers WHERE sheet_id = ?", (sheet_id,))
                res = cur.fetchone()
                next_idx = (res[0] if res[0] is not None else -1) + 1
                
                # Update DB
                conn.execute("INSERT INTO old_sheet_headers (sheet_id, col_index, header_name) VALUES (?,?,?)", (sheet_id, next_idx, text))
                conn.execute(f"ALTER TABLE old_sheet_data_{sheet_id} ADD COLUMN \"{text}\" TEXT")
                conn.commit()
                conn.close()
                
                # Update UI
                col_idx = table.columnCount()
                table.insertColumn(col_idx)
                table.setHorizontalHeaderItem(col_idx, QTableWidgetItem(text))
                
                # Add delegate
                delegate = EditorDelegate(self.parent_app.lisan_font_family, "text", table)
                table.setItemDelegateForColumn(col_idx, delegate)
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not add column: {e}")

    def delete_column(self):
        current_widget = self.tabs.currentWidget()
        if not current_widget: return
        table = current_widget.findChild(QTableWidget)
        if not table: return
        
        col_idx = table.currentColumn()
        if col_idx < 0:
            QMessageBox.warning(self, "Selection", "Please select a column to delete.")
            return
            
        sheet_id = table.property("sheet_id")
        header_name = table.horizontalHeaderItem(col_idx).text()
        
        reply = QMessageBox.question(self, "Delete Column", f"Are you sure you want to delete column '{header_name}'? This cannot be undone.", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                conn = sqlite3.connect(DB_FILE)
                # Drop column from data table
                conn.execute(f"ALTER TABLE old_sheet_data_{sheet_id} DROP COLUMN \"{header_name}\"")
                
                # Update headers meta: delete entry and shift indices
                conn.execute("DELETE FROM old_sheet_headers WHERE sheet_id = ? AND col_index = ?", (sheet_id, col_idx))
                conn.execute("UPDATE old_sheet_headers SET col_index = col_index - 1 WHERE sheet_id = ? AND col_index > ?", (sheet_id, col_idx))
                
                conn.commit()
                conn.close()
                
                table.removeColumn(col_idx)
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not delete column: {e}")

    def delete_sheet(self):
        idx = self.tabs.currentIndex()
        if idx < 0: return
        
        name = self.tabs.tabText(idx)
        current_widget = self.tabs.currentWidget()
        table = current_widget.findChild(QTableWidget)
        sheet_id = table.property("sheet_id")
        
        reply = QMessageBox.question(self, "Delete Sheet", f"Are you sure you want to delete sheet '{name}'? This cannot be undone.", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                conn = sqlite3.connect(DB_FILE)
                conn.execute(f"DROP TABLE IF EXISTS old_sheet_data_{sheet_id}")
                conn.execute("DELETE FROM old_sheets_meta WHERE id = ?", (sheet_id,))
                conn.execute("DELETE FROM old_sheet_headers WHERE sheet_id = ?", (sheet_id,))
                conn.commit()
                conn.close()
                self.tabs.removeTab(idx)
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))

    def update_font_size(self, size):
        for i in range(self.tabs.count()):
            tab = self.tabs.widget(i)
            table = tab.findChild(QTableWidget)
            if table and hasattr(self.parent_app, 'lisan_font_family'):
                table.setFont(QFont(self.parent_app.lisan_font_family, int(size)))

    def header_menu(self, pos, table, sheet_id):
        idx = table.horizontalHeader().logicalIndexAt(pos)
        menu = QMenu()
        rename_act = menu.addAction("Rename Column")
        act = menu.exec(table.horizontalHeader().mapToGlobal(pos))
        
        if act == rename_act:
            text, ok = QInputDialog.getText(self, "Rename Column", "New Name:", text=table.horizontalHeaderItem(idx).text())
            if ok and text:
                old_name = table.horizontalHeaderItem(idx).text()
                try:
                    conn = sqlite3.connect(DB_FILE)
                    conn.execute("UPDATE old_sheet_headers SET header_name = ? WHERE sheet_id = ? AND col_index = ?", (text, sheet_id, idx))
                    # Rename column in data table
                    conn.execute(f"ALTER TABLE old_sheet_data_{sheet_id} RENAME COLUMN \"{old_name}\" TO \"{text}\"")
                    conn.commit()
                    conn.close()
                    table.horizontalHeaderItem(idx).setText(text)
                except Exception as e:
                    print(e)
                    
    def export_selection_current_sheet(self):
        current_widget = self.tabs.currentWidget()
        if not current_widget: return
        table = current_widget.findChild(QTableWidget)
        if not table: return
        
        rows = sorted(list(set(index.row() for index in table.selectedIndexes())))
        if not rows:
            QMessageBox.warning(self, "Export", "Please select rows to export.")
            return

        # Get Headers
        headers = []
        for c in range(table.columnCount()):
            headers.append(table.horizontalHeaderItem(c).text())
            
        dlg = ColumnSelectionDialog(headers, self, show_tracks_option=False)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_indices = dlg.get_selected_indices()
            if not selected_indices:
                QMessageBox.warning(self, "Export", "No columns selected.")
                return
                
            path = QFileDialog.getSaveFileName(self, "Export Selection", "OldSheet_Export.xlsx", "Excel Files (*.xlsx)")[0]
            if not path:
                return
                
            try:
                data = []
                final_headers = [headers[i] for i in selected_indices]
                
                for r in rows:
                    row_data = []
                    for i in selected_indices:
                        it = table.item(r, i)
                        row_data.append(it.text() if it else "")
                    data.append(row_data)
                
                df = pd.DataFrame(data, columns=final_headers)
                df.to_excel(path, index=False)
                QMessageBox.information(self, "Export", "Export successful.")
            except Exception as e:
                QMessageBox.critical(self, "Error", str(e))


# --- SPLASH SCREEN & DATA LOADER ---

class DataLoader(QThread):
    finished = pyqtSignal(object) # Returns DataFrame
    
    def run(self):
        try:
            conn = sqlite3.connect(DB_FILE)
            # Fetch ALL data sorted by Audio No ASC
            df = pd.read_sql_query("SELECT * FROM events ORDER BY CAST(audio_no AS INTEGER) ASC", conn)
            conn.close()
            self.finished.emit(df)
        except Exception as e:
            print(f"Loading Error: {e}", flush=True)
            traceback.print_exc()
            self.finished.emit(pd.DataFrame())

            self.finished.emit(pd.DataFrame())

    # clear_filters REMOVED from here (wrong class)

class ModernSplash(QSplashScreen):
    def __init__(self):
        final_path = SPLASH_PATH
        img = QImage()
        if final_path:
            img.load(final_path)
            
        if img.isNull():
            pixmap = QPixmap(600, 400)
            pixmap.fill(QColor("#1a202c"))
        else:
            # High Quality Scaling - Reduced Size as requested or default (Original 960x600 -> New ~500x312)
            img = img.scaled(500, 312, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
            pixmap = QPixmap.fromImage(img)
            
        final_pixmap = QPixmap(pixmap.size())
        final_pixmap.fill(Qt.GlobalColor.transparent)
        painter = QPainter(final_pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        
        # 1. Draw rounded clipping path
        path = QPainterPath()
        path.addRoundedRect(QRectF(pixmap.rect()).adjusted(1, 1, -1, -1), 15, 15)
        painter.setClipPath(path)
        
        # 2. Draw base pixmap
        painter.drawPixmap(0, 0, pixmap)
        
        # 3. Brightness Boost (Overlay method - extremely fast)
        # Using a white overlay with ~10% opacity boosts brightness without being dull
        painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_Screen)
        painter.setBrush(QColor(255, 255, 255, 30)) # Low opacity white
        painter.setPen(Qt.PenStyle.NoPen)
        painter.drawRect(pixmap.rect())
        
        painter.end()
        
        super().__init__(final_pixmap)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setWindowOpacity(0.0)
        
        self.fade_in_anim = QPropertyAnimation(self, b"windowOpacity")
        self.fade_out_anim = QPropertyAnimation(self, b"windowOpacity")
        self.main_window_callback = None
        
        # --- PROGRESS BAR UI ---
        # We need a layout on the splash screen. 
        # Since QSplashScreen is a widget, we can set a layout.
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 30) # Bottom margin for bar
        layout.addStretch() # Push everything down
        
        # Status Label
        self.status_label = QLabel("Starting...")
        self.status_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.status_label.setStyleSheet("""
            color: #e2e8f0; 
            font-weight: bold; 
            font-size: 12px;
            background-color: rgba(15, 23, 42, 180);
            border-radius: 4px;
            padding: 2px 8px;
        """)
        # Wrap label in a container to center it horizontally without stretching full width if preferred,
        # but centered text is fine.
        layout.addWidget(self.status_label, 0, Qt.AlignmentFlag.AlignCenter)
        
        layout.addSpacing(5)
        
        # Progress Bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        
        # Cyan Style matching Dark Theme
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: none;
                background-color: rgba(255, 255, 255, 50);
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #38B2AC;
                border-radius: 3px;
            }
        """)
        
        layout.addWidget(self.progress_bar)

    def update_progress(self, val, msg=""):
        self.progress_bar.setValue(val)
        if msg:
            self.status_label.setText(msg)
        QApplication.processEvents()

    def fade_in(self):
        self.fade_in_anim.setDuration(1000)
        self.fade_in_anim.setStartValue(0.0)
        self.fade_in_anim.setEndValue(1.0)
        self.fade_in_anim.setEasingCurve(QEasingCurve.Type.OutQuad)
        self.fade_in_anim.start()

    def fade_out(self, callback):
        self.main_window_callback = callback
        self.fade_out_anim.setDuration(800)
        self.fade_out_anim.setStartValue(1.0)
        self.fade_out_anim.setEndValue(0.0)
        self.fade_out_anim.setEasingCurve(QEasingCurve.Type.InQuad)
        self.fade_out_anim.finished.connect(self.on_fade_out_finished)
        self.fade_out_anim.start()

    def on_fade_out_finished(self):
        if self.main_window_callback:
            self.main_window_callback()
        self.close()


class ColumnSelectionDialog(QDialog):
    def __init__(self, headers, parent=None, show_tracks_option=True, default_selection=None):
        super().__init__(parent)
        self.setWindowTitle("Select Columns to Export")
        self.setFixedSize(400, 500)
        self.setStyleSheet("background-color: white; color: #2d3748;")
        self.headers = headers
        self.result_mode = "selected" # Default mode
        self.show_tracks_option = show_tracks_option
        
        layout = QVBoxLayout(self)
        
        lbl = QLabel("Choose columns to include in the export:")
        lbl.setStyleSheet("font-weight: bold; font-size: 14px;")
        layout.addWidget(lbl)
        
        # Helper Buttons
        h_btn_layout = QHBoxLayout()
        btn_sel_all = QPushButton("Select All")
        btn_sel_all.setObjectName("SecondaryBtn")
        btn_sel_all.clicked.connect(self.select_all)
        
        btn_desel_all = QPushButton("Deselect All")
        btn_desel_all.setObjectName("SecondaryBtn")
        btn_desel_all.clicked.connect(self.deselect_all)
        
        h_btn_layout.addWidget(btn_sel_all)
        h_btn_layout.addWidget(btn_desel_all)
        layout.addLayout(h_btn_layout)
        
        # List of Checkboxes
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_content)
        self.checkboxes = []
        
        for i, h in enumerate(headers):
            chk = QCheckBox(h)
            # Default Selection Logic
            if default_selection:
                 if h in default_selection:
                     chk.setChecked(True)
                 else:
                     chk.setChecked(False)
            else:
                chk.setChecked(True) # Default all if no list provided
            
            self.checkboxes.append(chk)
            self.scroll_layout.addWidget(chk)
            
        self.scroll_layout.addStretch()
        self.scroll.setWidget(self.scroll_content)
        layout.addWidget(self.scroll)
        
        # Tracks Option
        self.chk_include_tracks = None
        if self.show_tracks_option:
            self.chk_include_tracks = QCheckBox("Stats: Include Audio Tracks Data (Creates 2nd Sheet)")
            self.chk_include_tracks.setStyleSheet("font-weight: bold; color: #2b6cb0;")
            layout.addWidget(self.chk_include_tracks)
        
        # Action Buttons
        btn_layout = QVBoxLayout()
        btn_layout.setSpacing(10)
        
        self.btn_distinct = QPushButton("Export with Distinct")
        self.btn_distinct.setObjectName("PrimaryBtn")
        self.btn_distinct.setStyleSheet("background-color: #2b6cb0; color: white; font-weight: bold; padding: 6px;")
        self.btn_distinct.clicked.connect(self.accept_distinct)
        
        self.btn_selected = QPushButton("Export Selected Columns")
        self.btn_selected.setObjectName("PrimaryBtn")
        # self.btn_selected.setStyleSheet("background-color: #2b6cb0; color: white; font-weight: bold; padding: 6px;")
        self.btn_selected.clicked.connect(self.accept_selected)
        
        btn_cancel = QPushButton("Cancel")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(self.btn_distinct)
        btn_layout.addWidget(self.btn_selected)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
        
    def select_all(self):
        for chk in self.checkboxes: chk.setChecked(True)
        
    def deselect_all(self):
        for chk in self.checkboxes: chk.setChecked(False)
        
    def accept_distinct(self):
        self.result_mode = "distinct"
        self.accept()
        
    def accept_selected(self):
        self.result_mode = "selected"
        self.accept()
        
    def get_selected_indices(self):
        return [i for i, chk in enumerate(self.checkboxes) if chk.isChecked()]
            
    def is_tracks_included(self):
        if self.chk_include_tracks:
            return self.chk_include_tracks.isChecked()
        return False

# --- 4. MAIN APPLICATION WINDOW ---
class ArchiveApp(QMainWindow):
    def __init__(self, splash=None):
        super().__init__()
        
        if splash: splash.update_progress(10, "Initializing UI...")
        self.setWindowTitle("DH Archive")
        self.setMinimumSize(1300, 900)
        self.undo_stack = []
        self.unsaved_changes = False
        self.sidebar_collapsed = True
        self.open_tracks = {}
        self.master_locked = True # Default Locked
        
        # Pagination / Loading State
        self.is_loading_more = False
        self.current_page = 1
        self.rows_per_page = 1000 # 1000 rows as requested
        self.total_records = 0
        self.total_pages = 1
        
        # --- SETTINGS VARIABLES ---
        self.is_dark_mode = False
        self.current_font_size = 14
        self.auto_backup = True # Default On
        self.confirm_delete = True # Default On
        self.backup_dir = DEFAULT_BACKUP_DIR # Init with default
        self.startup_view = 0 # Default: Master Sheet (0)
        self.saved_col_widths = []
        self.row_height_cache = {}
        self._is_loading_data = False  # Flag to prevent resize during load
        self._last_column_widths = []  # Track column widths to detect actual changes
        
        # Debounced Resize for Master - 1000ms to wait until user finishes dragging
        self.master_resize_timer = QTimer(self)
        self.master_resize_timer.setSingleShot(True)
        self.master_resize_timer.setInterval(1000)  # Wait 1 second after last resize event
        self.master_resize_timer.timeout.connect(self.resize_master_rows)
        # We will connect this in init_ui or similar once table is active
        
        self.alternating_rows = False # Default Off
        
        # Smart Data Entry Flags
        self.repeat_person = "Empty"
        self.repeat_occasion = "Empty"
        self.repeat_category = "Empty"
        self.repeat_place = "Empty"
        self.repeat_date = "Empty"
        
        # Search variables
        self.highlight_search = True
        self.current_search_term = ""
        
        self.lisan_font_family = "Arial"
        if os.path.exists(FONT_PATH):
            fid = QFontDatabase.addApplicationFont(FONT_PATH)
            if fid != -1:
                self.lisan_font_family = QFontDatabase.applicationFontFamilies(fid)[0]
                
        self.master_lists = {}
        self.row_height_cache = {} # audio_no -> height
        self.global_filters_dirty = {} # tab_index -> bool
        
        # Ensure keys exist
        for i in [2, 3, 4, 5, 6, 13, 16]:
            self.master_lists[i] = []

        
        # Initialize Defaults (Before load_settings)
        self.repeat_person = "Empty"
        self.repeat_occasion = "Empty"
        self.repeat_category = "Empty"
        self.repeat_place = "Empty"
        self.repeat_date = "Empty"
        self.repeat_outof = "Empty"
        self.repeat_a_v = "Empty"
        
        self.repeat_a_v = "Empty"
        
        if splash: splash.update_progress(30, "Connecting Database...")
        self.init_db_and_migrate()
        
        if splash: splash.update_progress(50, "Caching Data...")
        self.refresh_master_cache() # Moved before setup_main_area to populate dropdowns
        
        if splash: splash.update_progress(60, "Loading Settings...")
        self.load_settings() # LOAD SETTINGS ON STARTUP
        self.setup_styles()
        
        # Start Backup Background Worker if enabled
        if self.auto_backup:
            self.backup_worker = BackupWorker(self.backup_dir)
            self.backup_worker.finished.connect(self.log_message)
            self.backup_worker.start()
        
        central = QWidget()
        self.setCentralWidget(central)
        self.main_layout = QHBoxLayout(central)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        if splash: splash.update_progress(70, "Building Interface...")
        self.setup_sidebar()
        self.setup_main_area()
        
        if splash: splash.update_progress(90, "Loading Data...")
        self.load_data()
        self.switch_page(self.startup_view) # Auto-switch to preferred start page
        
        if splash: splash.update_progress(100, "Ready!")
        
        self.table.itemChanged.connect(self.handle_item_changed)
        self.table.currentCellChanged.connect(self.handle_cell_clicked)
       
        # self.table.verticalScrollBar().valueChanged.connect(self.check_scroll_load)
        h_bar = self.table.horizontalScrollBar()

        
        self.folder_shortcut = QShortcut(QKeySequence("Ctrl+Shift+O"), self)
        self.folder_shortcut.activated.connect(self.open_entry_folder)
        
        # UNDO SHORTCUT
        self.undo_shortcut_key = QShortcut(QKeySequence.StandardKey.Undo, self)
        self.undo_shortcut_key.activated.connect(self.perform_undo)
        
        # Apply visual settings initially
        self.apply_visual_settings()
        
        # --- NEW SAVING LOGIC ---
        self.shortcut_save = QShortcut(QKeySequence("Ctrl+S"), self)
        self.shortcut_save.activated.connect(self.save_changes)
        
        # Auto-Update Check
        self.check_for_updates(silent=True)

    def check_for_updates(self, silent=False):
        """Runs the update checker in background."""
        self.update_thread = UpdateChecker()
        self.update_thread.update_available.connect(self.on_update_available)
        if not silent:
            # We use lambda for simple slots, but be careful with cleanup if needed
            self.update_thread.no_update.connect(lambda: QMessageBox.information(self, "No Updates", "You are on the latest version."))
            self.update_thread.error_occurred.connect(lambda e: QMessageBox.warning(self, "Update Error", f"Could not check for updates:\n{e}"))
        self.update_thread.start()
        
    def on_update_available(self, new_version, notes, download_url):
        """Show dialog to user."""
        msg = QMessageBox(self)
        msg.setWindowTitle("Update Available")
        msg.setText(f"A new version (v{new_version}) is available!")
        msg.setInformativeText(f"Release Notes:\n{notes}\n\nDo you want to update now?")
        msg.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        msg.setDefaultButton(QMessageBox.StandardButton.Yes)
        
        if msg.exec() == QMessageBox.StandardButton.Yes:
            self.start_update_download(download_url)
            
    def start_update_download(self, url):
        """Starts downloading the update."""
        self.progress_dialog = QProgressDialog("Downloading Update...", "Cancel", 0, 100, self)
        self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
        self.progress_dialog.show()
        
        self.downloader = UpdateDownloader(url)
        self.downloader.progress.connect(self.progress_dialog.setValue)
        self.downloader.finished.connect(self.on_download_finished)
        self.downloader.error.connect(lambda e: QMessageBox.critical(self, "Download Error", f"Failed to download:\n{e}"))
        self.downloader.start()
        
    def on_download_finished(self, zip_path):
        """Install the update."""
        self.progress_dialog.close()
        
        # Confirm Installation
        reply = QMessageBox.question(self, "Install Update", 
                                   "Download complete. The app must restart to apply the update.\nProceed?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                                   
        if reply == QMessageBox.StandardButton.Yes:
            success, msg = install_update(zip_path)
            if not success:
                QMessageBox.critical(self, "Installation Failed", msg)
            else:
                # App should close now
                QApplication.quit()

    @contextlib.contextmanager
    def busy_cursor(self, button=None):
        """Standard indicator for long operations."""
        # Disable button to prevent double-click
        if button and hasattr(button, 'setEnabled'):
            button.setEnabled(False)
            
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            yield
        finally:
            QApplication.restoreOverrideCursor()
            if button and hasattr(button, 'setEnabled'):
                button.setEnabled(True)

    def mark_unsaved(self):
        if not self.unsaved_changes:
            self.unsaved_changes = True
            self.setWindowTitle(self.windowTitle() + " *")

    def save_changes(self):
        # LOCK GUARD: Absolutely no saving when locked
        if hasattr(self, 'master_locked') and self.master_locked:
            self.log_message("Save Blocked: Master is Locked")
            return

        self.save_data(silent=False)
        # Also save all "Other Sheets"
        if hasattr(self, 'old_sheets_page'):
            self.old_sheets_page.save_all_sheets()
            
        self.unsaved_changes = False
        current_title = self.windowTitle().split(" *")[0]
        self.setWindowTitle(current_title)

    def check_unsaved_changes(self):
        """Returns True if safe to proceed, False if user cancelled."""
        if not self.unsaved_changes:
            return True
        
        reply = QMessageBox.question(self, "Unsaved Changes", 
                                     "You have unsaved changes. Do you want to save them?",
                                     QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Discard | QMessageBox.StandardButton.Cancel)
                                     
        if reply == QMessageBox.StandardButton.Save:
            self.save_changes()
            return True
        elif reply == QMessageBox.StandardButton.Discard:
            self.unsaved_changes = False
            current_title = self.windowTitle().replace(" *", "")
            self.setWindowTitle(current_title)
            return True
        else:
            return False

    def closeEvent(self, event):
        if not self.check_unsaved_changes():
            event.ignore()
            return

        # Save UI State Persistence
        try:
            # 1. Column Widths for ALL Tabs
            all_widths = {}
            for i in range(self.sheets_tabs.count()):
                tab_name = self.sheets_tabs.tabText(i)
                table = self.sheets_tabs.widget(i)
                if hasattr(table, "columnCount"):
                    cols = []
                    for c in range(table.columnCount()):
                        cols.append(table.columnWidth(c))
                    all_widths[tab_name] = cols
            
            self.save_setting("ui_column_widths_json", json.dumps(all_widths))
            
            # 2. Row Height Cache (Master Only)
            if hasattr(self, 'row_height_cache') and self.row_height_cache:
                self.save_setting("ui_row_height_cache_json", json.dumps(self.row_height_cache))
                
        except Exception as e:
            print(f"Error saving UI state on close: {e}")

        super().closeEvent(event)

    def log_message(self, msg):
        """Standardizes logging to the debug console."""
        timestamp = datetime.now().strftime("%H:%M:%S")
        if hasattr(self, 'debug_console'):
            self.debug_console.append(f"[{timestamp}] {msg}")


    # --- MASTER LOCK LOGIC ---
    def toggle_master_lock(self):
        """Toggles the Master Database Lock."""
        # If currently locked, user wants to UNLOCK
        if self.btn_master_lock.isChecked():
            # Was unlocked, now locking
            # Check unsaved changes before locking
            if self.unsaved_changes:
                reply = QMessageBox.question(self, "Unsaved Changes", 
                                             "You have unsaved changes. Closing the Master Lock will save them. Continue?", 
                                             QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Discard | QMessageBox.StandardButton.Cancel)
                
                if reply == QMessageBox.StandardButton.Cancel:
                    self.btn_master_lock.setChecked(False) # Keep Unlocked
                    return
                elif reply == QMessageBox.StandardButton.Save:
                    self.save_changes()
                else: 
                    self.unsaved_changes = False # Discard
                    # DISCARD LOGIC: Reload to visually revert changes
                    self.load_data()
            
            # Check Other Sheets unsaved?
            
            # Check Other Sheets unsaved?
            # They use explicit save now, but if we want to be safe...
            # For now, we assume explicit save button usage on Other Sheets.
            
            self.master_locked = True
            self.btn_master_lock.setText("🔒 MASTER LOCKED")
            self.btn_master_lock.setStyleSheet("background-color: #742a2a; color: white; font-weight: bold; border: 2px solid #e53e3e; border-radius: 6px;")
            self.update_lock_state()
            
        else:
            # Was locked, now unlocking
            # Ask confirmation
            reply = QMessageBox.question(self, "Unlock Database", "Are you sure you want to UNLOCK the Master Database?\nThis allows editing and modification.", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                self.master_locked = False
                self.btn_master_lock.setText("🔓 MASTER UNLOCKED")
                self.btn_master_lock.setStyleSheet("background-color: #2f855a; color: white; font-weight: bold; border: 2px solid #48bb78; border-radius: 6px;")
                self.update_lock_state()
            else:
                 self.btn_master_lock.setChecked(True) # Keep Locked

    def update_lock_state(self):
        """Enforces lock state across the app."""
        if self.master_locked:
            # Master Sheet
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            self.btn_add.setEnabled(False)
            # self.btn_save.setEnabled(False) # Maybe keep save enabled if there were changes? No, lock implies saved.
            
            # Other Sheets (Global Lock)
            if hasattr(self, 'other_sheets_manager'):
                self.other_sheets_manager.btn_lock.setEnabled(False)
                self.other_sheets_manager.reset_lock() # Force lock
            
            # Master Lists Global Lock
            if hasattr(self, 'btn_master_lists_lock'):
                self.btn_master_lists_lock.setEnabled(False)
                self.btn_master_lists_lock.setChecked(True)
                self.btn_master_lists_lock.setText("🔒 Locked")
                self.toggle_master_lists_global_lock(True)

            # Mappings Global Lock
            if hasattr(self, 'btn_mappings_lock'):
                self.btn_mappings_lock.setEnabled(False)
                self.btn_mappings_lock.setChecked(True)
                self.btn_mappings_lock.setText("🔒 Locked")
                self.toggle_mappings_global_lock(True)

        else:
            # Master Sheet
            self.table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.SelectedClicked | QAbstractItemView.EditTrigger.AnyKeyPressed)
            self.btn_add.setEnabled(True)
            
            # Other Sheets
            if hasattr(self, 'other_sheets_manager'):
                self.other_sheets_manager.btn_lock.setEnabled(True)
            
            # Restore Master Lists & Mappings locks
            if hasattr(self, 'btn_master_lists_lock'):
                self.btn_master_lists_lock.setEnabled(True)
            
            if hasattr(self, 'btn_mappings_lock'):
                self.btn_mappings_lock.setEnabled(True)

        # Propagate Lock to active widgets (Tracks Widget)
        for r in range(self.table.rowCount()):
            cw = self.table.cellWidget(r, 0)
            if hasattr(cw, 'set_locked'):
                cw.set_locked(self.master_locked)

    def save_undo_state(self, target_table=None):
        """Saves current table state (DB + UI) to undo stack."""
        if target_table is None: target_table = self.table
        
        is_master = (target_table == self.table)
        sheet_id = target_table.property("sheet_id")
        
        state_data = []
        for r in range(target_table.rowCount()):
            if is_master:
                cw = target_table.cellWidget(r, 0)
                if isinstance(cw, TrackManagerWidget): continue
            
            # Get ROWID/ID
            row_item = target_table.item(r, 1 if is_master else 0)
            row_id = row_item.data(Qt.ItemDataRole.UserRole) if row_item else None
            
            row_vals = []
            start_col = 1 if is_master else 0
            for c in range(start_col, target_table.columnCount()):
                it = target_table.item(r, c)
                row_vals.append(it.text() if it else "")
                
            state_data.append({"id": row_id, "vals": row_vals})
            
        self.undo_stack.append({
            "type": "master" if is_master else "other",
            "sheet_id": sheet_id,
            "data": state_data
        })
        if len(self.undo_stack) > 20:
            self.undo_stack.pop(0)

    def perform_undo(self):
        """Restores the table and database to the previous state."""
        if not self.undo_stack:
            return
            
        last_state = self.undo_stack.pop()
        state_type = last_state.get("type", "master")
        state = last_state.get("data", [])
        sheet_id = last_state.get("sheet_id")
        
        conn = sqlite3.connect(DB_FILE)
        try:
            cursor = conn.cursor()
            cursor.execute("BEGIN TRANSACTION")
            
            if state_type == "master":
                cursor.execute("DELETE FROM events")
                for row in state:
                    vals = row["vals"]
                    placeholders = ",".join(["?"] * 16)
                    cursor.execute(f"INSERT INTO events VALUES ({placeholders})", tuple(vals))
                cursor.execute("COMMIT")
                conn.close()
                self.load_data()
                self.log_message("Undo: Master Sheet restored.")
            else:
                # Restore Other Sheet
                cursor.execute(f"DELETE FROM old_sheet_data_{sheet_id}")
                for row in state:
                    rid = row["id"]
                    vals = row["vals"]
                    placeholders = ",".join(["?"] * len(vals))
                    # Check columns
                    table_info = cursor.execute(f"PRAGMA table_info(old_sheet_data_{sheet_id})").fetchall()
                    col_names = [info[1] for info in table_info]
                    col_str = ", ".join([f'"{c}"' for c in col_names[:len(vals)]])
                    
                    if rid is not None:
                        # Restore with original ROWID
                        cursor.execute(f"INSERT INTO old_sheet_data_{sheet_id}(rowid, {col_str}) VALUES (?, {placeholders})", (rid, *vals))
                    else:
                        cursor.execute(f"INSERT INTO old_sheet_data_{sheet_id}({col_str}) VALUES ({placeholders})", tuple(vals))
                
                cursor.execute("COMMIT")
                conn.close()
                
                # Refresh UI
                if hasattr(self, 'old_sheets_page'):
                    self.old_sheets_page.load_sheets()
                self.log_message(f"Undo: Other Sheet {sheet_id} restored.")
                
            self.mark_unsaved()
        except Exception as e:
            self.log_message(f"Undo Error: {e}")
            if conn: conn.rollback()
        finally:
            if conn: conn.close()

    def init_db_and_migrate(self):
        if not self.check_unsaved_changes():
             return
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        # Migration: tracks_summary -> Tracks, av -> "AV"
        try:
            # Check if old columns exist
            c.execute("PRAGMA table_info(events)")
            cols = [row[1] for row in c.fetchall()]
            if "tracks_summary" in cols:
                c.execute("ALTER TABLE events RENAME COLUMN tracks_summary TO Tracks")
            if "av" in cols:
                c.execute("ALTER TABLE events RENAME COLUMN av TO \"AV\"")
            conn.commit()
        except Exception as e:
            print(f"Migration error: {e}")

        c.execute("CREATE TABLE IF NOT EXISTS events (audio_no INTEGER PRIMARY KEY, person TEXT, occasion TEXT, category TEXT, place TEXT, country TEXT, hijri_date TEXT, esavi_date TEXT, year TEXT, out_of TEXT, remarks TEXT, Tracks TEXT, AV TEXT, cass_no TEXT, came_from TEXT, incomplete TEXT)")
        c.execute("CREATE TABLE IF NOT EXISTS dropdown_options (id INTEGER PRIMARY KEY AUTOINCREMENT, category TEXT, value TEXT)")
        c.execute("CREATE TABLE IF NOT EXISTS locations (id INTEGER PRIMARY KEY AUTOINCREMENT, place TEXT, country TEXT)")
        c.execute("CREATE TABLE IF NOT EXISTS tracks (id INTEGER PRIMARY KEY AUTOINCREMENT, event_id INTEGER, track_no TEXT, track_name TEXT, FOREIGN KEY(event_id) REFERENCES events(audio_no))")
        c.execute("CREATE TABLE IF NOT EXISTS folder_mappings (id INTEGER PRIMARY KEY AUTOINCREMENT, start_no INTEGER, drive_info TEXT, folder_name TEXT, full_link TEXT, mapping_type TEXT)")
        c.execute("CREATE TABLE IF NOT EXISTS app_settings (setting_key TEXT PRIMARY KEY, setting_value TEXT)")
        # --- OLD SHEETS TABLES ---
        c.execute("CREATE TABLE IF NOT EXISTS old_sheets_meta (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, created_at TEXT)")
        c.execute("CREATE TABLE IF NOT EXISTS old_sheet_headers (sheet_id INTEGER, col_index INTEGER, header_name TEXT, PRIMARY KEY(sheet_id, col_index))")
        # -------------------------
        conn.commit()
        conn.close()

    def load_settings(self):
        """Loads settings from the DB."""
        conn = sqlite3.connect(DB_FILE)
        try:
            # Theme
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='dark_mode'").fetchone()
            if res and res[0] == 'true':
                self.is_dark_mode = True
            else:
                self.is_dark_mode = False
            
            # Font Size
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='font_size'").fetchone()
            if res:
                self.current_font_size = int(res[0])
            else:
                self.current_font_size = 14
                
            # Auto Backup
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='auto_backup'").fetchone()
            if res:
                self.auto_backup = (res[0] == 'true')
                
            # Confirm Delete
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='confirm_delete'").fetchone()
            if res:
                self.confirm_delete = (res[0] == 'true')

            # Backup Path
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='backup_path'").fetchone()
            if res and res[0]:
                self.backup_dir = res[0]

            # Startup View
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='startup_view'").fetchone()
            if res and res[0]:
                self.startup_view = int(res[0])

            # Column Widths (New Robust JSON Path)
            self.saved_ui_widths = {}
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='ui_column_widths_json'").fetchone()
            if res and res[0]:
                try:
                    self.saved_ui_widths = json.loads(res[0])
                    # Sync to master legacy width var if present
                    if "Master Sheet" in self.saved_ui_widths:
                        self.saved_col_widths = self.saved_ui_widths["Master Sheet"]
                    else:
                        self.saved_col_widths = []
                except:
                    self.saved_ui_widths = {}
                    self.saved_col_widths = []
            else:
                # Fallback to legacy single-string if JSON not found
                res_legacy = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='col_widths'").fetchone()
                if res_legacy and res_legacy[0]:
                    try: self.saved_col_widths = [int(w) for w in res_legacy[0].split(',')]
                    except: self.saved_col_widths = []
                else:
                    self.saved_col_widths = []

            # Row Height Cache Persistence
            self.row_height_cache = {}
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='ui_row_height_cache_json'").fetchone()
            if res and res[0]:
                try:
                    self.row_height_cache = json.loads(res[0])
                except Exception as e:
                    self.row_height_cache = {}

            # Highlight Search
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='highlight_search'").fetchone()
            if res:
                self.highlight_search = (res[0] == 'true')
                
            # Alternating Rows
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='alternating_rows'").fetchone()
            
            # Load Defaults
            res = conn.execute("SELECT setting_key, setting_value FROM app_settings WHERE setting_key LIKE 'default_%'").fetchall()
            for key, val in res:
                if key == "default_person": self.repeat_person = val
                elif key == "default_occasion": self.repeat_occasion = val
                elif key == "default_category": self.repeat_category = val
                elif key == "default_place": self.repeat_place = val
                elif key == "default_date": self.repeat_date = val
                elif key == "default_outof": self.repeat_outof = val
                elif key == "default_av": self.repeat_a_v = val # New Load

            if res:
                self.alternating_rows = (res[0] == 'true')

            # Smart Entry Defaults
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_person'").fetchone()
            if res: self.repeat_person = res[0]
            else: self.repeat_person = "Empty"
            
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_occasion'").fetchone()
            if res: self.repeat_occasion = res[0]
            else: self.repeat_occasion = "Empty"
            
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_category'").fetchone()
            if res: self.repeat_category = res[0]
            else: self.repeat_category = "Empty"
            
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_place'").fetchone()
            if res: self.repeat_place = res[0]
            else: self.repeat_place = "Empty"
            
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_date'").fetchone()
            if res: self.repeat_date = res[0]
            else: self.repeat_date = "Empty"

            # Out of
            res = conn.execute("SELECT setting_value FROM app_settings WHERE setting_key='repeat_outof'").fetchone()
            if res: self.repeat_outof = res[0]
            else: self.repeat_outof = "Empty"
            
        except:
            pass
        finally:
            conn.close()

    def save_setting(self, key, value):
        """Saves a single setting to the DB."""
        conn = sqlite3.connect(DB_FILE)
        conn.execute("INSERT OR REPLACE INTO app_settings (setting_key, setting_value) VALUES (?, ?)", (key, str(value)))
        conn.commit()
        conn.close()

    def setup_styles(self):
        """Unified light styling with Guaranteed Sky Blue highlights."""
        if self.is_dark_mode:
            self.setStyleSheet(STYLESHEET_DARK)
            palette = self.palette()
            palette.setColor(QPalette.ColorRole.Highlight, QColor("#2d3748")) # Dark Highlight Selection
            palette.setColor(QPalette.ColorRole.HighlightedText, QColor("white")) # Force White Text
            self.setPalette(palette)
        else:
            self.setStyleSheet(STYLESHEET_LIGHT)
            palette = self.palette()
            palette.setColor(QPalette.ColorRole.Highlight, QColor("#e3f2fd"))
            palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#2d3748"))
            self.setPalette(palette)
            
    def apply_visual_settings(self):
        """Applies visual settings like alternating rows to all tables."""
        # Main Table
        self.table.setAlternatingRowColors(self.alternating_rows)
        
        # Search Tables (if initialized)
        if hasattr(self, 'track_search_tab'):
            # Track Search results table (single table)
            table = self.track_search_tab.search_results
            if table:
                table.setAlternatingRowColors(self.alternating_rows)
                    
        if hasattr(self, 'track_search_tab'):
            self.track_search_tab.search_results.setAlternatingRowColors(self.alternating_rows)
            
        # Master Lists
        for tbl in self.master_tab_tables.values():
            tbl.setAlternatingRowColors(self.alternating_rows)
            
        # Settings Table
        self.val_table.setAlternatingRowColors(self.alternating_rows)
        
        # Open Track Managers
        for manager in self.open_tracks.values():
            manager.track_table.setAlternatingRowColors(self.alternating_rows)

    def setup_sidebar(self):
        self.sidebar = QFrame()
        self.sidebar.setObjectName("Sidebar")
        self.sidebar.setFixedWidth(60) # Start Collapsed
        layout = QVBoxLayout(self.sidebar)
        layout.setContentsMargins(10, 20, 10, 20)
        layout.setSpacing(15)
        
        # New Header Layout for Icon + Title
        header_layout = QHBoxLayout()
        
        self.btn_collapse = QPushButton("☰")
        self.btn_collapse.setObjectName("CollapseBtn")
        self.btn_collapse.setFixedSize(50, 40)
        self.btn_collapse.clicked.connect(self.toggle_sidebar)
        
        self.logo_label = QLabel("DH AV Archive")
        self.logo_label.setObjectName("SidebarTitle")
        self.logo_label.setStyleSheet("font-size: 18px; font-weight: 800; margin-bottom: 10px;")
        self.logo_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        self.logo_label.hide() # Hide Initially
        
        header_layout.addWidget(self.btn_collapse)
        header_layout.addWidget(self.logo_label)
        header_layout.addStretch()
        
        layout.addLayout(header_layout)
        
        self.btn_sheet = QPushButton("📊") # Icons only
        self.btn_sheet.setObjectName("MenuBtn")
        self.btn_sheet.setCheckable(True)
        self.btn_sheet.setChecked(True)
        self.btn_sheet.clicked.connect(lambda: self.switch_page(0))
        layout.addWidget(self.btn_sheet)
        
        # Old Sheets Button Removed
        
        self.btn_tracks = QPushButton("🎵")
        self.btn_tracks.setObjectName("MenuBtn")
        self.btn_tracks.setCheckable(True)
        self.btn_tracks.setChecked(False)
        self.btn_tracks.clicked.connect(lambda: self.switch_page(1))
        layout.addWidget(self.btn_tracks)
        
        # FIXED: Sidebar Button name to match switch logic
        self.btn_master = QPushButton("📋")
        self.btn_master.setObjectName("MenuBtn")
        self.btn_master.setCheckable(True)
        self.btn_master.setChecked(False)
        self.btn_master.clicked.connect(lambda: self.switch_page(2)) # Index 2
        layout.addWidget(self.btn_master)

        
        layout.addStretch()
        
        self.btn_settings = QPushButton("⚙️")
        self.btn_settings.setObjectName("MenuBtn")
        self.btn_settings.setCheckable(True)
        self.btn_settings.setChecked(False)
        self.btn_settings.clicked.connect(lambda: self.switch_page(3)) # Index 3
        layout.addWidget(self.btn_settings)
        
        self.debug_console = QTextEdit()
        self.debug_console.setObjectName("DebugConsole") # Removed inline style
        self.debug_console.setReadOnly(True)
        self.debug_console.setFixedHeight(150)
        self.debug_console.hide() # Hide Initially
        layout.addWidget(self.debug_console)
        
        self.main_layout.addWidget(self.sidebar)

    def toggle_sidebar(self):
        """Animates the sidebar and cleans up text to icons."""
        sw = self.sidebar.width()
        if self.sidebar_collapsed:
            ew = 260
            self.logo_label.show()
            self.debug_console.show()
            self.btn_sheet.setText("  📊  Sheets")
            # self.btn_old removed
            self.btn_tracks.setText("  🎵  Tracks")
            self.btn_master.setText("  📋  Master Lists")
            self.btn_settings.setText("  ⚙️  Settings")
        else:
            ew = 60
            self.logo_label.hide()
            self.debug_console.hide()
            # UPDATED: Just icons when collapsed
            self.btn_sheet.setText("📊")
            # self.btn_old removed
            self.btn_tracks.setText("🎵")
            self.btn_master.setText("📋")
            self.btn_settings.setText("⚙️")
            
        self.animation = QPropertyAnimation(self.sidebar, b"minimumWidth")
        self.animation.setDuration(200)
        self.animation.setStartValue(sw)
        self.animation.setEndValue(ew)
        self.animation.setEasingCurve(QEasingCurve.Type.InOutQuart)
        
        self.animation2 = QPropertyAnimation(self.sidebar, b"maximumWidth")
        self.animation2.setDuration(200)
        self.animation2.setStartValue(sw)
        self.animation2.setEndValue(ew)
        self.animation2.setEasingCurve(QEasingCurve.Type.InOutQuart)
        
        self.animation.start()
        self.animation.finished.connect(self.on_sidebar_animation_finished)
        self.animation2.start()
        self.sidebar_collapsed = not self.sidebar_collapsed

    def on_sidebar_animation_finished(self):
        """Forces update of sticky headers and table layout after sidebar toggle."""

        table = self.find_active_table()
        if table:
            table.viewport().update()

    def relock_tabs(self):
        """Forces all designated tabs to lock when navigating away or changing window focus."""
        # 1. Master Lists
        if hasattr(self, 'master_tab_buttons'):
            for cat, btns in self.master_tab_buttons.items():
                btns['lock'].setChecked(True)
                btns['lock'].setText("🔒 Locked")
                btns['lock'].setStyleSheet("")
                btns['add'].setEnabled(False)
                btns['del'].setEnabled(False)
                btns['import'].setEnabled(False)
            
            for table in self.master_tab_tables.values():
                table.set_lock_state(True)
                
        # 2. Folder Mappings
        if hasattr(self, 'b_lock_folders'):
            self.b_lock_folders.setChecked(True)
            self.b_lock_folders.setText("🔒 Locked")
            self.b_lock_folders.setStyleSheet("")
            self.fb_add.setEnabled(False)
            self.fb_del.setEnabled(False)
            self.fb_imp.setEnabled(False)
            # MasterListTable inherits QTableWidget but we added set_lock_state to it?
            # Wait, previously I saw MasterListTable def. Let's check if it has set_lock_state.
            # Yes, lines 1428.
            self.val_table.set_lock_state(True)
            
        # 3. Mp3 Mappings
        if hasattr(self, 'b_lock_mp3'):
            self.b_lock_mp3.setChecked(True)
            self.b_lock_mp3.setText("🔒 Locked")
            self.b_lock_mp3.setStyleSheet("")
            self.fb_add_mp3.setEnabled(False)
            self.fb_del_mp3.setEnabled(False)
            self.fb_imp_mp3.setEnabled(False)
            self.mp3_table.set_lock_state(True)

    def changeEvent(self, event):
        if event.type() == QEvent.Type.ActivationChange:
            if not self.isActiveWindow():
                 pass # Auto-lock disabled
        super().changeEvent(event)
            
    def switch_page(self, index): 
        # Auto-lock disabled
        self.stack.setCurrentIndex(index)
        self.btn_sheet.setChecked(index == 0)
        self.btn_tracks.setChecked(index == 1)
        # self.btn_old removed
        
        if hasattr(self, 'btn_master'):
            self.btn_master.setChecked(index == 2)
            
        self.btn_settings.setChecked(index == 3)
        
        # Load specific data for pages
        if index == 2: # Master Lists
             self.load_master_tabs_data()
        
        if index == 3: # Settings
            self.load_settings_data()

    def update_app_font_size(self, size):
        """Updates the global font size and refreshes the current view."""
        self.current_font_size = size
        self.lbl_font_size_display.setText(f"{size}px")
        
        # Save setting
        self.save_setting("font_size", str(size))
        
        # Refresh ALL views
        self.load_data()
        self.load_master_tabs_data() 
        self.load_settings_data()
        
        # Update Old Sheets if initialized
        if hasattr(self, 'old_sheets_page'):
             self.old_sheets_page.update_font_size(size)

        if hasattr(self, 'track_search_tab') and self.track_search_tab.search_input.text():
             self.track_search_tab.run_track_search()

    def toggle_theme_mode(self, is_checked):
        self.is_dark_mode = is_checked
        self.save_setting("dark_mode", "true" if is_checked else "false")
        self.setup_styles()
        # Force a refresh of specific widgets that might not auto-update style
        self.switch_page(self.stack.currentIndex())
    
    def toggle_auto_backup(self, is_checked):
        self.auto_backup = is_checked
        self.save_setting("auto_backup", "true" if is_checked else "false")
        
    def toggle_confirm_delete(self, is_checked):
        self.confirm_delete = is_checked
        self.save_setting("confirm_delete", "true" if is_checked else "false")
    
    def toggle_highlight_search(self, is_checked):
        self.highlight_search = is_checked
        self.save_setting("highlight_search", "true" if is_checked else "false")
        
    def toggle_alternating_rows(self, is_checked):
        self.alternating_rows = is_checked
        self.save_setting("alternating_rows", "true" if is_checked else "false")
        self.apply_visual_settings()
        
    def set_default_person(self, text):
        self.repeat_person = text
        self.save_setting("default_person", text)

    def set_default_occasion(self, text):
        self.repeat_occasion = text
        self.save_setting("default_occasion", text)

    def set_default_category(self, text):
        self.repeat_category = text
        self.save_setting("default_category", text)

    def set_default_place(self, text):
        self.repeat_place = text
        self.save_setting("default_place", text)
        
    def set_default_date(self, text):
        self.repeat_date = text
        self.save_setting("default_date", text)

    def set_default_outof(self, text):
        self.repeat_outof = text
        self.save_setting("default_outof", text)

    def set_default_a_v(self, text):
        self.repeat_a_v = text
        self.save_setting("default_av", text)

    def save_startup_view_pref(self, index):
        self.startup_view = index
        self.save_setting("startup_view", str(index))
        
    def prune_backups(self):
        """Enforces retention policy: Keep all < 4 weeks, older keep 1 per month."""
        try:
            if not os.path.exists(self.backup_dir): return
            
            # 1. Gather all backups
            allowed_exts = ['.db', '.zip', '.sql']
            files = [f for f in os.listdir(self.backup_dir) if any(f.endswith(e) for e in allowed_exts)]
            
            # Parse Dates: filename convention often contains YYYY-MM-DD or YYYYMMDD
            # archive_backup_YYYY-MM-DD_HH-MM-SS.db
            file_dates = []
            for f in files:
                path = os.path.join(self.backup_dir, f)
                # Try to extract date from filename first, fallback to mtime
                # Regex for YYYY-MM-DD
                import re
                match = re.search(r'(\d{4}-\d{2}-\d{2})', f)
                dt = None
                if match:
                    try:
                        dt = datetime.strptime(match.group(1), "%Y-%m-%d")
                    except: pass
                
                if not dt:
                    # Fallback to file modified time
                    dt = datetime.fromtimestamp(os.path.getmtime(path))
                
                file_dates.append({'name': f, 'path': path, 'date': dt})
            
            # Sort by date descending (newest first)
            file_dates.sort(key=lambda x: x['date'], reverse=True)
            
            cutoff = datetime.now() - timedelta(weeks=4)
            seen_months = set()
            
            for item in file_dates:
                if item['date'] > cutoff:
                    # Keep recent
                    continue
                else:
                    # Older than 4 weeks
                    m_key = (item['date'].year, item['date'].month)
                    if m_key in seen_months:
                        # We already have a backup for this month (a newer one, since we sorted desc)
                        # DELETE THIS ONE
                        try:
                            os.remove(item['path'])
                            print(f"Pruned old backup: {item['name']}")
                        except Exception as e:
                            print(f"Failed to prune {item['name']}: {e}")
                    else:
                        # Keep this one as the monthly snapshot
                        seen_months.add(m_key)
                        
        except Exception as e:
            self.log_message(f"Prune Error: {e}")

    def manual_backup(self, tag="manual"):
        """Triggers a manual backup immediately."""
        self.prune_backups() # Cleanup before making new one
        
        self.log_message("Starting backup...")
        # Check if backup dir exists
        if not os.path.exists(self.backup_dir):
            try:
                os.makedirs(self.backup_dir)
            except:
                self.log_message("Error: Link to Backup Folder is broken.")
                return

        # Pass prefix based on tag
        prefix = f"backup_{tag}"
        self.backup_worker_manual = BackupWorker(self.backup_dir, prefix=prefix)
        self.backup_worker_manual.finished.connect(lambda path: self.log_message(f"Backup saved: {os.path.basename(path)}"))
        self.backup_worker_manual.start()
        
    def open_backup_folder(self):
        if not os.path.exists(self.backup_dir):
            os.makedirs(self.backup_dir)
        subprocess.run(["open", self.backup_dir])

    def select_backup_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Backup Folder")
        if folder:
            self.txt_backup_path.setText(folder)
            self.backup_dir = folder
            self.save_setting("backup_path", folder)
            self.log_message(f"Backup location changed to: {folder}")

    def export_selection(self):
        # Default Columns requested by User
        default_cols = ["Audio No", "Person", "Occasion", "Category", "Place", "Country", "Hijri Date", "Esavi Date", "Year", "Remarks"]
        
        # Get selected rows
        rows = sorted(list(set(index.row() for index in self.table.selectedIndexes())))
        if not rows:
            QMessageBox.warning(self, "Export Selection", "Please select rows to export.")
            return
            
        # Dialog for Column Selection
        # Headers: self.headers[1:] (Skip Tracks)
        export_headers = self.headers[1:]
        dlg = ColumnSelectionDialog(export_headers, self, default_selection=default_cols)
        
        if dlg.exec() == QDialog.DialogCode.Accepted:
            selected_indices = dlg.get_selected_indices()
            include_tracks = dlg.is_tracks_included()
            is_distinct = (dlg.result_mode == "distinct")
            
            if not selected_indices:
                QMessageBox.warning(self, "Export Selection", "No columns selected.")
                return
                
            # Fixed Path to Desktop/Audio MP3/List.xlsx as requested
            # Note: User provided absolute path /Users/mahad/Desktop/Audio MP3/
            base_dir = "/Users/mahad/Desktop/Audio MP3/"
            if not os.path.exists(base_dir):
                try:
                    os.makedirs(base_dir)
                except Exception as e:
                    QMessageBox.warning(self, "Export Error", f"Could not create directory {base_dir}:\n{e}")
                    return

            path = os.path.join(base_dir, "List.xlsx")
            
            # If user wants to change location, we could ask, but request said "keep location... to desktop"
            # We will use this fixed path.
            
            try:
                # 1. Prepare Master Data
                data = []
                final_headers = [export_headers[i] for i in selected_indices]
                selected_audio_nos = []
                
                # We need all data first to filter for distinct if needed
                # So we gather data for ALL rows first
                
                for r in rows:
                    audio_item = self.table.item(r, 1)
                    if not audio_item: 
                        continue 
                    
                    # Capture Audio No for Tracks Export (if needed)
                    # For Distinct export, we might filter some out, so we'll handle tracks matching later if needed
                    # OR we just export tracks for selected rows regardless of distinct filtering?
                    # User didn't specify behavior for Tracks with Distinct.
                    # Assuming Distinct applies to the Master Sheet rows exported.
                    
                    row_data = []
                    # We need to grab data corresponding to final_headers
                    for i in selected_indices:
                        col_idx = i + 1 
                        it = self.table.item(r, col_idx)
                        row_data.append(it.text() if it else "")
                    data.append(row_data)

                if not data:
                    QMessageBox.warning(self, "Export Selection", "No valid data rows selected.")
                    return

                # 2. Apply Distinct Logic if requested
                final_data = []
                if is_distinct:
                    subset_cols = ["Person", "Category", "Place", "Country", "Hijri Date", "Esavi Date"]
                    # Find indices
                    chk_indices = [i for i, h in enumerate(final_headers) if h in subset_cols]
                    
                    if chk_indices:
                        seen = set()
                        for row_vals in data:
                            # Create key tuple from subset columns
                            key = tuple(row_vals[i] for i in chk_indices)
                            if key not in seen:
                                seen.add(key)
                                final_data.append(row_vals)
                    else:
                        final_data = data
                else:
                    final_data = data
                
                # Recalculate selected_audio_nos based on filtered data
                final_audio_nos = []
                if include_tracks:
                    if "Audio No" in final_headers:
                        idx_audio = final_headers.index("Audio No")
                        final_audio_nos = [str(row[idx_audio]) for row in final_data]
                    else:
                        pass

                # 3. Write to Excel with Formatting Preservation
                
                file_exists = os.path.exists(path)
                
                if file_exists:
                    try:
                        wb = openpyxl.load_workbook(path)
                        if "Master" in wb.sheetnames:
                            ws = wb["Master"]
                        else:
                            ws = wb.create_sheet("Master")
                            
                        # Clear data below header (row 2+)
                        if ws.max_row > 1:
                            ws.delete_rows(2, ws.max_row - 1)
                            
                        # Write new data
                        for values in final_data:
                            ws.append(values)
                            
                        wb.save(path)
                        self.log_message(f"Exported {len(final_data)} rows to existing file: {path}")
                        
                    except Exception as e:
                        # Fallback if load fails (e.g. file open)
                        raise Exception(f"Failed to update existing Excel: {e}")
                else:
                    # Create New
                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "Master"
                    
                    # Headers
                    ws.append(final_headers)
                    
                    # Data
                    for row in final_data:
                        ws.append(row)
                        
                    # Styling
                    # Font: Lisan or Arial, Size 12/14? User said "Same fonts same size".
                    # I'll try to match standard app look: Arial 12/14
                    
                    ft = Font(name="Arial", size=12)
                    for row in ws.iter_rows():
                        for cell in row:
                            cell.font = ft
                            
                    # Header Style
                    header_ft = Font(name="Arial", size=12, bold=True)
                    for cell in ws[1]:
                        cell.font = header_ft
                        
                    wb.save(path)
                    self.log_message(f"Created new export file: {path}")

                
                # Handle Tracks Sheet (Append or Replace?)
                # User asked to erase data below headers. 
                # If we have tracks, we should probably do similar logic for Tracks sheet.
                if include_tracks and final_audio_nos:
                    # Fetch Tracks matches
                    placeholders = ','.join(['?'] * len(final_audio_nos))
                    query = f"SELECT event_id, track_no, track_name FROM tracks WHERE event_id IN ({placeholders}) ORDER BY event_id, id"
                    
                    conn = sqlite3.connect(DB_FILE)
                    tracks_data = conn.execute(query, final_audio_nos).fetchall()
                    conn.close()
                    
                    # Open workbook again (or keep open? Better to separate steps to ensure save)
                    wb = openpyxl.load_workbook(path)
                    if "Tracks" in wb.sheetnames:
                        ws_t = wb["Tracks"]
                        if ws_t.max_row > 1:
                            ws_t.delete_rows(2, ws_t.max_row - 1)
                    else:
                        ws_t = wb.create_sheet("Tracks")
                        ws_t.append(["Audio No", "Track No", "Track Name"])
                        # Header Style
                        header_ft = Font(name="Arial", size=12, bold=True)
                        for cell in ws_t[1]:
                            cell.font = header_ft
                            
                    # Write Tracks
                    ft = Font(name="Arial", size=12)
                    for t_row in tracks_data:
                        ws_t.append(list(t_row))
                        # Style last row
                        for cell in ws_t[ws_t.max_row]:
                            cell.font = ft
                            
                    wb.save(path)

                QMessageBox.information(self, "Export", f"Export successful to:\n{path}")
            except Exception as e:
                # traceback.print_exc()
                QMessageBox.critical(self, "Export Error", str(e))

    def setup_main_area(self):
        self.stack = QStackedWidget()
        
        # --- PAGE 0: SHEETS Widget (Consolidated) ---
        self.sheet_page = QWidget()
        main_layout = QVBoxLayout(self.sheet_page)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(5)
        
        # --- TOOLBAR LINE 1: Search + Management + Lock ---
        line1 = QHBoxLayout()
        line1.setSpacing(6)
        
        # Search Box
        self.txt_search_master = QLineEdit()
        self.txt_search_master.setPlaceholderText("Search...")
        self.txt_search_master.setFixedWidth(220)
        self.txt_search_master.setStyleSheet("padding: 6px; border-radius: 4px; border: 1px solid #cbd5e0;")
        self.txt_search_master.returnPressed.connect(self.search_active_sheet)
        line1.addWidget(self.txt_search_master)
        
        self.btn_search_master = QPushButton("🔍") # Condensed text
        self.btn_search_master.setFixedWidth(40)
        self.btn_search_master.setObjectName("PrimaryBtn")
        self.btn_search_master.clicked.connect(self.search_active_sheet)
        line1.addWidget(self.btn_search_master)
        
        self.btn_clear_search = QPushButton("✖") # Condensed text
        self.btn_clear_search.setObjectName("SecondaryBtn")
        self.btn_clear_search.clicked.connect(self.smart_clear_search)
        line1.addWidget(self.btn_clear_search)
        
        # Add Button (Context Aware)
        self.btn_add = QPushButton("➕ Add Rows") 
        self.btn_add.setObjectName("PrimaryBtn")
        self.btn_add.clicked.connect(self.handle_add_rows)
        line1.addWidget(self.btn_add)
        
        line1.addSpacing(15)
        
        # Management Helpers
        self.btn_open_folder_main = QPushButton("📂 Folder")
        self.btn_open_folder_main.setObjectName("SecondaryBtn") 
        self.btn_open_folder_main.clicked.connect(self.open_entry_folder)
        line1.addWidget(self.btn_open_folder_main)
        
        self.btn_open_mp3_main = QPushButton("🎵 Mp3")
        self.btn_open_mp3_main.setObjectName("SecondaryBtn")
        self.btn_open_mp3_main.clicked.connect(self.open_entry_mp3)
        line1.addWidget(self.btn_open_mp3_main)
        
        self.btn_export_sel = QPushButton("📊 Export")
        self.btn_export_sel.setObjectName("SecondaryBtn")
        self.btn_export_sel.clicked.connect(self.export_active_sheet_selection)
        line1.addWidget(self.btn_export_sel)
        
        self.btn_refresh_sheet = QPushButton("🔄 Refresh")
        self.btn_refresh_sheet.setObjectName("SecondaryBtn") 
        self.btn_refresh_sheet.clicked.connect(self.refresh_active_sheet)
        line1.addWidget(self.btn_refresh_sheet)

        line1.addStretch()
        
        # GLOBAL LOCK BUTTON
        self.btn_master_lock = QPushButton("🔒 MASTER LOCKED")
        self.btn_master_lock.setCheckable(True)
        self.btn_master_lock.setChecked(True)
        self.btn_master_lock.setFixedSize(150, 32)
        self.btn_master_lock.setStyleSheet("background-color: #742a2a; color: white; font-weight: bold; border: 2px solid #e53e3e; border-radius: 6px;")
        self.btn_master_lock.clicked.connect(self.toggle_global_lock)
        line1.addWidget(self.btn_master_lock)
        
        main_layout.addLayout(line1)
        
        if not hasattr(self, 'filter_state'):
            self.filter_state = {} 
        
        self.filter_history = [] 

        # Add Advanced Filter Button to Line 1
        self.btn_advanced_filter = QPushButton("Filter")
        self.btn_advanced_filter.setObjectName("PrimaryBtn")
        self.btn_advanced_filter.setStyleSheet("background-color: #2b6cb0; color: white; padding: 6px;")
        self.btn_advanced_filter.clicked.connect(self.open_advanced_filter)
        
        # Insert FILTER button into Line 1 (Index 3, after Clear Search)
        line1.insertWidget(3, self.btn_advanced_filter) 

        # Add Smart Copy Button
        self.btn_smart_copy = QPushButton("Smart Copy")
        self.btn_smart_copy.setStyleSheet("background-color: #6b46c1; color: white; padding: 6px; font-weight: bold;")
        self.btn_smart_copy.clicked.connect(self.smart_copy_selection)
        line1.insertWidget(4, self.btn_smart_copy)
        
        # Update Row Count Label - Move to Line 1
        # self.lbl_row_count = QLabel("Rows: 0") # REMOVED from original position
        # self.lbl_row_count.setStyleSheet("font-weight: bold; color: #DC143C; font-size: 14px; margin-left: 10px;")
        # line1.insertWidget(8, self.lbl_row_count) 

        # Add Row Count AFTER Master Lock
        self.lbl_row_count = QLabel("Row Count : 0/0")
        self.lbl_row_count.setStyleSheet("font-weight: bold; color: #E53E3E; font-size: 14px; margin-left: 10px;")
        line1.addWidget(self.lbl_row_count) 
        
        # --- TABS FOR SHEETS ---
        self.sheets_tabs = QTabWidget()
        self.sheets_tabs.setStyleSheet("""
            QTabWidget::pane { border: 1px solid #cbd5e0; background: white; }
            QTabBar::tab { background: #edf2f7; color: #4a5568; padding: 8px 15px; margin-right: 2px; border-top-left-radius: 4px; border-top-right-radius: 4px; }
            QTabBar::tab:selected { background: white; color: #2d3748; font-weight: bold; border-bottom: 2px solid #3182ce; }
        """)
        
        # --- TAB 0: MASTER SHEET ---
        self.table = ArchiveTable(self)
        self.table.setObjectName("MasterTable")
        # Initial Lock State
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers) # Set to NoEditTriggers because master_locked is True by default
        
        # New Header Structure: "+" at index 0. Reorder "Tracks" out or merged? 
        # User said "Plus button in start of each row... open tracks window". 
        # So we likely replace the old "Tracks" column (which was index 0) with this "+" interactive column.
        # Headers: (+), Audio No, Person, Occasion, Category, Place, Country, Hijri, Esavi, Year, Out, Remarks, Summary, AV, Cass, Came, Inc
        
        self.headers = ["+", "Audio No", "Person", "Occasion", "Category", "Place", "Country", "Hijri Date", "Esavi Date", "Year", "Out of", "Remarks", "Tracks", "AV", "Cass No", "Came From", "Incomplete"]
        self.table.setColumnCount(len(self.headers))
        self.table.setHorizontalHeaderLabels(self.headers)
        
        self.table.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectItems)
        self.table.verticalHeader().setVisible(False) # We will manage rows
        
        # Triggers
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.SelectedClicked | QAbstractItemView.EditTrigger.AnyKeyPressed)
        
        # Layout behaviors - ENABLE word wrap for proper height calculation with wrapped text
        self.table.setWordWrap(True)
        # Interactive Resize for columns
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        self.table.horizontalHeader().setStretchLastSection(True)
        
        # Debounced Row Resizing on Column Adjustment
        self.table.horizontalHeader().sectionResized.connect(lambda: self.master_resize_timer.start())
        self.table.setColumnWidth(0, 40) # Small width for "+"
        
        # Resize Rows Interactive? User said "rows should always resize according the text".
        # So manual resize might be overwritten. We will set ResizeToContents logic later.
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive) 
        
        # Connect Signals
        self.update_delegates()
        self.table.cellClicked.connect(self.on_table_cell_clicked)
        self.table.itemChanged.connect(self.handle_item_changed)
        
        self.sheets_tabs.addTab(self.table, "Master Sheet")
        self.sheets_tabs.currentChanged.connect(self.on_sheets_tab_changed)
        
        # Load Other Sheets
        self.load_other_sheets()
        
        main_layout.addWidget(self.sheets_tabs)


        
        # --- PAGE 1: TRACKS (Renamed from Search Engine) ---
        self.tracks_page = QWidget()
        tl_layout = QVBoxLayout(self.tracks_page)
        tl_layout.setContentsMargins(0, 0, 0, 0)
        
        self.tracks_tabs = QTabWidget()
        
        # Tab 1: Tracks Sheet (ID=1)
        # We use OtherSheetTable logic but instantiated here manually
        self.tracks_sheet_table = OtherSheetTable(self)
        self.tracks_sheet_table.setProperty("table_name", "tracks") # For DB updates
        self.tracks_sheet_table.setAlternatingRowColors(self.alternating_rows)
        
        # Initial Locked State
        if getattr(self, 'master_locked', True):
            self.tracks_sheet_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        else:
            self.tracks_sheet_table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed)
            
        # Delegate
        track_delegate = EditorDelegate(self.lisan_font_family, "text", self.tracks_sheet_table)
        self.tracks_sheet_table.setItemDelegate(track_delegate)
        self.tracks_sheet_table.itemChanged.connect(lambda item: self.save_other_sheet_cell(item))
        # Initial Load
        self.load_tracks_sheet_data(self.tracks_sheet_table) # New helper method needed or reuse existing logic
        
        self.tracks_tabs.addTab(self.tracks_sheet_table, "Tracks Sheet")

        # Tab 2: Track Search (Existing)
        self.track_search_tab = TrackSearchTab(self)
        self.tracks_tabs.addTab(self.track_search_tab, "Track Search")
        
        tl_layout.addWidget(self.tracks_tabs)
        

        
        # --- PAGE 3: SETTINGS (REDESIGNED) ---
        self.settings_page = QWidget()
        set_layout = QVBoxLayout(self.settings_page)
        set_layout.setContentsMargins(30, 30, 30, 30)
        self.settings_tabs = QTabWidget()



        
        # Use a Scroll Area for the settings tab content
        self.tools_tab = QWidget()
        tools_layout = QVBoxLayout(self.tools_tab)
        
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        
        scroll_content = QWidget()
        tl = QVBoxLayout(scroll_content)
        tl.setSpacing(20)
        tl.setContentsMargins(10, 10, 20, 10)
        
        # -- Card 1: Appearance --
        self.group_appearance = QGroupBox("Appearance")
        app_layout = QGridLayout()
        app_layout.setContentsMargins(15, 20, 15, 20)
        app_layout.setSpacing(15)
        
        # Dark Mode
        lbl_dark = QLabel("Dark Mode")
        lbl_dark_sub = QLabel("Switch between light and dark themes.")
        lbl_dark_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        self.chk_theme = ToggleSwitch() # Changed from QCheckBox to custom ToggleSwitch
        self.chk_theme.setChecked(self.is_dark_mode) # Set initial state
        self.chk_theme.toggled.connect(self.toggle_theme_mode)
        
        app_layout.addWidget(lbl_dark, 0, 0)
        app_layout.addWidget(lbl_dark_sub, 1, 0)
        app_layout.addWidget(self.chk_theme, 0, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Alternating Rows
        lbl_alt = QLabel("Alternating Row Colors")
        lbl_alt_sub = QLabel("Shade every other row for readability.")
        lbl_alt_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        self.chk_alt = ToggleSwitch()
        self.chk_alt.setChecked(self.alternating_rows)
        self.chk_alt.toggled.connect(self.toggle_alternating_rows)
        
        app_layout.addWidget(lbl_alt, 2, 0)
        app_layout.addWidget(lbl_alt_sub, 3, 0)
        app_layout.addWidget(self.chk_alt, 2, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Highlight Search
        lbl_hl = QLabel("Highlight Search Matches")
        lbl_hl_sub = QLabel("Highlight text matching search queries.")
        lbl_hl_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        self.chk_highlight = ToggleSwitch()
        self.chk_highlight.setChecked(self.highlight_search)
        self.chk_highlight.toggled.connect(self.toggle_highlight_search)
        
        app_layout.addWidget(lbl_hl, 4, 0)
        app_layout.addWidget(lbl_hl_sub, 5, 0)
        app_layout.addWidget(self.chk_highlight, 4, 1, 2, 1, Qt.AlignmentFlag.AlignRight)

        # Text Size
        lbl_font = QLabel("Interface Text Size")
        self.slider_font = QSlider(Qt.Orientation.Horizontal)
        self.slider_font.setRange(10, 24)
        self.slider_font.setValue(self.current_font_size) # Set initial state
        self.slider_font.setFixedWidth(200)
        self.slider_font.valueChanged.connect(self.update_app_font_size)
        
        self.lbl_font_size_display = QLabel(f"{self.current_font_size}px")
        self.lbl_font_size_display.setStyleSheet("font-weight: bold; color: #3182ce;")
        
        app_layout.addWidget(lbl_font, 6, 0)
        app_layout.addWidget(self.slider_font, 6, 1)
        app_layout.addWidget(self.lbl_font_size_display, 6, 2)
        
        self.group_appearance.setLayout(app_layout)
        tl.addWidget(self.group_appearance)
        
        # -- Card 2: Workflow & Safety (NEW) --
        self.group_workflow = QGroupBox("Workflow & Safety")
        wf_layout = QGridLayout()
        wf_layout.setContentsMargins(15, 20, 15, 20)
        wf_layout.setSpacing(15)
        
        # Auto Backup
        lbl_backup = QLabel("Auto-Backup on Startup")
        lbl_backup_sub = QLabel("Automatically create daily/weekly backups.")
        lbl_backup_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        self.chk_backup = ToggleSwitch()
        self.chk_backup.setChecked(self.auto_backup)
        self.chk_backup.toggled.connect(self.toggle_auto_backup)
        
        wf_layout.addWidget(lbl_backup, 0, 0)
        wf_layout.addWidget(lbl_backup_sub, 1, 0)
        wf_layout.addWidget(self.chk_backup, 0, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Backup Location
        lbl_loc = QLabel("Backup Location")
        self.txt_backup_path = QLineEdit(self.backup_dir)
        self.txt_backup_path.setReadOnly(True)
        self.txt_backup_path.setStyleSheet("color: #718096;")
        btn_change_loc = QPushButton("Change...")
        btn_change_loc.clicked.connect(self.select_backup_folder)
        btn_change_loc.setObjectName("SecondaryBtn")
        
        wf_layout.addWidget(lbl_loc, 2, 0)
        wf_layout.addWidget(self.txt_backup_path, 3, 0)
        wf_layout.addWidget(btn_change_loc, 3, 1)

        # Manual Backup Controls
        btn_backup_now = QPushButton("Backup Now")
        btn_backup_now.clicked.connect(self.manual_backup)
        btn_backup_now.setObjectName("SecondaryBtn")
        
        btn_open_backups = QPushButton("Open Folder")
        btn_open_backups.clicked.connect(self.open_backup_folder)
        btn_open_backups.setObjectName("SecondaryBtn")
        
        wf_layout.addWidget(btn_backup_now, 4, 0)
        wf_layout.addWidget(btn_open_backups, 4, 1)
        
        # Confirm Delete
        lbl_del = QLabel("Confirm Deletion")
        lbl_del_sub = QLabel("Ask for confirmation before deleting items.")
        lbl_del_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        self.chk_confirm = ToggleSwitch()
        self.chk_confirm.setChecked(self.confirm_delete)
        self.chk_confirm.toggled.connect(self.toggle_confirm_delete)
        
        wf_layout.addWidget(lbl_del, 5, 0)
        wf_layout.addWidget(lbl_del_sub, 6, 0)
        wf_layout.addWidget(self.chk_confirm, 5, 1, 2, 1, Qt.AlignmentFlag.AlignRight)

        # Startup View
        lbl_start = QLabel("Startup View")
        lbl_start_sub = QLabel("Choose which page opens on launch.")
        lbl_start_sub.setStyleSheet("color: #718096; font-size: 11px;")
        
        bg_start = QButtonGroup(self)
        self.rad_start_master = QRadioButton("Sheets")
        self.rad_start_search = QRadioButton("Search")
        
        # Check current Setting
        if self.startup_view == 1:
            self.rad_start_search.setChecked(True)
        else:
            self.rad_start_master.setChecked(True)
            
        bg_start.addButton(self.rad_start_master)
        bg_start.addButton(self.rad_start_search)
        
        self.rad_start_master.toggled.connect(lambda: self.save_setting("startup_view", "0"))
        self.rad_start_search.toggled.connect(lambda: self.save_setting("startup_view", "1"))
        
        wf_layout.addWidget(lbl_start, 7, 0)
        wf_layout.addWidget(lbl_start_sub, 8, 0)
        start_opts = QVBoxLayout()
        start_opts.addWidget(self.rad_start_master)
        start_opts.addWidget(self.rad_start_search)
        wf_layout.addLayout(start_opts, 7, 1, 2, 1)
        
        self.group_workflow.setLayout(wf_layout)
        tl.addWidget(self.group_workflow)
        
        # -- Card: New Entry Defaults --
        self.group_defaults = QGroupBox("New Entry Defaults")
        def_layout = QGridLayout()
        def_layout.setContentsMargins(15, 20, 15, 20)
        def_layout.setSpacing(15)

        # Helper to add Dropdown rows
        def add_dropdown_row(layout, row, label, current_val, items, slot):
            l = QLabel(label)
            cb = QComboBox()
            # Set font to Fatemi for dropdowns
            cb.setFont(QFont(self.lisan_font_family, self.current_font_size))
            cb.setFixedWidth(200)
            cb.addItems(items)
            cb.setCurrentText(current_val)
            cb.currentTextChanged.connect(slot)
            layout.addWidget(l, row, 0)
            layout.addWidget(cb, row, 1, Qt.AlignmentFlag.AlignRight)

        # Build items lists
        base_opts = ["Empty", "Repeat Previous"]
        list_person = base_opts + self.master_lists.get(2, [])
        list_occasion = base_opts + self.master_lists.get(3, [])
        list_category = base_opts + self.master_lists.get(4, [])
        list_place = base_opts + self.master_lists.get(5, [])
        list_date = base_opts # Dates have no master list
        list_outof = base_opts + [str(i) for i in range(1, 11)]
        list_av = base_opts + self.master_lists.get(13, []) # New A/V List

        add_dropdown_row(def_layout, 0, "Person", self.repeat_person, list_person, self.set_default_person)
        add_dropdown_row(def_layout, 1, "Occasion", self.repeat_occasion, list_occasion, self.set_default_occasion)
        add_dropdown_row(def_layout, 2, "Category", self.repeat_category, list_category, self.set_default_category)
        add_dropdown_row(def_layout, 3, "Place", self.repeat_place, list_place, self.set_default_place)
        add_dropdown_row(def_layout, 4, "Date", self.repeat_date, list_date, self.set_default_date)
        add_dropdown_row(def_layout, 5, "Out of", self.repeat_outof, list_outof, self.set_default_outof)
        add_dropdown_row(def_layout, 6, "AV", self.repeat_a_v, list_av, self.set_default_a_v) # New Row

        self.group_defaults.setLayout(def_layout)
        tl.addWidget(self.group_defaults)
        
        # -- Card 3: Data Management --
        self.group_data = QGroupBox("Data Management")
        data_layout = QGridLayout()
        data_layout.setContentsMargins(15, 20, 15, 20)
        data_layout.setSpacing(15)
        
        # Refresh Master Lists
        lbl_ref = QLabel("Refresh Master Lists")
        lbl_ref_sub = QLabel("Reload master lists from database to fix dropdowns.")
        lbl_ref_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_ref = QPushButton("Refresh Lists")
        btn_ref.clicked.connect(self.refresh_master_lists_action)
        btn_ref.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_ref, 0, 0)
        data_layout.addWidget(lbl_ref_sub, 1, 0)
        data_layout.addWidget(btn_ref, 0, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Rebuild DB
        lbl_reb = QLabel("Rebuild Database")
        lbl_reb_sub = QLabel("Run migration scripts to ensure DB schema is up to date.")
        lbl_reb_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_reb = QPushButton("Rebuild DB")
        btn_reb.clicked.connect(self.init_db_and_migrate)
        btn_reb.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_reb, 2, 0)
        data_layout.addWidget(lbl_reb_sub, 3, 0)
        data_layout.addWidget(btn_reb, 2, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Export .xlsx
        lbl_exp = QLabel("Export Data")
        lbl_exp_sub = QLabel("Save the entire Master Sheet to Excel.")
        lbl_exp_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_xls = QPushButton("Export .xlsx")
        btn_xls.clicked.connect(self.export_to_excel)
        btn_xls.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_exp, 4, 0)
        data_layout.addWidget(lbl_exp_sub, 5, 0)
        data_layout.addWidget(btn_xls, 4, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Export DB
        lbl_exdb = QLabel("Export Database")
        lbl_exdb_sub = QLabel("Save the entire .db file for backup/transfer.")
        lbl_exdb_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_exdb = QPushButton("Export DB")
        btn_exdb.clicked.connect(self.export_database)
        btn_exdb.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_exdb, 6, 0)
        data_layout.addWidget(lbl_exdb_sub, 7, 0)
        data_layout.addWidget(btn_exdb, 6, 1, 2, 1, Qt.AlignmentFlag.AlignRight)

        # Import DB (Moved Here)
        lbl_imdb = QLabel("Import Database")
        lbl_imdb_sub = QLabel("Overwrite current DB with an external .db file.")
        lbl_imdb_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_imdb = QPushButton("Import DB")
        btn_imdb.clicked.connect(self.import_database)
        btn_imdb.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_imdb, 8, 0)
        data_layout.addWidget(lbl_imdb_sub, 9, 0)
        data_layout.addWidget(btn_imdb, 8, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        # Optimize
        lbl_vac = QLabel("Database Maintenance")
        lbl_vac_sub = QLabel("Compact database file to save space.")
        lbl_vac_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_vac = QPushButton("Run Vacuum")
        btn_vac.clicked.connect(self.optimize_database)
        btn_vac.setObjectName("SecondaryBtn")
        
        data_layout.addWidget(lbl_vac, 10, 0)
        data_layout.addWidget(lbl_vac_sub, 11, 0)
        data_layout.addWidget(btn_vac, 10, 1, 2, 1, Qt.AlignmentFlag.AlignRight)

        # Full Import
        lbl_fimp = QLabel("Import Full Data")
        lbl_fimp_sub = QLabel("Wipe database and import comprehensive Excel file.")
        lbl_fimp_sub.setStyleSheet("color: #718096; font-size: 11px;")
        btn_fimp = QPushButton("Import Full Data")
        btn_fimp.clicked.connect(self.import_full_excel_backup)
        btn_fimp.setObjectName("SecondaryBtn")

        data_layout.addWidget(lbl_fimp, 12, 0)
        data_layout.addWidget(lbl_fimp_sub, 13, 0)
        data_layout.addWidget(btn_fimp, 12, 1, 2, 1, Qt.AlignmentFlag.AlignRight)
        
        self.group_data.setLayout(data_layout)
        tl.addWidget(self.group_data)
        
        tl.addStretch()
        
        # Update Button
        btn_update = QPushButton("Check for Updates")
        btn_update.setObjectName("SecondaryBtn")
        btn_update.clicked.connect(lambda: self.check_for_updates(silent=False))
        tl.addWidget(btn_update)
        
        scroll_area.setWidget(scroll_content)
        tools_layout.addWidget(scroll_area)
        self.settings_tabs.addTab(self.tools_tab, "General Settings")

        # --- MASTER LISTS TAB (Moved from Page 2) ---
        self.master_lists_page = QWidget()
        ml_layout = QVBoxLayout(self.master_lists_page)
        ml_layout.setContentsMargins(20, 20, 20, 20)
        
        ml_top_controls = QHBoxLayout()
        self.btn_master_lists_lock = QPushButton("🔒 Locked")
        self.btn_master_lists_lock.setObjectName("LockBtn")
        self.btn_master_lists_lock.setCheckable(True)
        self.btn_master_lists_lock.setChecked(True)
        self.btn_master_lists_lock.clicked.connect(self.toggle_master_lists_global_lock)
        ml_top_controls.addWidget(self.btn_master_lists_lock)
        
        ml_top_controls.addSpacing(20)
        self.btn_refresh_master = QPushButton("🔄 Refresh Lists")
        self.btn_refresh_master.setObjectName("SecondaryBtn")
        self.btn_refresh_master.clicked.connect(self.refresh_master_lists_action)
        ml_top_controls.addWidget(self.btn_refresh_master)
        ml_top_controls.addStretch()
        ml_layout.addLayout(ml_top_controls)

        self.master_tab_widget = QTabWidget()
        self.master_tab_tables = {}
        self.master_tab_buttons = {} 
        categories = ["Person", "Occasion", "Category", "Locations", "AV", "Incomplete", "Created"]
        
        for cat in categories:
            tab = QWidget()
            tl = QVBoxLayout(tab)
            tl.setContentsMargins(10, 10, 10, 10)
            btn_h = QHBoxLayout()
            b_add = QPushButton(f"➕ Add {cat}")
            b_add.setEnabled(False) 
            b_add.clicked.connect(lambda checked, c=cat: self.add_master_tab_item(c))
            
            b_del = QPushButton(f"🗑️ Delete Selected")
            b_del.setEnabled(False) 
            b_del.clicked.connect(lambda checked, c=cat: self.delete_master_tab_item(c))
            
            b_import = QPushButton("📥 Import Excel")
            b_import.setObjectName("SecondaryBtn")
            b_import.clicked.connect(lambda checked, c=cat: self.import_master_list(c))
            b_import.setEnabled(False)
            
            b_export = QPushButton("📊 Export")
            b_export.setObjectName("SecondaryBtn")
            b_export.clicked.connect(lambda checked, c=cat: self.export_simple_table_selection(self.master_tab_tables[c], f"MasterList_{c}"))
            
            self.master_tab_buttons[cat] = {'add': b_add, 'del': b_del, 'import': b_import}
            
            btn_h.addSpacing(20)
            btn_h.addWidget(b_import)
            btn_h.addWidget(b_export)
            btn_h.addWidget(b_add)
            btn_h.addWidget(b_del)
            btn_h.addStretch()
            tl.addLayout(btn_h)
            
            t_tbl = MasterListTable(self)
            if cat == "Locations":
                t_tbl.setColumnCount(2)
                t_tbl.setHorizontalHeaderLabels(["Place", "Country"])
            else:
                t_tbl.setColumnCount(1)
                t_tbl.setHorizontalHeaderLabels([cat])
            
            t_tbl.setItemDelegate(EditorDelegate(self.lisan_font_family, "text", t_tbl))
            t_tbl.itemChanged.connect(lambda item, c=cat: self.save_master_tab_edit(item, c))
            
            tl.addWidget(t_tbl)
            self.master_tab_tables[cat] = t_tbl
            self.master_tab_widget.addTab(tab, cat)

        self.fb_add = QPushButton("➕ Add Folder Mapping")
        self.fb_add.setEnabled(False)
        self.fb_add.clicked.connect(self.add_settings_item)
        
        self.fb_del = QPushButton("🗑️ Delete Selected")
        self.fb_del.setEnabled(False)
        self.fb_del.clicked.connect(self.delete_settings_item)
        
        self.fb_imp = QPushButton("📂 Import Excel")
        self.fb_imp.setEnabled(False)
        self.fb_imp.clicked.connect(self.import_folder_info)
        self.fb_imp.setObjectName("SecondaryBtn")
        
        self.fb_export = QPushButton("📊 Export")
        self.fb_export.setObjectName("SecondaryBtn")
        self.fb_export.clicked.connect(lambda: self.export_simple_table_selection(self.val_table, "Folder_Mappings"))

        self.btn_mappings_lock = QPushButton("🔒 Locked")
        self.btn_mappings_lock.setObjectName("LockBtn")
        self.btn_mappings_lock.setCheckable(True)
        self.btn_mappings_lock.setChecked(True)
        self.btn_mappings_lock.clicked.connect(self.toggle_mappings_global_lock)

        ml_layout.addWidget(self.master_tab_widget)
        # MOVED to Stack Page (requested by User)
        # self.master_lists_page is already defined above at line 5139
        # self.settings_tabs.addTab(self.master_list_tab, "Master Lists")

        # TAB 2: Folder Mappings
        self.folder_tab = QWidget()
        fl = QVBoxLayout(self.folder_tab)
        fb_h = QHBoxLayout()
        fb_h.addWidget(self.btn_mappings_lock)
        fb_h.addSpacing(20)
        fb_h.addWidget(self.fb_imp)
        fb_h.addWidget(self.fb_export)
        fb_h.addWidget(self.fb_add)
        fb_h.addWidget(self.fb_del)
        fb_h.addStretch()
        fl.addLayout(fb_h)
        
        self.val_table = MasterListTable(self)
        self.val_table.setColumnCount(4)
        self.val_table.setHorizontalHeaderLabels(["Start No", "Drive", "Name", "Link"])
        self.val_table.cellDoubleClicked.connect(self.handle_settings_double_click)
        self.val_table.itemChanged.connect(self.save_settings_edit)
        self.val_table.setItemDelegate(EditorDelegate(self.lisan_font_family, "text", self.val_table))
        fl.addWidget(self.val_table)
        self.settings_tabs.addTab(self.folder_tab, "Folder Mappings")

        # TAB 3: Mp3 Folder Mappings
        self.mp3_tab = QWidget()
        ml = QVBoxLayout(self.mp3_tab)
        mb_h = QHBoxLayout()
        
        self.fb_add_mp3 = QPushButton("➕ Add Mp3 Mapping")
        self.fb_add_mp3.setEnabled(False)
        self.fb_add_mp3.clicked.connect(self.add_mp3_settings_item)
        
        self.fb_del_mp3 = QPushButton("🗑️ Delete Selected")
        self.fb_del_mp3.setEnabled(False)
        self.fb_del_mp3.clicked.connect(self.delete_mp3_settings_item)
        
        self.fb_imp_mp3 = QPushButton("📂 Import Excel")
        self.fb_imp_mp3.setEnabled(False)
        self.fb_imp_mp3.clicked.connect(self.import_mp3_folder_info)
        self.fb_imp_mp3.setObjectName("SecondaryBtn")
        
        self.fb_export_mp3 = QPushButton("📊 Export")
        self.fb_export_mp3.setObjectName("SecondaryBtn")
        self.fb_export_mp3.clicked.connect(lambda: self.export_simple_table_selection(self.mp3_table, "Mp3_Mappings"))
        
        mb_h = QHBoxLayout()
        mb_h.addSpacing(20) # Spacer if needed
        mb_h.addWidget(self.fb_imp_mp3)
        mb_h.addWidget(self.fb_export_mp3)
        mb_h.addWidget(self.fb_add_mp3)
        mb_h.addWidget(self.fb_del_mp3)
        mb_h.addStretch()
        ml.addLayout(mb_h)
        
        self.mp3_table = MasterListTable(self)
        self.mp3_table.setColumnCount(4)
        self.mp3_table.setHorizontalHeaderLabels(["Start No", "Drive", "Name", "Link"])
        self.mp3_table.cellDoubleClicked.connect(self.handle_mp3_settings_double_click)
        self.mp3_table.itemChanged.connect(self.save_mp3_settings_edit)
        self.mp3_table.setItemDelegate(EditorDelegate(self.lisan_font_family, "text", self.mp3_table))
        ml.addWidget(self.mp3_table)
        self.settings_tabs.addTab(self.mp3_tab, "Mp3 Mappings")

        
        set_layout.addWidget(self.settings_tabs)
        
        self.stack.addWidget(self.sheet_page)
        self.stack.addWidget(self.tracks_page)
        
        # Old Sheets Page Removed

        self.stack.addWidget(self.master_lists_page) # Index 3
        
        self.stack.addWidget(self.settings_page) # Index 4 now

        self.main_layout.addWidget(self.stack)



    def refresh_master_sheet_action(self):
        """Reloads the Master Sheet with confirmation."""
        if not self.check_unsaved_changes():
            return
        # Force initial load state (Last 50)
        self.load_data()
        self.log_message("Master Sheet refreshed from database.")
        QMessageBox.information(self, "Refresh", "Master Sheet data has been successfully reloaded.")

    def refresh_master_lists_action(self):
        """Reloads the Master Lists with confirmation."""
        if not self.check_unsaved_changes():
            return
        self.load_master_tabs_data()
        self.log_message("Master Lists refreshed from database.")
        QMessageBox.information(self, "Refresh", "Master Lists have been successfully refreshed.")

    def toggle_master_lists_global_lock(self, checked):
        """Toggles lock for all Master List tables."""
        if not checked: # User trying to UNLOCK
             # Check Master Lock
             if hasattr(self, 'master_locked') and self.master_locked:
                 QMessageBox.warning(self, "Master Locked", "Master Database is Locked.")
                 self.btn_master_lists_lock.setChecked(True)
                 return
             
             reply = QMessageBox.question(self, "Confirm Unlock", "Unlock all Master Lists?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
             if reply != QMessageBox.StandardButton.Yes:
                 self.btn_master_lists_lock.setChecked(True)
                 return
             self.btn_master_lists_lock.setText("🔓 Unlocked")
        else:
             self.btn_master_lists_lock.setText("🔒 Locked")

        for cat, tbl in self.master_tab_tables.items():
            tbl.set_lock_state(checked)
            btns = self.master_tab_buttons.get(cat, {})
            if btns.get('add'): btns['add'].setEnabled(not checked)
            if btns.get('del'): btns['del'].setEnabled(not checked)
            if btns.get('import'): btns['import'].setEnabled(not checked)

    def toggle_mappings_global_lock(self, checked):
        """Toggles lock for Folder and Mp3 mappings."""
        if not checked: # User trying to UNLOCK
             if hasattr(self, 'master_locked') and self.master_locked:
                 QMessageBox.warning(self, "Master Locked", "Master Database is Locked.")
                 self.btn_mappings_lock.setChecked(True)
                 return
             
             reply = QMessageBox.question(self, "Confirm Unlock", "Unlock all Mappings?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
             if reply != QMessageBox.StandardButton.Yes:
                 self.btn_mappings_lock.setChecked(True)
                 return
             self.btn_mappings_lock.setText("🔓 Unlocked")
        else:
             self.btn_mappings_lock.setText("🔒 Locked")

        # Tables
        if hasattr(self, 'val_table'): self.val_table.set_lock_state(checked)
        if hasattr(self, 'mp3_table'): self.mp3_table.set_lock_state(checked)
        
        # Buttons
        for b in [self.fb_add, self.fb_del, self.fb_imp, self.fb_add_mp3, self.fb_del_mp3, self.fb_imp_mp3]:
            if b: b.setEnabled(not checked)

    def toggle_tab_lock(self, key, is_checked, btn_lock, btn_add, btn_del, target_dict, btn_import=None):
        # Confirmation for Unlocking
        if not is_checked: # User trying to UNLOCK
             reply = QMessageBox.question(self, "Confirm Unlock", f"Are you sure you want to unlock {key}?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
             if reply == QMessageBox.StandardButton.No:
                 # Revert
                 btn_lock.blockSignals(True)
                 btn_lock.setChecked(True)
                 btn_lock.blockSignals(False)
                 return

        tbl = target_dict[key]
        if is_checked:
            btn_lock.setText("🔒 Locked")
            btn_lock.setStyleSheet("") 
            btn_add.setEnabled(False)
            btn_del.setEnabled(False)
            if btn_import:
                btn_import.setEnabled(False)
            tbl.set_lock_state(True)
        else:
            btn_lock.setText("🔓 Unlocked")
            btn_lock.setStyleSheet("background-color: #fef08a; color: #854d0e; border: 1px solid #fde047;")
            btn_add.setEnabled(True)
            btn_del.setEnabled(True)
            if btn_import:
                btn_import.setEnabled(True)
            tbl.set_lock_state(False)

    def optimize_database(self):
        conn = sqlite3.connect(DB_FILE)
        conn.execute("VACUUM")
        conn.close()
        self.log_message("Database optimized (VACUUM completed).")
        QMessageBox.information(self, "System", "Database compacted successfully.")

    def export_to_excel(self):
        path = QFileDialog.getSaveFileName(self, "Export Archive", "", "Excel Files (*.xlsx)")[0]
        if path:
            conn = sqlite3.connect(DB_FILE)
            df = pd.read_sql_query("SELECT * FROM events ORDER BY CAST(audio_no AS INTEGER) ASC", conn)
            df.to_excel(path, index=False)
            conn.close()
            self.log_message(f"Exported database to: {path}")
            QMessageBox.information(self, "System", "Export complete.")

    def export_database(self):
        """Copies the actual .db file to a user-selected location."""
        path = QFileDialog.getSaveFileName(self, "Export Database File", "archive_database.db", "SQLite Database (*.db)")[0]
        if not path:
            return
            
        try:
            shutil.copy2(DB_FILE, path)
            self.log_message(f"Database exported to: {path}")
            QMessageBox.information(self, "Success", "Database file exported successfully.")
        except Exception as e:
            QMessageBox.critical(self, "Export Failed", f"Could not export database:\n{e}")

    def export_simple_table_selection(self, table, default_name="Export"):
        """Export selected rows from a simple table widget to Excel."""
        if not table:
            QMessageBox.warning(self, "Export", "No table available for export.")
            return
        
        # Get selected rows
        selected_rows = sorted(list(set(index.row() for index in table.selectedIndexes())))
        
        # If no selection, export all rows
        if not selected_rows:
            reply = QMessageBox.question(
                self, "Export All", 
                "No rows selected. Export all rows?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                selected_rows = list(range(table.rowCount()))
            else:
                return
        
        # Get column count and headers
        col_count = table.columnCount()
        headers = [table.horizontalHeaderItem(i).text() if table.horizontalHeaderItem(i) else f"Column {i}" 
                  for i in range(col_count)]
        
        # Extract data
        data = []
        for row in selected_rows:
            row_data = []
            for col in range(col_count):
                item = table.item(row, col)
                row_data.append(item.text() if item else "")
            data.append(row_data)
        
        # Create DataFrame
        df = pd.DataFrame(data, columns=headers)
        
        # Save to file
        path = QFileDialog.getSaveFileName(self, "Export to Excel", f"{default_name}.xlsx", "Excel Files (*.xlsx)")[0]
        if path:
            try:
                df.to_excel(path, index=False)
                QMessageBox.information(self, "Success", f"Exported {len(data)} rows to:\n{path}")
            except Exception as e:
                QMessageBox.critical(self, "Export Failed", f"Could not export:\n{e}")


    def import_full_excel_backup(self):
        """Totally replaces the database with data from a structured Excel file."""
        msg = "You are about to completely WIPE the current database and replace it with data from an Excel file.\n\n" \
              "Requirements:\n" \
              "- 'Master' sheet for main events.\n" \
              "- 'Person', 'Occasion', 'Category', 'Location' for master lists.\n" \
              "- 'INFO SHEET', 'INFO SHEET MP3' for folder mappings.\n" \
              "- 'Tracks' for track details.\n" \
              "- All other sheets will be imported as Other Sheets.\n\n" \
              "A backup will be created before proceeding. Continue?"
        
        reply = QMessageBox.warning(self, "FULL DATA IMPORT", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.No:
            return

        path = QFileDialog.getOpenFileName(self, "Select Excel Backup", "", "Excel Files (*.xlsx *.xlsm)")[0]
        if not path:
            return

        # 1. Read Sheet Names FIRST (Lightweight)
        try:
            xls = pd.ExcelFile(path)
            sheet_names = xls.sheet_names
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not read Excel file:\n{e}")
            return
            
        # 2. Show Selection Dialog
        dlg = ImportSelectionDialog(sheet_names, self)
        if not dlg.exec():
            return
            
        selected_sheets = dlg.selected_sheets
        if not selected_sheets:
            return

        # 3. Auto Backup
        self.manual_backup() 
        QMessageBox.information(self, "Backup Created", "A safety backup has been created in the Backups folder.\n\nThe import will now begin. Please wait and do not force quit.")

        # 4. Setup Progress Dialog
        self.progress = QProgressDialog("Initializing Import...", "Cancel", 0, 100, self)
        self.progress.setWindowTitle("Importing Data")
        self.progress.setWindowModality(Qt.WindowModality.WindowModal)
        self.progress.setMinimumDuration(0)
        self.progress.setCancelButton(None) # Disable cancel to prevent corruption
        self.progress.setValue(0)
        
        # 5. Start Worker
        self.import_worker = ImportWorker(path, DB_FILE, selected_sheets)
        self.import_worker.progress_update.connect(lambda msg, val: (self.progress.setLabelText(msg), self.progress.setValue(val)))
        self.import_worker.finished_success.connect(self.on_import_success)
        self.import_worker.finished_error.connect(self.on_import_error)
        self.import_worker.start()

    def on_import_success(self):
        self.progress.setValue(100)
        self.progress.close()
        
        # 0=MasterSheet, 1=Search, 2=Lists, 3=OldSheets, 4=Settings
        self.current_page_index = 0 
        
        # Pagination State
        self.current_page = 1
        self.rows_per_page = 100
        self.total_records = 0
        self.total_pages = 1
        
        # Reload Everything
        self.init_db_and_migrate()
        self.load_data()
        self.load_master_tabs_data()
        
        if hasattr(self, 'old_sheets_page'):
             self.old_sheets_page.load_sheets()
             self.old_sheets_page.update_font_size(self.current_font_size)
        
        self.load_settings() # Reload settings (themes etc might be reset if DB wiped, though usually settings table persisted? Actually we wiped app_settings? No, we didn't wipe app_settings in the worker! Good catch.)
        # Wait, I need to check if we wiped app_settings. In worker I wiped: events, dropdown_options, locations, folder_mappings, tracks, old_sheets_*.
        # I did NOT wipe app_settings. So settings should be fine.
        
        # Reload Generic Sheets Tabs
        self.load_other_sheets()
        
        QMessageBox.information(self, "Import Complete", "The database has been fully rebuilt from the Excel file.")

    def on_import_error(self, err_msg):
        self.progress.close()
        QMessageBox.critical(self, "Import Failed", f"An error occurred during import:\n{err_msg}")

    def import_database(self):
        """Imports a .db file, backing up the current one first, then overwriting."""
        path = QFileDialog.getOpenFileName(self, "Select Database to Restore", "", "SQLite Database (*.db)")[0]
        if not path:
            return
            
        # 1. Validation: Is it a valid SQLite file with our table?
        try:
            conn = sqlite3.connect(path)
            conn.execute("SELECT count(*) FROM events")
            conn.close()
        except Exception as e:
            QMessageBox.critical(self, "Invalid File", f"The selected file does not appear to be a valid archive database.\nError: {e}")
            return
            
        # 2. Confirmation
        msg = "You are about to OVERWRITE the current database with the selected file.\n\n" \
              "This action will replace all current data.\n" \
              "A backup of your current data will be created automatically before proceeding.\n\n" \
              "Are you sure you want to continue?"
        reply = QMessageBox.question(self, "Confirm Restore", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.No:
            return
            
        # 3. Auto-Backup Current
        try:
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            backup_folder = os.path.join(self.backup_dir, "Imports")
            if not os.path.exists(backup_folder):
                os.makedirs(backup_folder)
                
            backup_path = os.path.join(backup_folder, f"pre_import_backup_{timestamp}.db")
            shutil.copy2(DB_FILE, backup_path)
            self.log_message(f"Pre-import backup saved to: {backup_path}")
        except Exception as e:
            QMessageBox.critical(self, "Backup Failed", f"Could not create safety backup. Import aborted.\n{e}")
            return
            
        # 4. Overwrite and Reload
        try:
            # Connect existing connections (if any, though we close them in functions)
            # Just to be safe, we close connection in finally blocks, so we should be good.
            
            shutil.copy2(path, DB_FILE)
            
            # Reload EVERYTHING
            self.init_db_and_migrate() # Ensure schema is good if importing older db
            self.load_data()
            self.load_master_tabs_data()
            self.load_settings_data()
            self.refresh_master_cache()
            self.update_app_font_size(self.current_font_size) # Pass current font size
            
            self.log_message("Database restored successfully.")
            QMessageBox.information(self, "Success", "Database imported and restored successfully.")
            
        except Exception as e:
            QMessageBox.critical(self, "Restoration Failed", f"Critical error during import:\n{e}\n\nPlease check the Backups folder to restore the pre-import backup manually if needed.")

    def load_master_list_with_counts(self, category, table, conn):
        """Load master list with usage statistics across all data sheets."""
        try:
            table.blockSignals(True)
            table.setRowCount(0)
            
            # 1. Get all sheet names with their display names in creation order
            # The old_sheets_meta table has: id (matches table name old_sheet_data_{id}), name (display name)
            sheet_meta_query = "SELECT id, name FROM old_sheets_meta ORDER BY id"  # Order by ID for creation order
            sheet_metadata = conn.execute(sheet_meta_query).fetchall()
            
            # Create mapping: table_name -> display_name and preserve order
            sheet_mapping = {}
            sheet_display_names = []  # Preserve creation order
            for sheet_id, display_name in sheet_metadata:
                table_name = f"old_sheet_data_{sheet_id}"
                sheet_mapping[table_name] = display_name
                sheet_display_names.append(display_name)
            
            # 2. Set up table columns based on category
            if category == "Locations":
                # Columns: Place, Country, Master, [sheet1], [sheet2], ...
                headers = ["Place", "Country", "Master"] + sheet_display_names
                table.setColumnCount(len(headers))
                table.setHorizontalHeaderLabels(headers)
                
                # Query: Get locations with counts
                loc_query = "SELECT id, place, country FROM locations ORDER BY place"
                locations = conn.execute(loc_query).fetchall()
                
                table.setRowCount(len(locations))
                
                for row_idx, (loc_id, place, country) in enumerate(locations):
                    # Column 0: Place
                    place_val = str(place) if place else ""
                    it_place = LisanTableItem(place_val, self.lisan_font_family, self.current_font_size)
                    it_place.setData(Qt.ItemDataRole.UserRole, loc_id)
                    it_place.setData(OriginalValueRole, place_val)
                    table.setItem(row_idx, 0, it_place)
                    
                    # Column 1: Country
                    country_val = str(country) if country else ""
                    it_country = LisanTableItem(country_val, self.lisan_font_family, self.current_font_size)
                    it_country.setData(OriginalValueRole, country_val)
                    table.setItem(row_idx, 1, it_country)
                    
                    # Column 2: Master count (from events table)
                    master_count_query = "SELECT COUNT(*) FROM events WHERE place = ?"
                    master_count = conn.execute(master_count_query, (place_val,)).fetchone()[0]
                    it_master_count = QTableWidgetItem(str(master_count))
                    it_master_count.setFlags(it_master_count.flags() ^ Qt.ItemFlag.ItemIsEditable)  # Read-only
                    it_master_count.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                    table.setItem(row_idx, 2, it_master_count)
                    
                    # Columns 3+: Individual sheet counts
                    for col_offset, display_name in enumerate(sheet_display_names):
                        sheet_count = 0
                        try:
                            # Find the corresponding table name and sheet_id for this display name
                            table_name = [t for t, d in sheet_mapping.items() if d == display_name][0]
                            sheet_id = int(table_name.split('_')[-1])  # Extract ID from old_sheet_data_X
                            
                            # Find which col_X corresponds to the "place" header in this sheet
                            # Query old_sheet_headers to find the col_index for "place"
                            header_query = "SELECT col_index FROM old_sheet_headers WHERE sheet_id = ? AND LOWER(header_name) = 'place'"
                            result = conn.execute(header_query, (sheet_id,)).fetchone()
                            
                            if result:
                                col_index = result[0]
                                col_name = f"col_{col_index}"
                                
                                # Now query the sheet table using the correct column name
                                sheet_count_query = f"SELECT COUNT(*) FROM `{table_name}` WHERE `{col_name}` = ?"
                                sheet_count = conn.execute(sheet_count_query, (place_val,)).fetchone()[0]
                        except Exception:
                            pass  # Silently skip if column doesn't exist
                        
                        it_sheet_count = QTableWidgetItem(str(sheet_count))
                        it_sheet_count.setFlags(it_sheet_count.flags() ^ Qt.ItemFlag.ItemIsEditable)  # Read-only
                        it_sheet_count.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                        table.setItem(row_idx, 3 + col_offset, it_sheet_count)
                        
            else:
                # For dropdown-based categories (Person, Occasion, Category, AV, Incomplete, Created)
                # Columns: [Category Name], Master, [sheet1], [sheet2], ...
                headers = [category, "Master"] + sheet_display_names
                table.setColumnCount(len(headers))
                table.setHorizontalHeaderLabels(headers)
                
                # Map category to event column name
                column_map = {
                    "Person": "person",
                    "Occasion": "occasion",
                    "Category": "category",
                    "AV": "av",
                    "Incomplete": "incomplete",
                    "Created": "created"
                }
                event_column = column_map.get(category, category.lower())
                
                # Get all values for this category
                values_query = "SELECT id, value FROM dropdown_options WHERE category = ? ORDER BY value"
                values = conn.execute(values_query, (category,)).fetchall()
                
                table.setRowCount(len(values))
                
                for row_idx, (val_id, value) in enumerate(values):
                    # Column 0: Category value
                    val_str = str(value) if value else ""
                    it_val = LisanTableItem(val_str, self.lisan_font_family, self.current_font_size)
                    it_val.setData(Qt.ItemDataRole.UserRole, val_id)
                    it_val.setData(OriginalValueRole, val_str)
                    table.setItem(row_idx, 0, it_val)
                    
                    # Column 1: Master count (from events table)
                    try:
                        master_count_query = f"SELECT COUNT(*) FROM events WHERE {event_column} = ?"
                        master_count = conn.execute(master_count_query, (val_str,)).fetchone()[0]
                    except:
                        master_count = 0  # Column doesn't exist in events table
                    
                    it_master_count = QTableWidgetItem(str(master_count))
                    it_master_count.setFlags(it_master_count.flags() ^ Qt.ItemFlag.ItemIsEditable)  # Read-only  
                    it_master_count.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                    table.setItem(row_idx, 1, it_master_count)
                    
                    # Columns 2+: Individual sheet counts
                    for col_offset, display_name in enumerate(sheet_display_names):
                        sheet_count = 0
                        try:
                            # Find the corresponding table name and sheet_id for this display name
                            table_name = [t for t, d in sheet_mapping.items() if d == display_name][0]
                            sheet_id = int(table_name.split('_')[-1])  # Extract ID from old_sheet_data_X
                            
                            # Find which col_X corresponds to the event_column header in this sheet
                            # Query old_sheet_headers to find the col_index for this category's column
                            header_query = "SELECT col_index FROM old_sheet_headers WHERE sheet_id = ? AND LOWER(header_name) = ?"
                            result = conn.execute(header_query, (sheet_id, event_column.lower())).fetchone()
                            
                            if result:
                                col_index = result[0]
                                col_name = f"col_{col_index}"
                                
                                # Now query the sheet table using the correct column name
                                sheet_count_query = f"SELECT COUNT(*) FROM `{table_name}` WHERE `{col_name}` = ?"
                                sheet_count = conn.execute(sheet_count_query, (val_str,)).fetchone()[0]
                        except Exception:
                            pass  # Silently skip if column doesn't exist
                        
                        it_sheet_count = QTableWidgetItem(str(sheet_count))
                        it_sheet_count.setFlags(it_sheet_count.flags() ^ Qt.ItemFlag.ItemIsEditable)  # Read-only
                        it_sheet_count.setTextAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
                        table.setItem(row_idx, 2 + col_offset, it_sheet_count)
            
            # Apply column sizing
            header = table.horizontalHeader()
            # First column(s) - the value column(s)
            if category == "Locations":
                header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)  # Place
                header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)  # Country
                count_start_col = 2
            else:
                header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)  # Value
                count_start_col = 1
            
            # All count columns - resize to contents
            for col_idx in range(count_start_col, len(headers)):
                header.setSectionResizeMode(col_idx, QHeaderView.ResizeMode.ResizeToContents)
                
        except Exception as e:
            print(f"Error loading master list with counts for {category}: {e}")
            traceback.print_exc()
        finally:
            table.blockSignals(False)


    def load_master_tabs_data(self):
        """Loads all master list tabs with usage statistics."""
        try:
            # Added timeout to prevent infinite hang if DB is locked
            conn = sqlite3.connect(DB_FILE, timeout=10)
            
            for cat, tbl in self.master_tab_tables.items():
                try:
                    self.load_master_list_with_counts(cat, tbl, conn)
                except Exception as e:
                    print(f"Error loading {cat} master list: {e}")
                    traceback.print_exc()
                
            conn.close()
        except Exception as e:
            print(f"CRITICAL ERROR in load_master_tabs_data: {e}")
            traceback.print_exc()

    def add_master_tab_item(self, category):
        conn = sqlite3.connect(DB_FILE)
        if category == "Locations":
            conn.execute("INSERT INTO locations (place, country) VALUES ('New Place', 'New Country')")
        else:
            conn.execute("INSERT INTO dropdown_options (category, value) VALUES (?, 'New Item')", (category,))
        conn.commit()
        conn.close()
        self.load_master_tabs_data()
        self.refresh_master_cache()

    def delete_master_tab_item(self, category):
        """Batch delete logic for Master Lists."""
        tbl = self.master_tab_tables[category]
        
        # Get unique rows sorted descending to avoid index shifting issues
        rows = sorted(list(set(index.row() for index in tbl.selectedIndexes())), reverse=True)
        
        if not rows:
            return
            
        count = len(rows)
        msg = f"Are you sure you want to permanently delete {count} entries from the {category} list?"
        reply = QMessageBox.question(self, "Delete Entries", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.No:
            return
            
        # --- OPTIMIZATION START ---
        tbl.setUpdatesEnabled(False)
        tbl.blockSignals(True)
        
        try:
            conn = sqlite3.connect(DB_FILE)
            
            for row in rows:
                db_id = tbl.item(row, 0).data(Qt.ItemDataRole.UserRole)
                if db_id is not None:
                    if category == "Locations":
                        conn.execute("DELETE FROM locations WHERE id=?", (db_id,))
                    else:
                        conn.execute("DELETE FROM dropdown_options WHERE id=?", (db_id,))
                
                # Remove directly from UI instead of reloading everything
                tbl.removeRow(row)
            
            conn.commit()
            conn.close()
        finally:
            tbl.blockSignals(False)
            tbl.setUpdatesEnabled(True)
        # --- OPTIMIZATION END ---
        
        # We still need to update the cache for dropdowns, but no need to reload UI
        self.refresh_master_cache()
        self.refresh_master_lists_action()

    def import_master_list(self, category):
        """Import unique new entries from Excel into the specified Master List."""
        path = QFileDialog.getOpenFileName(self, f"Import {category} List", "", "Excel/CSV (*.xlsx *.csv)")[0]
        if not path:
            return

        try:
            # 1. Read File
            df = pd.read_csv(path) if path.endswith('.csv') else pd.read_excel(path)
            
            # 2. Extract Potential Values
            # Strategy: Look for 'Name', 'Value', 'Place' (if location). Fallback to first column.
            target_col = None
            cols = [str(c).strip().lower() for c in df.columns]
            
            # Identify useful columns
            if category == "Locations":
                # Need Place and Country
                place_col = df.columns[0] # Default
                country_col = None
                
                for c in df.columns:
                    cl = str(c).lower().strip()
                    if "place" in cl or "city" in cl:
                        place_col = c
                    elif "country" in cl:
                        country_col = c
                        
                # Fetch existing locations to avoid duplicates
                conn = sqlite3.connect(DB_FILE)
                existing_places = {r[0].strip().lower() for r in conn.execute("SELECT place FROM locations").fetchall()}
                conn.close()
                
                added_count = 0
                skipped_count = 0
                
                db_data = []
                for _, row in df.iterrows():
                    p = str(row[place_col]).strip()
                    if not p or p.lower() == "nan": continue
                    
                    if p.lower() in existing_places:
                        skipped_count += 1
                        continue
                        
                    c = ""
                    if country_col:
                        c = str(row[country_col]).strip()
                        if c.lower() == "nan": c = ""
                    
                    db_data.append((p, c))
                    added_count += 1
                    existing_places.add(p.lower()) # Add to set to prevent dupes within the file itself
                
                if db_data:
                    conn = sqlite3.connect(DB_FILE)
                    conn.executemany("INSERT INTO locations (place, country) VALUES (?, ?)", db_data)
                    conn.commit()
                    conn.close()
                    
            else:
                # Normal Dropdowns
                col_name = df.columns[0]
                for c in df.columns:
                    if "name" in str(c).lower() or "value" in str(c).lower() or category.lower() in str(c).lower():
                        col_name = c
                        break
                
                conn = sqlite3.connect(DB_FILE)
                # Get existing for this category
                existing = {r[0].strip().lower() for r in conn.execute("SELECT value FROM dropdown_options WHERE category=?", (category,)).fetchall()}
                conn.close()
                
                added_count = 0
                skipped_count = 0
                db_data = []
                
                for _, row in df.iterrows():
                    val = str(row[col_name]).strip()
                    if not val or val.lower() == "nan": continue
                    
                    if val.lower() in existing:
                        skipped_count += 1
                        continue
                    
                    db_data.append((category, val))
                    added_count += 1
                    existing.add(val.lower())
                    
                if db_data:
                    conn = sqlite3.connect(DB_FILE)
                    conn.executemany("INSERT INTO dropdown_options (category, value) VALUES (?, ?)", db_data)
                    conn.commit()
                    conn.close()

            # 3. Finish
            self.refresh_master_lists_action()
            QMessageBox.information(self, "Import Successful", f"Imported: {added_count}\nSkipped (Duplicates): {skipped_count}")

        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import:\n{str(e)}")

    def save_master_tab_edit(self, item, category):
        """Global spelling update logic when a Master list entry is changed."""
        tbl = self.master_tab_tables[category]
        if tbl.is_locked:
            return
            
        db_id = tbl.item(item.row(), 0).data(Qt.ItemDataRole.UserRole)
        if db_id is None:
            return
            
        new_val = item.text().strip()
        old_val = item.data(OriginalValueRole) # Get Original from custom role
        
        if not old_val: # Fallback if missing
             old_val = item.data(Qt.ItemDataRole.DisplayRole)
        
        if old_val == new_val:
            return

        conn = sqlite3.connect(DB_FILE)
        if category == "Locations":
            p = tbl.item(item.row(), 0).text()
            c = tbl.item(item.row(), 1).text()
            conn.execute("UPDATE locations SET place=?, country=? WHERE id=?", (p, c, db_id))
        else:
            conn.execute("UPDATE dropdown_options SET value=? WHERE id=?", (new_val, db_id))
        conn.commit()
        
        # Update Original Value Role so next edit is tracked correctly
        item.setData(OriginalValueRole, new_val)

        col_name = category.lower()
        if category == "AV":
            col_name = "AV"
        elif category == "Locations":
            col_name = "place"
        elif category == "Incomplete":
            col_name = "incomplete"
            
        try:
            check_query = f"SELECT COUNT(*) FROM events WHERE {col_name} = ?"
            count = conn.execute(check_query, (old_val,)).fetchone()[0]
            if count > 0:
                # Wording fixed as requested
                msg = f"Found {count} entries in the Master Sheet using the spelling '{old_val}'.\n\nWould you like to update all existing entries in the Master Sheet to match this new Change?"
                reply = QMessageBox.question(self, "Global Change", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                if reply == QMessageBox.StandardButton.Yes:
                    
                    # --- UNDO CAPTURE ---
                    self.save_undo_state()
                    # --------------------
                    
                    conn.execute(f"UPDATE events SET {col_name} = ? WHERE {col_name} = ?", (new_val, old_val))
                    conn.commit()
                    self.load_data()
        except:
            pass

        conn.close()
        self.refresh_master_cache()
        item.setData(Qt.ItemDataRole.DisplayRole, new_val)

        
        # Update Row Count
        if hasattr(self, 'update_row_count'):
             self.update_row_count()

        # Update Row Count
        if hasattr(self, 'update_row_count'):
             self.update_row_count()

    def add_rows_master(self, count):
        """Adds specified number of rows."""
        
        # --- UNDO CAPTURE ---
        self.save_undo_state()
        
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        
        # 1. Next Audio No
        max_no = 0
        try:
            cursor.execute("SELECT MAX(audio_no) FROM events")
            result = cursor.fetchone()
            if result and result[0] is not None:
                max_no = int(result[0])
        except:
            pass
        
        # 2. Get Last Entry for Smart Fill
        prev_data = {}
        try:
            res = cursor.execute("SELECT * FROM events WHERE audio_no = ?", (max_no,)).fetchone()
            if res:
                for i in range(16):
                     prev_data[i+1] = res[i]
        except:
            pass
            
        # 3. Prepare Batch Data
        new_rows = []
        start_no = max_no + 1
        
        # Pre-calculate reusable values
        p_person = self.repeat_person if self.repeat_person != "Empty" and self.repeat_person != "Repeat Previous" else ""
        if self.repeat_person == "Repeat Previous": p_person = str(prev_data.get(2, ""))
        
        p_occasion = self.repeat_occasion if self.repeat_occasion != "Empty" and self.repeat_occasion != "Repeat Previous" else ""
        if self.repeat_occasion == "Repeat Previous": p_occasion = str(prev_data.get(3, ""))
        
        p_cat = self.repeat_category if self.repeat_category != "Empty" and self.repeat_category != "Repeat Previous" else ""
        if self.repeat_category == "Repeat Previous": p_cat = str(prev_data.get(4, ""))
        
        p_place = ""
        p_country = ""
        if self.repeat_place == "Repeat Previous":
            p_place = str(prev_data.get(5, ""))
            p_country = str(prev_data.get(6, ""))
        elif self.repeat_place != "Empty":
            p_place = self.repeat_place
            c_res = cursor.execute("SELECT country FROM locations WHERE place = ?", (p_place,)).fetchone()
            if c_res: p_country = c_res[0]
            
        p_h_date = str(prev_data.get(7, "")) if self.repeat_date == "Repeat Previous" else ""
        p_e_date = str(prev_data.get(8, "")) if self.repeat_date == "Repeat Previous" else ""
        p_year = str(prev_data.get(9, "")) if self.repeat_date == "Repeat Previous" else ""
        
        p_outof = self.repeat_outof if self.repeat_outof != "Empty" and self.repeat_outof != "Repeat Previous" else ""
        if self.repeat_outof == "Repeat Previous": p_outof = str(prev_data.get(10, ""))
        
        p_av = self.repeat_a_v if self.repeat_a_v != "Empty" and self.repeat_a_v != "Repeat Previous" else ""
        if self.repeat_a_v == "Repeat Previous": p_av = str(prev_data.get(13, ""))

        for i in range(count):
            row_vals = [""] * 16
            row_vals[0] = str(start_no + i)
            row_vals[1] = p_person
            row_vals[2] = p_occasion
            row_vals[3] = p_cat
            row_vals[4] = p_place
            row_vals[5] = p_country
            row_vals[6] = p_h_date
            row_vals[7] = p_e_date
            row_vals[8] = p_year
            row_vals[9] = p_outof
            row_vals[12] = p_av
            new_rows.append(tuple(row_vals))

        # 4. Bulk Insert
        placeholders = ",".join(["?"] * 16)
        cursor.executemany(f"INSERT INTO events VALUES ({placeholders})", new_rows)
        conn.commit()
        conn.close() 

        # 5. Refresh Logic
        # Jump to Last Page
        self.load_data() 
        
        self.current_page = 999999
        self.load_data()
        self.table.scrollToBottom()
        QMessageBox.information(self, "Success", f"Added {count} new entries.")

    def refresh_master_cache(self):
        """Populate drop-downs from DB."""
        conn = sqlite3.connect(DB_FILE)
        
        # Load Defaults
        self.master_lists = {}
        # Keys: 2:Person, 3:Occasion, 4:Category, 5:Place, 9:Year, 13:AV, 16:Incomplete
        
        # 1. Dropdown Options (Person, Occasion, Category, AV)
        # Note: 'AV' is strictly from master list? Or from events?
        # User said "AV and incomplete has a dropdown and so it has a master list."
        # Assuming 'AV' category exists in dropdown_options.
        
        rows = conn.execute("SELECT category, value FROM dropdown_options ORDER BY value").fetchall()
        
        self.master_lists[2] = [] # Person
        self.master_lists[3] = [] # Occasion
        self.master_lists[4] = [] # Category
        self.master_lists[13] = [] # AV
        
        for cat, val in rows:
            if cat == "Person": self.master_lists[2].append(val)
            elif cat == "Occasion": self.master_lists[3].append(val)
            elif cat == "Category": self.master_lists[4].append(val)
            elif cat in ["AV", "A/V"]: self.master_lists[13].append(val)

        # 2. Place (from locations)

        places = conn.execute("SELECT DISTINCT place FROM locations ORDER BY place").fetchall()
        self.master_lists[5] = [p[0] for p in places if p[0]]
        
        # 3. Year (from events distinct)
        years = conn.execute("SELECT DISTINCT year FROM events WHERE year IS NOT NULL AND year != '' ORDER BY year").fetchall()
        self.master_lists[9] = [str(y[0]) for y in years]
        
        # 4. Incomplete (from events distinct OR master list?)
        # User said "AV and incomplete has a dropdown and so it has a master list."
        # If Incomplete has a master list, check dropdown_options 'Incomplete'.
        # If not found, use distinct values.
        
        inc_opts = [val for cat, val in rows if cat == "Incomplete"]
        if inc_opts:
            self.master_lists[16] = inc_opts
        else:
            # Fallback
            incs = conn.execute("SELECT DISTINCT incomplete FROM events WHERE incomplete IS NOT NULL AND incomplete != ''").fetchall()
            self.master_lists[16] = [str(i[0]) for i in incs]

        # Add "(Blanks)" option to all lists
        for k in self.master_lists:
            if "(Blanks)" not in self.master_lists[k]:
                self.master_lists[k].insert(0, "(Blanks)")
        
        conn.close()

    def smart_copy_selection(self):
        """Copies selected entries to Desktop/Audio MP3 sorted by Year/Date."""
        selected_rows = sorted(set(index.row() for index in self.table.selectedIndexes()))
        if not selected_rows:
            QMessageBox.warning(self, "No Selection", "Please select rows to copy.")
            return

        base_dest = os.path.expanduser("~/Desktop/Audio MP3")
        if not os.path.exists(base_dest):
            try:
                os.makedirs(base_dest)
            except OSError as e:
                QMessageBox.critical(self, "Error", f"Could not create destination folder:\n{base_dest}\n{e}")
                return

        # 1. Gather Data
        entries = []
        conn = sqlite3.connect(DB_FILE)
        
        for row in selected_rows:
            # Audio No (Col 1)
            item_id = self.table.item(row, 1)
            audio_no = item_id.text() if item_id else None
            
            # Year (Col 9)
            item_year = self.table.item(row, 9)
            year = item_year.text().strip() if item_year else ""
            
            # Hijri Date (Col 7) - Expected DD/MM/YYYY
            item_date = self.table.item(row, 7)
            h_date = item_date.text().strip() if item_date else ""
            
            if not audio_no: continue
            
            # Parse Date for Sorting
            sort_date = (9999, 12, 31) # Default to end if invalid
            day_str, month_str = "00", "00"
            
            if h_date:
                try:
                    parts = h_date.split('/')
                    if len(parts) == 3:
                        d, m, y = map(int, parts)
                        sort_date = (y, m, d)
                        day_str = f"{d:02d}"
                        month_str = f"{m:02d}"
                except:
                    pass
            
            # Fallback Year if not in date
            if not year and sort_date[0] != 9999:
                year = str(sort_date[0])
            if not year: year = "Unknown"
            
            folder_name = f"{year}_{day_str}-{month_str}"
            
            # Resolve Source Path
            src_path = None
            try:
                # Reuse logic from open_entry_folder but direct DB query
                mapping = conn.execute("SELECT full_link FROM folder_mappings WHERE mapping_type='folder' AND start_no <= ? ORDER BY start_no DESC LIMIT 1", (audio_no,)).fetchone()
                if mapping:
                    root_path = urllib.parse.unquote(str(mapping[0]).replace('file://', ''))
                    if os.path.exists(root_path):
                        # Search for specific folder
                        for entry in os.listdir(root_path):
                            if entry.startswith('.'): continue
                            full_entry = os.path.join(root_path, entry)
                            if not os.path.isdir(full_entry): continue
                            
                            found = False
                            if entry == str(audio_no): found = True
                            elif entry.startswith(f"{audio_no} ") or entry.startswith(f"{audio_no}.") or entry.startswith(f"{audio_no}_"): found = True
                            elif entry.startswith(f"{str(audio_no).zfill(2)} ") or entry.startswith(f"{str(audio_no).zfill(3)} "): found = True
                            
                            if found:
                                src_path = full_entry
                                break
                        
                        # If strict subfolder not found, use root? No, risky. 
                        # User wants "copy the main entry folders".
                        if not src_path:
                             # Fallback: If the root path ITSELF ends with the audio number (unlikely but possible) elements
                             pass
            except Exception as e:
                print(f"Path resolution error for {audio_no}: {e}")
                
            entries.append({
                'audio_no': audio_no,
                'sort_key': (year, sort_date),
                'folder_name': folder_name,
                'src_path': src_path
            })
            
        conn.close()
        
        # 2. Sort
        entries.sort(key=lambda x: x['sort_key'])
        
        # 3. Copy
        progress = QProgressDialog("Copying Folders...", "Cancel", 0, len(entries), self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        
        success_count = 0
        fail_count = 0
        skipped_count = 0
        
        for i, entry in enumerate(entries):
            if progress.wasCanceled():
                break
            
            progress.setValue(i)
            audio_no = entry['audio_no']
            src = entry['src_path']
            
            if not src:
                skipped_count += 1
                print(f"Skipping {audio_no}: Source not found.")
                continue
                
            dest_dir = os.path.join(base_dest, entry['folder_name'])
            
            # Ensure YYYY_DD-MM folder exists
            if not os.path.exists(dest_dir):
                os.makedirs(dest_dir)
                
            # Copy source folder INTO dest_dir
            # e.g. Dest/2023_01-01/105 My Event
            try:
                folder_basename = os.path.basename(src)
                final_dest = os.path.join(dest_dir, folder_basename)
                
                if os.path.exists(final_dest):
                    # Already exists
                    skipped_count += 1
                else:
                    shutil.copytree(src, final_dest)
                    success_count += 1
            except Exception as e:
                print(f"Copy Failed {audio_no}: {e}")
                fail_count += 1
                
        progress.setValue(len(entries))
        
        QMessageBox.information(self, "Smart Copy Complete", 
                                f"Processed {len(entries)} entries.\n"
                                f"Copied: {success_count}\n"
                                f"Skipped: {skipped_count}\n"
                                f"Failed: {fail_count}\n\n"
                                f"Location: {base_dest}")

    def open_advanced_filter(self):
        """Opens the new AdvancedFilterDialog."""
        dlg = AdvancedFilterDialog(self, self.filter_state, self.filter_history, db_path=DB_FILE)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            # Update Filters
            self.filter_state = dlg.local_filters
            
            # Update History
            # Check if this state is already the last one (avoid dupe)
            # Ensure filter_state is NOT empty (and has actual values)
            has_values = any(v for v in self.filter_state.values())
            if self.filter_state and has_values:
                if not self.filter_history or self.filter_history[-1] != self.filter_state:
                    self.filter_history.append(self.filter_state)
                    if len(self.filter_history) > 10:
                        self.filter_history.pop(0) # Keep last 10
            
            # Reload
            self.load_data()

    def open_filter_dialog(self, key, name):
        # Legacy method - redirect to Advanced Filter if ever called
        self.open_advanced_filter()


    def search_entries_master(self):
        """Trigger reload with current filters."""
        # Update Search Term for Highlighting
        if hasattr(self, 'txt_search_master'):
            self.current_search_term = self.txt_search_master.text().strip()
            
        # Reset page to 1 for new search
        self.current_page = 1 
        self.load_data()

    def smart_clear_search(self):
        """Clear search inputs and jump to the currently selected row."""
        button = getattr(self, 'btn_clear_search', None)
        with self.busy_cursor(button):
            # 1. Capture Selection
            target_audio_no = None
            current_row = self.table.currentRow()
            if current_row >= 0:
                item = self.table.item(current_row, 1)
                if item:
                    target_audio_no = item.data(Qt.ItemDataRole.UserRole) or item.text()
            
            # 2. Reset Inputs
            self.txt_search_master.clear()
            self.current_search_term = "" 
            
            if hasattr(self, 'filter_state'):
                self.filter_state = {}
            if hasattr(self, 'clear_all_filter_ui'):
                self.clear_all_filter_ui()
            
            # 3. Reload Data
            self.load_data()
            
            # 4. Restore Selection
            if target_audio_no:
                for r in range(self.table.rowCount()):
                    it = self.table.item(r, 1)
                    if it and it.text() == str(target_audio_no):
                        self.table.setCurrentCell(r, 1)
                        self.table.scrollToItem(it)
                        break

    def toggle_invalid_filter(self):
        """Toggles the display of invalid entries."""
        if self.btn_show_invalid.isChecked():
            self.btn_show_invalid.setStyleSheet("background-color: #feb2b2; color: #9b2c2c; border: 1px solid #742a2a; font-weight: bold;")
        else:
            self.btn_show_invalid.setStyleSheet("")
        self.search_entries_master()

    def scroll_to_top(self):
        self.table.scrollToTop()
        
    def scroll_to_bottom(self):
        self.table.scrollToBottom()


    # --- PAGINATION METHODS ---
    def next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.load_data()

    def prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.load_data()
            
    def update_pagination_ui(self):
        if hasattr(self, 'lbl_page_info'):
             self.lbl_page_info.setText(f"Page {self.current_page} of {self.total_pages}")
             self.btn_prev_page.setEnabled(self.current_page > 1)
             self.btn_next_page.setEnabled(self.current_page < self.total_pages)

    def on_data_loaded(self, rows, total_count, time_taken):
        """Worker Finished Signal Handler."""
        self.is_loading_more = False # Reset flag if we used it (though we aren't using infinite scroll anymore)
        self.loading_worker = None # Cleanup
        
        self.total_records = total_count
        import math
        self.total_pages = math.ceil(self.total_records / self.rows_per_page)
        if self.total_pages < 1: self.total_pages = 1
        
        # UI Updates
        self.update_pagination_ui()
        self.log_message(f"Loaded {len(rows)} rows in {time_taken:.2f}s (Total: {self.total_records})")
        
        self.table.setUpdatesEnabled(False)
        self.table.blockSignals(True)
        # self.table.setSortingEnabled(False) # Already default
        
        self.table.setRowCount(0)
        
        # Populate
        current_rows = self.table.rowCount()
        for idx, row_data in enumerate(rows):
            r = current_rows + idx
            self.table.insertRow(r)
            
            # --- AUDIO NO (Col 1) READ-ONLY ENFORCEMENT ---
            # We need to manually create the item for Audio No if add_table_row doesn't do it specifically or to override it.
            
            self.table.setRowHeight(r, 45) # Fixed height for speed
            self.add_table_row(r, row_data)
            
            # Force Read-Only on Audio No (Col 1)
            item_audio = self.table.item(r, 1)
            if item_audio:
                # Remove ItemIsEditable flag & Enable so it can be selected but not edited
                item_audio.setFlags(item_audio.flags() & ~Qt.ItemFlag.ItemIsEditable) 

        # Update Row Count Logic (Moved here to ensure it runs after count is known)
        max_audio_no = 0
        try:
            # Query MAX audio_no from DB
            conn = sqlite3.connect(DB_FILE)
            max_res = conn.execute("SELECT MAX(audio_no) FROM events").fetchone()
            if max_res and max_res[0]:
                max_audio_no = int(max_res[0])
            conn.close()
        except:
            pass
            
        filtered_count = self.total_records
        if hasattr(self, 'lbl_row_count'):
            self.lbl_row_count.setText(f"Row Count : {filtered_count}/{max_audio_no}")

    # except Exception as e:
    #     print(f"Load Data Error: {e}")
    #     self.table.blockSignals(False)
    #     self.table.setUpdatesEnabled(True)

        self.table.blockSignals(False)
        self.table.setUpdatesEnabled(True)
        # --- RESIZE ROWS ON INITIAL LOAD ---
        # Temporarily Disabled for Stability Check
        # QTimer.singleShot(100, self.table.resizeRowsToContents)

    def clear_filters(self):
        """Resets all filter buttons and inputs."""
        # Reset Search
        if hasattr(self, "txt_search_master") and self.txt_search_master:
             self.txt_search_master.clear()
             
        # Reset logic state
        self.filter_state = {} 
        
        # Reset Buttons Visual State
        # If we have filter buttons stored as attributes:
        # self.btn_filter_person, self.btn_filter_occasion, etc.
        # We should uncheck them if they are checkable.
        filter_buttons = [
            "btn_filter_person", "btn_filter_occasion", "btn_filter_category", 
            "btn_filter_place", "btn_filter_year", "btn_filter_a_v", "btn_filter_incomplete", 
            "btn_show_invalid"
        ]
        for btn_name in filter_buttons:
            if hasattr(self, btn_name):
                btn = getattr(self, btn_name)
                if btn and btn.isCheckable():
                    btn.blockSignals(True)
                    btn.setChecked(False)
                    btn.blockSignals(False)

    def load_data(self):
        """Loads data from the events table into the main ArchiveTable."""
        button = getattr(self, 'btn_search_master', None)
        
        # Set loading flag to prevent resize triggers
        self._is_loading_data = True
        
        try:
            with self.busy_cursor(button):
                self.table.blockSignals(True)
                self.table.setUpdatesEnabled(False)
                self.table.setRowCount(0)
                self.open_tracks = {}

                try:
                    conn = sqlite3.connect(DB_FILE)
                    cursor = conn.cursor()
                    
                    query = "SELECT audio_no, person, occasion, category, place, country, hijri_date, esavi_date, year, out_of, remarks, Tracks, AV, cass_no, came_from, incomplete FROM events"
                    conditions = []
                    params = []
                    
                    if hasattr(self, 'txt_search_master') and self.txt_search_master.text().strip():
                        term = f"%{self.txt_search_master.text().strip()}%"
                        sub_conds = []
                        for col in ["person", "occasion", "category", "place", "remarks", "year", "audio_no"]:
                            sub_conds.append(f"{col} LIKE ?")
                            params.append(term)
                        conditions.append(f"({' OR '.join(sub_conds)})")
                        
                    filter_map = { 2: 'person', 3: 'occasion', 4: 'category', 5: 'place', 6: 'country', 9: 'year', 13: 'AV', 16: 'incomplete' }
                    if hasattr(self, 'filter_state'):
                        for key, col_name in filter_map.items():
                            if key in self.filter_state and self.filter_state[key]:
                                items = list(self.filter_state[key])
                                has_blanks = "(Blanks)" in items
                                real_items = [x for x in items if x != "(Blanks)"]
                                sub_conds = []
                                if real_items:
                                    placeholders = ",".join(["?"] * len(real_items))
                                    sub_conds.append(f"{col_name} IN ({placeholders})")
                                    params.extend(real_items)
                                if has_blanks:
                                    sub_conds.append(f"({col_name} IS NULL OR {col_name} = '')")
                                if sub_conds:
                                    conditions.append(f"({' OR '.join(sub_conds)})")
                    
                    if conditions:
                        query += " WHERE " + " AND ".join(conditions)
                    query += " ORDER BY CAST(audio_no AS INTEGER) ASC" 
                    
                    cursor.execute(query, params)
                    rows = cursor.fetchall()
                    
                    # Populate Table
                    self.table.setRowCount(len(rows))
                    for i, row_data in enumerate(rows):
                        self.add_table_row(i, row_data)
                        if i % 500 == 0: QApplication.processEvents()
                    conn.close()
                    
                    # Restore Widths
                    if hasattr(self, 'saved_col_widths') and self.saved_col_widths and len(self.saved_col_widths) == self.table.columnCount():
                         for i, w in enumerate(self.saved_col_widths):
                             if w > 0: self.table.setColumnWidth(i, w)
                    else:
                         self.table.resizeColumnsToContents()
                    
                    # Optimized Row Resizing - Always recalculate first time, then use cache
                    if not self.row_height_cache:
                        # First load - calculate all heights
                        self.table.resizeRowsToContents()
                        for r in range(self.table.rowCount()):
                            it = self.table.item(r, 1)
                            if it: self.row_height_cache[it.text()] = self.table.rowHeight(r)
                    else:
                        # Use cached heights but validate they match current content
                        applied_count = 0
                        calculated_count = 0
                        for r in range(self.table.rowCount()):
                            it = self.table.item(r, 1)
                            if it:
                                cached_height = self.row_height_cache.get(it.text())
                                if cached_height:
                                    self.table.setRowHeight(r, cached_height)
                                    applied_count += 1
                                else:
                                    # New row not in cache - calculate and add
                                    self.table.resizeRowToContents(r)
                                    self.row_height_cache[it.text()] = self.table.rowHeight(r)
                                    calculated_count += 1

                except Exception as e:
                    print(f"Error loading Master data: {e}")
                    traceback.print_exc()
                finally:
                    self.table.blockSignals(False)
                    self.table.setUpdatesEnabled(True)

                # Lazy Update Others (Background Tabs)
                self.filter_other_sheets(lazy=True)
                
                # Initialize column width tracking AFTER load to prevent spurious resize triggers
                self._last_column_widths = [self.table.columnWidth(i) for i in range(self.table.columnCount())]
        finally:
            # Always clear the loading flag
            self._is_loading_data = False
            
    def resize_master_rows(self):
        """Responsive row resizing for Master table with progress updates."""
        # Guard: Don't resize during data loading
        if self._is_loading_data:
            return
        
        # Guard: Check if column widths actually changed
        current_widths = [self.table.columnWidth(i) for i in range(self.table.columnCount())]
        if hasattr(self, '_last_column_widths') and current_widths == self._last_column_widths:
            return
        
        # Determine which columns changed
        columns_changed = set()
        if hasattr(self, '_last_column_widths') and len(self._last_column_widths) == len(current_widths):
            for i, (old_w, new_w) in enumerate(zip(self._last_column_widths, current_widths)):
                if old_w != new_w:
                    columns_changed.add(i)
        
        self._last_column_widths = current_widths[:]
        
        try:
            self.statusBar().showMessage("Updating row heights...")
            self.table.setUpdatesEnabled(False)
            
            row_count = self.table.rowCount()
            rows_updated = 0
            
            for r in range(row_count):
                # Get audio_no for cache lookup
                it = self.table.item(r, 1)
                if not it:
                    continue
                    
                audio_no = it.text()
                
                # Only recalculate if we don't have a reliable cached value
                # OR if the Remarks column (which has wrapped text) changed width
                remarks_col_changed = 11 in columns_changed  # Column 11 is "Remarks"
                
                if audio_no not in self.row_height_cache or remarks_col_changed:
                    # Recalculate this row
                    self.table.resizeRowToContents(r)
                    new_height = self.table.rowHeight(r)
                    self.row_height_cache[audio_no] = new_height
                    rows_updated += 1
                # else: keep the cached height, don't recalculate
                
                if r % 100 == 0:
                    QApplication.processEvents()
                    self.statusBar().showMessage(f"Updating row heights... {r}/{row_count}")
            
            # Save the updated cache to database
            if hasattr(self, 'row_height_cache') and self.row_height_cache and rows_updated > 0:
                import json
                cache_json = json.dumps(self.row_height_cache)
                self.save_setting("ui_row_height_cache_json", cache_json)
                    
        finally:
            self.table.setUpdatesEnabled(True)
            self.statusBar().showMessage(f"Updated {rows_updated} row heights", 2000)

        if hasattr(self, 'update_lock_state'):
            self.update_lock_state()
        self.update_row_count()

    def update_row_count(self):
        """Updates the total rows label with visible count."""
        if hasattr(self, 'lbl_row_count'):
            visible_count = 0
            search_text = self.txt_search_master.text().strip() if hasattr(self, 'txt_search_master') else ""
            filter_active = any(v for v in self.filter_state.values()) if hasattr(self, 'filter_state') else False
            
            if not search_text and not filter_active:
                visible_count = self.table.rowCount()
            else:
                for r in range(self.table.rowCount()):
                     if not self.table.isRowHidden(r):
                         visible_count += 1
            
            max_audio_no = 0
            if hasattr(self, 'max_audio_no_cache') and self.max_audio_no_cache:
                max_audio_no = self.max_audio_no_cache
            else:
                  try:
                     conn = sqlite3.connect(DB_FILE)
                     max_res = conn.execute("SELECT MAX(audio_no) FROM events").fetchone()
                     if max_res and max_res[0]:
                         max_audio_no = int(max_res[0])
                         self.max_audio_no_cache = max_audio_no
                     conn.close()
                  except: pass
            
            self.lbl_row_count.setText(f"Row Count : {visible_count}/{max_audio_no}")

    def on_sheets_tab_changed(self, index):
        """Lazy update of background sheets when they become visible."""
        if index > 0 and self.global_filters_dirty.get(index):
            self.filter_single_sheet(index)

    def filter_other_sheets(self, lazy=False):
        """Applies current global filters logic."""
        if not hasattr(self, 'sheets_tabs'): return
        current_idx = self.sheets_tabs.currentIndex()
        for i in range(1, self.sheets_tabs.count()):
            self.global_filters_dirty[i] = True
            
        if lazy:
            if current_idx > 0:
                self.filter_single_sheet(current_idx)
            return

        # Explicit Full Update (e.g. from a "Refresh" button if ever needed)
        total_sheets = self.sheets_tabs.count() - 1
        progress = QProgressDialog("Updating all sheets...", "Cancel", 0, total_sheets, self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()
        
        for i in range(1, self.sheets_tabs.count()):
            if progress.wasCanceled(): break
            progress.setValue(i-1)
            self.filter_single_sheet(i)
        progress.setValue(total_sheets)

    def filter_single_sheet(self, tab_index):
        """Filters a specific sheet tab based on current global filter state."""
        table = self.sheets_tabs.widget(tab_index)
        if not isinstance(table, OtherSheetTable): return

        self.global_filters_dirty[tab_index] = False # Clear flag early
        
        # Pre-calculate Filters
        active_filters = []
        filter_active = False
        master_map = { 2: "Person", 3: "Occasion", 4: "Category", 5: "Place", 6: "Country", 9: "Year", 13: "AV", 16: "Incomplete" }
        
        if hasattr(self, 'filter_state'):
            for m_idx, values in self.filter_state.items():
                if values:
                    m_name = master_map.get(m_idx)
                    if m_name:
                        active_filters.append((m_name.lower(), values))
                        filter_active = True

        table.setUpdatesEnabled(False)
        try:
            if not filter_active:
                # Clear all filters
                for r in range(table.rowCount()):
                    table.setRowHidden(r, False)
                    if r % 2000 == 0: QApplication.processEvents()
            else:
                # Apply Filters
                if not hasattr(table, 'col_map_cache'):
                    table.col_map_cache = {}
                    for c in range(table.columnCount()):
                        it = table.horizontalHeaderItem(c)
                        if it: table.col_map_cache[it.text().lower()] = c
                
                table_filters = []
                for col_name, values in active_filters:
                    t_col = table.col_map_cache.get(col_name)
                    if t_col is not None:
                        table_filters.append((t_col, values))

                for r in range(table.rowCount()):
                    should_hide = False
                    for f_col, f_vals in table_filters:
                        base_item = table.item(r, f_col)
                        val = base_item.text().strip() if base_item else ""
                        if not val: val = "(Blanks)"
                        if val not in f_vals:
                            should_hide = True
                            break
                    table.setRowHidden(r, should_hide)
                    if r % 2000 == 0: QApplication.processEvents()
        finally:
            table.setUpdatesEnabled(True)

    def add_table_row(self, row_idx, row_data):
        """Populates a row. DB has 16 cols. Table has 17 cols (Plus Col at 0)."""
        # 0: Plus Button
        it0 = QTableWidgetItem("➕")
        it0.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        it0.setFlags(Qt.ItemFlag.ItemIsEnabled | Qt.ItemFlag.ItemIsSelectable) # Not editable
        self.table.setItem(row_idx, 0, it0)
        
        # Data Columns ( 1 to 16 )
        # DB: audio_no, person, occasion, category, place, country, hijri, esavi, year, out_of, remarks, Tracks, "A/V", cass, came, inc
        for db_col_idx, val in enumerate(row_data):
            table_col_idx = db_col_idx + 1
            if table_col_idx >= self.table.columnCount(): break
            
            final_val = str(val) if val is not None else ""
            item = LisanTableItem(final_val, self.lisan_font_family, self.current_font_size)
            
            # Store ID in Audio No column (Col 1)
            if db_col_idx == 0: 
                item.setData(Qt.ItemDataRole.UserRole, final_val)
            
            # Validation
            self.validate_table_item(item, table_col_idx)
                
            self.table.setItem(row_idx, table_col_idx, item)

    def validate_table_item(self, item, col_idx=None):
        """Checks if item value is valid based on dropdown lists."""
        if not item: return
        col = col_idx if col_idx is not None else item.column()
        # Col indices with dropdowns: 2:person, 3:occasion, 4:category, 5:place, 6:country, 13:AV, 16:incomplete
        if col in self.master_lists:
            val = item.text().strip()
            if val and val not in self.master_lists[col]:
                # INVALID: Turn Red
                item.setForeground(QColor("#e53e3e")) # Bright Red
                f = item.font()
                f.setBold(True)
                item.setFont(f)
            else:
                # VALID: Reset to default color (based on theme)
                color = "white" if self.is_dark_mode else "#2d3748"
                item.setForeground(QColor(color))
                f = item.font()
                f.setBold(False)
                item.setFont(f)


    def on_table_cell_clicked(self, row, col):
        """Handles click on the Expand (+) button column."""
        if col == 0:
            item = self.table.item(row, 0)
            if not item: return
            
            audio_no_item = self.table.item(row, 1)
            audio_no = audio_no_item.data(Qt.ItemDataRole.UserRole) if audio_no_item else None
            
            if not audio_no: return

            self.table.setUpdatesEnabled(False)
            try:
                text = item.text()
                if "➕" in text:
                    item.setText("➖")
                    self.expand_row(row, audio_no)
                else:
                    item.setText("➕")
                    if row + 1 < self.table.rowCount():
                       # Small check to ensure we are removing a widget row
                       if self.table.cellWidget(row + 1, 0) or self.table.columnSpan(row+1, 0) > 1:
                            self.table.removeRow(row + 1)
                            if audio_no in self.open_tracks:
                                del self.open_tracks[audio_no]
            finally:
                self.table.setUpdatesEnabled(True)
                self.table.viewport().update()

    def expand_row(self, parent_row, audio_no):
        """Inserts a new row below parent_row and sets widget."""
        expansion_row = parent_row + 1
        self.table.insertRow(expansion_row)
        
        # Create Widget
        tr_widget = TrackManagerWidget(audio_no, self)
        self.open_tracks[audio_no] = tr_widget
        
        # Span columns
        self.table.setSpan(expansion_row, 0, 1, self.table.columnCount())
        
        # Set Widget
        self.table.setCellWidget(expansion_row, 0, tr_widget)
        
        # Set Height
        self.table.setRowHeight(expansion_row, 400) # Initial Height

    def handle_item_changed(self, item):
        if not item: return
        row = item.row()
        col = item.column()
        
        # Auto-Resize Row logic
        self.table.resizeRowToContents(row)
        
        # Validation
        self.validate_table_item(item, col)
        
        self.mark_unsaved()
            

    def force_layout_update(self):

        """Forces table to recalculate layout and scroll."""
        self.table.viewport().update()
        self.table.updateGeometry()
        # Trigger a minor resize to force calculation
        s = self.table.size()
        self.table.resize(s.width(), s.height() - 1)
        self.table.resize(s.width(), s.height())
        
        if self.table.rowCount() > 0:
            self.table.scrollToBottom()
        self.table.blockSignals(False)
        self.table.setUpdatesEnabled(True)


    def prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.load_data()
            
    def next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.load_data()
            
    def goto_page(self):
        text = self.txt_page.text()
        if text.isdigit():
            pg = int(text)
            if 1 <= pg <= self.total_pages:
                self.current_page = pg
                self.load_data()
            else:
                self.txt_page.setText(str(self.current_page))
        else:
            self.txt_page.setText(str(self.current_page))



    # --- UNDO SYSTEM STUBS ---
    def save_undo_state(self, target_widget=None):
        """Placeholder to prevent crashes. Feature not fully implemented."""
        pass

    def push_undo_action(self, action_dict):
        """Placeholder to prevent crashes. Feature not fully implemented."""
        # print(f"Undo Action Pushed: {action_dict}") # Debug
        pass

    def insert_event(self, row_data):
        """Inserts a new event row into the database."""
        conn = sqlite3.connect(DB_FILE)
        try:
            placeholders = ",".join(["?"] * 16)
            # row_data should be a list of 16 elements matching the DB schema order
            conn.execute(f"INSERT INTO events VALUES ({placeholders})", tuple(row_data))
            conn.commit()
        except Exception as e:
            self.log_message(f"Insert Error: {e}")
            QMessageBox.warning(self, "Database Error", f"Could not insert row: {e}")
        finally:
            conn.close()

    def update_event_field(self, audio_no, col_idx, value):
        """Updates a specific field for a given audio_no."""
        db_cols = ["", "audio_no", "person", "occasion", "category", "place", "country", 
                   "hijri_date", "esavi_date", "year", "out_of", "remarks", 
                   "Tracks", "AV", "cass_no", "came_from", "incomplete"]
        
        if col_idx < 1 or col_idx >= len(db_cols):
            return

        col_name = db_cols[col_idx]
        conn = sqlite3.connect(DB_FILE)
        try:
            cur = conn.execute(f"UPDATE events SET {col_name} = ? WHERE audio_no = ?", (value, audio_no))
            conn.commit()
            if cur.rowcount == 0:
                pass
            else:
                 # Immediate Verification
                 ver = conn.execute(f"SELECT {col_name} FROM events WHERE audio_no = ?", (audio_no,)).fetchone()
                 
        except Exception as e:
            self.log_message(f"Update Error ({col_name}): {e}")
        finally:
            conn.close()

    def update_audio_no(self, old_no, new_no):
        """Updates the Primary Key audio_no and related foreign keys."""
        conn = sqlite3.connect(DB_FILE)
        try:
             # Check if new_no already exists
            count = conn.execute("SELECT COUNT(*) FROM events WHERE audio_no = ?", (new_no,)).fetchone()[0]
            if count > 0:
                return False, "Audio Number already exists."

            # Update DB
            cursor = conn.cursor()
            cursor.execute("BEGIN TRANSACTION")
            # Update Parent
            cursor.execute("UPDATE events SET audio_no = ? WHERE audio_no = ?", (new_no, old_no))
            # Update Children (manually cascading)
            cursor.execute("UPDATE tracks SET event_id = ? WHERE event_id = ?", (new_no, old_no))
            cursor.execute("COMMIT")
            return True, ""
        except Exception as e:
            if conn: conn.rollback()
            self.log_message(f"Update PK Error: {e}")
            return False, str(e)
        finally:
            conn.close()

    def handle_item_changed(self, item):
        col = item.column()
        row = item.row()
        text = item.text().strip()
        
        if col == 0:
            return

        # Duplicate Audio No protection
        if col == 1:
            # Get original ID from UserRole
            old_id = item.data(Qt.ItemDataRole.UserRole)
            
            if not text:
                # User cleared the cell. Revert to old_id immediately to prevent "datatype mismatch" on PK update.
                # If they want to delete, they should use the Delete button/shortcut.
                self.table.blockSignals(True)
                item.setText(str(old_id) if old_id else "")
                self.table.blockSignals(False)
                return

            if text:
                # Check duplication against DB (excluding self)
                # But we must be careful: if we are renaming, the DB still has OLD ID
                # If we are just refreshing the same ID, it's fine.
                
                if old_id and old_id != text:
                    # User is trying to change ID. Check if new ID exists.
                    conn = sqlite3.connect(DB_FILE)
                    count = conn.execute("SELECT COUNT(*) FROM events WHERE audio_no = ?", (text,)).fetchone()[0]
                    conn.close()
                    
                    if count > 0:
                         QMessageBox.warning(self, "Repeated Entry", f"The Audio Number '{text}' already exists.")
                         self.table.blockSignals(True)
                         item.setText(str(old_id)) # Revert
                         self.table.blockSignals(False)
                         return

        # Validation Logic for Master Lists
        valid = True
        
        # FAILSAFE: If Master Locked, Revert Immediately
        if hasattr(self, 'master_locked') and self.master_locked:
            self.table.blockSignals(True)
            # Revert is hard without old value, but since we are handling KeyPress/Double click prevention,
            # this shouldn't be reached. If it is, we prevent saving.
            self.log_message("Warning: Modification attempted while Locked.")
            self.table.blockSignals(False)
            return
            
        if col in self.master_lists:
            # Check if text is in the list
            if text and text not in self.master_lists[col]:
                valid = False

        self.table.blockSignals(True)
        
        # FIX: Do not replace the item instance while handling its signal (causes SegFault).
        # Update properties in-place.
        if not valid:
            item.setForeground(QColor("#e53e3e"))
            font = QFont(self.lisan_font_family, self.current_font_size)
            font.setBold(True)
            item.setFont(font)
        else:
             # Reset to normal
             # Reset to normal
             # Use default text/foreground color from palette/theme
             item.setData(Qt.ItemDataRole.ForegroundRole, None) # Clear any override
             font = QFont(self.lisan_font_family, self.current_font_size)
             font.setBold(False)
             item.setFont(font)
             
        # SAVE TO DB
        # Audio No is Col 1.
        # But `update_event_field` expects Audio No (PK) to identify row.
        # We need the Audio No of the ROW.
        # If we edited the Audio No itself (Col 1), we use the OLD ID to find it, or if new entry...
        
        # Get Primary Key (Audio No)
        it_id = self.table.item(row, 1)
        audio_no = it_id.text() if it_id else None
        
        # If we edited Audio No (Col 1), `text` IS the new Audio No.
        # We handled duplication check above. 
        # But we need to update the PK in DB.
        
        if col == 1:
            # PK Change
            # `handle_item_changed` already checked duplication.
            # We need the OLD ID to update it.
            # But wait, `update_audio_no` needs old_id.
            old_pk = item.data(Qt.ItemDataRole.UserRole)
            if old_pk and str(old_pk) != text:
                 success, err = self.update_audio_no(old_pk, text)
                 if success:
                     item.setData(Qt.ItemDataRole.UserRole, text) # Update cached ID
                 else:
                     # Revert done in check? No, we should revert here if failed
                     self.table.blockSignals(True)
                     item.setText(str(old_pk))
                     self.table.blockSignals(False)
        else:
            # Normal Field Update
            # Get Audio No from Col 1 (which acts as ID)
            # If Col 1 is empty, we can't update?
            # Assuming row has valid Audio No.
            
            # If we are editing a new empty row, it might not have Audio No yet.
            # But the synchronous load puts data.
            
            # Retrieve PK
            pk_item = self.table.item(row, 1)
            current_pk = pk_item.text() if pk_item else None
            
            if current_pk:
                self.update_event_field(current_pk, col, text)
        
        # self.table.setItem(row, col, new_item) <--- REMOVED TO PREVENT CRASH
        
        # Auto-Calc Dates
        if col == 7 and len(text) == 10:
            esavi = calculate_esavi(text)
            if esavi:
                # Update UI
                self.table.setItem(row, 8, LisanTableItem(esavi, self.lisan_font_family, self.current_font_size))
                self.table.setItem(row, 9, LisanTableItem(text.split('/')[-1], self.lisan_font_family, self.current_font_size))
                         
        elif col == 8 and len(text) == 10:
            hijri = calculate_hijri(text)
            if hijri:
                self.table.setItem(row, 7, LisanTableItem(hijri, self.lisan_font_family, self.current_font_size))
                self.table.setItem(row, 9, LisanTableItem(hijri.split('/')[-1], self.lisan_font_family, self.current_font_size))
                
        # self.table.resizeColumnToContents(col) # REMOVED
        
        # --- RESIZE ROW TO FIT WRAPPED TEXT ---
        self.table.resizeRowToContents(row)
        
        self.table.blockSignals(False)
    
        # --- MODIFIED: Mark as unsaved instead of saving immediately ---
        self.mark_unsaved()

    def handle_cell_clicked(self, r, c, pr, pc):
        pass # Removed save_data trigger on click to prevent full rewrites

    def save_data(self, silent=True):
        self.table.blockSignals(True)
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        try:
            cursor.execute("BEGIN TRANSACTION")
            # CRITICAL FIX: Do NOT delete all events. Only update/insert visible rows.
            # cursor.execute("DELETE FROM events") 
            
            for r in range(self.table.rowCount()):
                cw = self.table.cellWidget(r, 0)
                if not isinstance(cw, TrackManagerWidget):
                    row_data = [self.table.item(r, c).text() if self.table.item(r, c) else "" for c in range(1, self.table.columnCount())]
                    if row_data[0]:
                        # Use INSERT OR REPLACE to update existing or add new, without touching others.
                        cursor.execute("INSERT OR REPLACE INTO events VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", tuple(row_data))
            cursor.execute("COMMIT")
        except:
            pass
        finally:
            conn.close()
            self.table.blockSignals(False)

    def toggle_tracks(self, audio_no, target_table=None):
        if target_table is None:
            target_table = self.table
        target_row = -1
        for row in range(target_table.rowCount()):
            it = target_table.item(row, 1)
            if it and it.text().strip() == str(audio_no):
                target_row = row
                break
        if target_row == -1:
            return
        track_key = f"{id(target_table)}_{audio_no}"
        btn = target_table.cellWidget(target_row, 0)
        if track_key in self.open_tracks:
            target_table.removeRow(target_row + 1)
            btn.setText("+")
            del self.open_tracks[track_key]
        else:
            target_table.insertRow(target_row + 1)
            target_table.setRowHeight(target_row + 1, 400)
            target_table.setSpan(target_row + 1, 0, 1, target_table.columnCount())
            manager = TrackManagerWidget(audio_no, self)
            target_table.setCellWidget(target_row + 1, 0, manager)
            btn.setText("-")
            self.open_tracks[track_key] = manager

            # FIX: Delayed update to ensure layout settles
            QTimer.singleShot(20, lambda: target_table.viewport().update())
            QTimer.singleShot(20, lambda: target_table.updateGeometries())

    def find_active_table(self):
        idx = self.stack.currentIndex()
        if idx == 0:
            return self.table
        if idx == 1:
            # Check which sub-tab is active
            sub_idx = self.tracks_tabs.currentIndex()
            if sub_idx == 0: # Tracks Sheet
                return self.tracks_sheet_table
            elif sub_idx == 1: # Track Search
                return self.track_search_tab.search_results
        return None

    def open_entry_folder(self):
        active_table = self.find_active_table()
        if not active_table or active_table.currentRow() < 0:
            return
        audio_item = active_table.item(active_table.currentRow(), 1)
        if not audio_item:
            return
        
        try:
            audio_no = int(audio_item.text())
        except ValueError:
            return

        conn = sqlite3.connect(DB_FILE)
        mapping = conn.execute("SELECT full_link FROM folder_mappings WHERE mapping_type='folder' AND start_no <= ? ORDER BY start_no DESC LIMIT 1", (audio_no,)).fetchone()
        conn.close()
        
        if mapping:
            base_path = urllib.parse.unquote(str(mapping[0]).replace('file://', ''))
            if not os.path.exists(base_path):
                 QMessageBox.warning(self, "Folder Not Found", f"The mapped folder path does not exist:\n{base_path}")
                 return

            target_path = base_path
            target_found = False
            
            # Attempt to find a specific subfolder for this Audio No
            try:
                for entry in os.listdir(base_path):
                    if entry.startswith('.'): continue
                    full_entry = os.path.join(base_path, entry)
                    if not os.path.isdir(full_entry): continue
                    
                    found = False
                    if entry == str(audio_no): found = True
                    elif entry.startswith(f"{audio_no} ") or entry.startswith(f"{audio_no}.") or entry.startswith(f"{audio_no}_"): found = True
                    elif entry.startswith(f"{str(audio_no).zfill(2)} ") or entry.startswith(f"{str(audio_no).zfill(3)} "): found = True
                    
                    if found:
                        target_path = full_entry
                        target_found = True
                        break
            except Exception as e:
                print(f"Subfolder search error: {e}")

            if not target_found and target_path == base_path:
                 QMessageBox.warning(self, "Folder Not Found", f"Could not find a specific folder for Audio No: {audio_no}\nin {base_path}")
                 return

            subprocess.run(["open", target_path])
        else:
            QMessageBox.warning(self, "Not Mapped", f"No folder mapping found for Audio No: {audio_no}")

    def open_entry_mp3(self):
        active_table = self.find_active_table()
        if not active_table or active_table.currentRow() < 0:
            return
        audio_item = active_table.item(active_table.currentRow(), 1)
        if not audio_item:
            return
        
        try:
            audio_no = int(audio_item.text())
        except ValueError:
            return

        conn = sqlite3.connect(DB_FILE)
        mapping = conn.execute("SELECT full_link FROM folder_mappings WHERE mapping_type='mp3' AND start_no <= ? ORDER BY start_no DESC LIMIT 1", (audio_no,)).fetchone()
        conn.close()
        
        if mapping:
            base_path = urllib.parse.unquote(str(mapping[0]).replace('file://', ''))
            if not os.path.exists(base_path):
                QMessageBox.warning(self, "Path Not Found", f"The mapped folder path does not exist:\n{base_path}")
                return

            target_path = base_path
            target_found = False
            
            # Attempt to find a specific subfolder for this Audio No
            try:
                for entry in os.listdir(base_path):
                    if entry.startswith('.'): continue
                    full_entry = os.path.join(base_path, entry)
                    if not os.path.isdir(full_entry): continue
                    
                    found = False
                    if entry == str(audio_no): found = True
                    elif entry.startswith(f"{audio_no} ") or entry.startswith(f"{audio_no}.") or entry.startswith(f"{audio_no}_"): found = True
                    elif entry.startswith(f"{str(audio_no).zfill(2)} ") or entry.startswith(f"{str(audio_no).zfill(3)} "): found = True
                    
                    if found:
                        target_path = full_entry
                        target_found = True
                        break
            except Exception as e:
                print(f"Subfolder search error: {e}")

            if not target_found:
                 QMessageBox.warning(self, "MP3 Not Found", f"Could not find an MP3 folder for Audio No: {audio_no}\nin {base_path}")
                 return

            subprocess.run(["open", target_path])
        else:
            QMessageBox.warning(self, "Not Mapped", f"No MP3 mapping found for Audio No: {audio_no}")

    def load_settings_data(self):
        self.val_table.blockSignals(True)
        self.val_table.setRowCount(0)
        conn = sqlite3.connect(DB_FILE)
        res = conn.execute("SELECT start_no, full_link FROM folder_mappings WHERE mapping_type='folder' ORDER BY start_no ASC").fetchall()
        for i, db_row in enumerate(res):
            self.val_table.insertRow(i)
            # 2 Columns: Start No, Link
            for j in range(2):
                item = QTableWidgetItem(str(db_row[j]))
                item.setFont(QFont("Arial", 12))
                item.setData(Qt.ItemDataRole.UserRole, db_row[0])
                item.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
                self.val_table.setItem(i, j, item)
        conn.close()
        self.val_table.setColumnCount(2) # Force 2 columns
        self.val_table.setHorizontalHeaderLabels(["Start Number", "File Location"])
        self.val_table.resizeColumnsToContents()
        self.val_table.horizontalHeader().setStretchLastSection(True)
        self.val_table.blockSignals(False)
        self.load_mp3_settings_data()
    
    def jump_to_audio_no(self, target_audio_no):
        """Switches to Master Sheet and scrolls to the specific audio number using Pagination."""

        # Clear search highlights and input
        self.current_search_term = ""
        self.highlight_search = False
        if hasattr(self, 'txt_search_master'):
            self.txt_search_master.clear()
        
        self.switch_page(0) # Go to Master Sheet
        
        target_str = str(target_audio_no).strip()
        
        # 1. Find the global index of this Audio No
        conn = sqlite3.connect(DB_FILE)
        cursor = conn.cursor()
        
        # Check if it exists first
        exists = cursor.execute("SELECT 1 FROM events WHERE audio_no = ?", (target_str,)).fetchone()
        
        if not exists:
            conn.close()
            QMessageBox.information(self, "Not Found", f"Audio No {target_str} not found in the Master Sheet.")
            return
            
        # Determine global index
        # User view is sorted by Audio No ASC.
        # If Audio No is numeric strings ("1", "2", "10"), standard string sort is "1", "10", "2".
        # BUT if user sees "1", "2", "10", then sorting is NUMERIC.
        # We must align our count logic with the sort logic.
        
        try:
            # OPTIMIZED: Remove CAST(audio_no AS INTEGER)
            # Since we validated exists with string lookup, and we want numeric count:
            val_int = int(target_str)
            # We must be careful: audio_no is stored as INTEGER in DB?
            # Schema says: (0, 'audio_no', 'INTEGER', 0, None, 1) -> It IS integer.
            # So direct comparison works if we pass integer.
            
            count = cursor.execute("SELECT COUNT(*) FROM events WHERE audio_no < ?", (val_int,)).fetchone()[0]
        except:
             count = 0
            
        conn.close()
        
        # 2. Setup Table State
        self.current_page = 1 # Not used by load_data but keeps state clean
        
        # Clear Filters to ensure item is visible
        if hasattr(self, 'filter_state'):
            self.filter_state = {}
            if hasattr(self, 'table'):
                 # Reset Header Icons (Visual Only)
                 # Ideally we should call a method to reset headers, 
                 # but for now forcing reload clears the view.
                 pass

        # 3. Load Data
        self.load_data()
        
        row_in_page = count
        
        # 4. Scroll and Select
        # Use QTimer to ensure layout is complete before scrolling/selecting
        def scroll_and_select():
            self.table.clearSelection()
            self.table.clearFocus()
            
            # Find item again
            # Use Column 1 (Audio No) as it holds the key
            item = self.table.item(row_in_page, 1)
            
            # Additional Safety Check: Verify Item Text Matches Target
            # Because if sort order changed or threading issue, row_in_page might be offset.
            # But with single threaded load_data, it should be exact.
            
            if item:
                self.table.scrollToItem(item, QAbstractItemView.ScrollHint.PositionAtCenter)
                self.table.selectRow(row_in_page)
                self.table.setCurrentItem(item)
                self.table.setFocus()
        
        # Reduced delay (was 300) - 100ms should be enough if load_data is blocking
        QTimer.singleShot(100, scroll_and_select)

    def jump_to_old_sheet_entry(self, sheet_id, row_id):
        """Switches to Other Sheets page, finds the correct tab, and scrolls to the row."""
        # 1. Switch to Page 2 (Other Sheets)
        self.switch_page(2)
        
        # 2. Find the correct tab
        # We need to iterate tabs in self.old_sheets_page.tabs
        tabs_widget = self.old_sheets_page.tabs
        target_tab_idx = -1
        target_table = None
        
        for i in range(tabs_widget.count()):
            w = tabs_widget.widget(i)
            # Find the table inside this widget
            tbl = w.findChild(QTableWidget)
            if tbl:
                # Check property "sheet_id"
                sid = tbl.property("sheet_id")
                # Property returns QVariant/int, ensure comparison is safe
                if str(sid) == str(sheet_id):
                    target_tab_idx = i
                    target_table = tbl
                    break
        
        if target_tab_idx == -1 or not target_table:
            QMessageBox.warning(self, "Not Found", "Could not find the target sheet.")
            return

        # 3. Select Tab
        tabs_widget.setCurrentIndex(target_tab_idx)
        
        # 4. Find Row by ROWID (UserRole)
        # Scan table rows
        target_row = -1
        for r in range(target_table.rowCount()):
            # Check col 0 for UserRole
            item = target_table.item(r, 0)
            if item:
                rid = item.data(Qt.ItemDataRole.UserRole)
                if str(rid) == str(row_id):
                    target_row = r
                    break
                    
        if target_row != -1:
            # Scroll and Select
            target_table.clearSelection()
            target_table.clearFocus()
            
            item = target_table.item(target_row, 0)
            target_table.scrollToItem(item, QAbstractItemView.ScrollHint.PositionAtCenter)
            target_table.selectRow(target_row)
            target_table.setCurrentItem(item)
            target_table.setFocus()
        else:
             QMessageBox.warning(self, "Not Found", "Could not find the target entry in the sheet.")

    def handle_settings_double_click(self, row, col):
        if self.val_table.is_locked:
            return
        if col == 1: # Link column
            folder = QFileDialog.getExistingDirectory(self, "Select Folder", "/Volumes")
            if folder:
                self.val_table.item(row, col).setText(f"file://{folder}/")
                # Auto-save happens via on_item_changed or we trigger it explicitly?
                # The existing code calls save_settings_edit in on_item_changed usually
                # But here we set text programmatically, which triggers signal.
                self.save_settings_edit(self.val_table.item(row, col))

    def save_settings_edit(self, item):
        if self.val_table.is_locked:
            return
        db_id = item.data(Qt.ItemDataRole.UserRole)
        if db_id is None:
            return
        conn = sqlite3.connect(DB_FILE)
        r = item.row()
        s, l = [self.val_table.item(r, j).text() for j in range(2)]
        conn.execute("UPDATE folder_mappings SET start_no=?, full_link=? WHERE start_no=? AND mapping_type='folder'", (s, l, db_id))
        conn.commit()
        conn.close()

    def add_settings_item(self):
        dialog = FolderMappingDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            d = dialog.get_data()
            conn = sqlite3.connect(DB_FILE)
            conn.execute("INSERT INTO folder_mappings (start_no, full_link, mapping_type) VALUES (?,?, 'folder')", (d["start_no"], d["link"]))
            conn.commit()
            conn.close()
            self.load_settings_data()

    def delete_settings_item(self):
        row = self.val_table.currentRow()
        if row < 0:
            return
        db_id = self.val_table.item(row, 0).data(Qt.ItemDataRole.UserRole)
        if QMessageBox.question(self, "Confirm", "Delete this mapping?") == QMessageBox.StandardButton.Yes:
            conn = sqlite3.connect(DB_FILE)
            conn.execute("DELETE FROM folder_mappings WHERE start_no=? AND mapping_type='folder'", (db_id,))
            conn.commit()
            conn.close()
            self.load_settings_data()

    def import_folder_info(self):
        path = QFileDialog.getOpenFileName(self, "Select Info Sheet", "", "Excel/CSV (*.xlsx *.csv)")[0]
        if not path:
            return
        try:
            df = pd.read_csv(path) if path.endswith('.csv') else pd.read_excel(path)
            df.columns = [str(c).strip().capitalize() for c in df.columns]
            conn = sqlite3.connect(DB_FILE)
            for _, row in df.iterrows():
                num = int(row.get('Number', row.get('Start_no', 0)))
                link = str(row.get('Link', row.get('File Location', '')))
                if num:
                    conn.execute("INSERT OR REPLACE INTO folder_mappings (start_no, full_link, mapping_type) VALUES (?,?, 'folder')", 
                                 (num, link))
            conn.commit()
            conn.close()
            self.load_settings_data()
        except:
            pass

    def refresh_master_cache(self):
        conn = sqlite3.connect(DB_FILE)
        try:
            for idx, name in [(2, "Person"), (3, "Occasion"), (4, "Category"), (13, "AV"), (16, "Incomplete")]:
                query = "SELECT value FROM dropdown_options WHERE category=?"
                res = conn.execute(query, (name,)).fetchall()
                self.master_lists[idx] = [str(r[0]) for r in res]
            
            # Locations (Place)
            res_place = conn.execute("SELECT DISTINCT place FROM locations ORDER BY place ASC").fetchall()
            self.master_lists[5] = [str(r[0]) for r in res_place if r[0]]

            # Locations (Country) - New for Index 6
            res_country = conn.execute("SELECT DISTINCT country FROM locations WHERE country != '' ORDER BY country ASC").fetchall()
            self.master_lists[6] = [str(r[0]) for r in res_country if r[0]]
            
        except:
            pass
        finally:
            conn.close()
        if hasattr(self, 'track_search_tab'):
            self.track_search_tab.update_dropdown_items()

    def update_delegates(self):
        self.delegates = {
            "text": EditorDelegate(self.lisan_font_family, "text", self.table),
            "person": EditorDelegate(self.lisan_font_family, "person", self.table, self.master_lists),
            "occasion": EditorDelegate(self.lisan_font_family, "occasion", self.table, self.master_lists),
            "category": EditorDelegate(self.lisan_font_family, "category", self.table, self.master_lists),
            "place": EditorDelegate(self.lisan_font_family, "place", self.table, self.master_lists),
            "country": EditorDelegate(self.lisan_font_family, "country", self.table, self.master_lists), # New
            "AV": EditorDelegate(self.lisan_font_family, "AV", self.table, self.master_lists),
            "incomplete": EditorDelegate(self.lisan_font_family, "incomplete", self.table, self.master_lists)
        }
        col_map = {2: "person", 3: "occasion", 4: "category", 5: "place", 6: "country", 13: "AV", 16: "incomplete"} # Changed to 16
        for c in range(len(self.headers)):
            mode = col_map.get(c, "text")
            self.table.setItemDelegateForColumn(c, self.delegates.get(mode, self.delegates["text"]))

    # --- MP3 MAPPING LOGIC ---
    def load_mp3_settings_data(self):
        self.mp3_table.blockSignals(True)
        self.mp3_table.setRowCount(0)
        conn = sqlite3.connect(DB_FILE)
        res = conn.execute("SELECT start_no, full_link FROM folder_mappings WHERE mapping_type='mp3' ORDER BY start_no ASC").fetchall()
        for i, db_row in enumerate(res):
            self.mp3_table.insertRow(i)
            for j in range(2):
                item = QTableWidgetItem(str(db_row[j]))
                item.setFont(QFont("Arial", 12))
                item.setData(Qt.ItemDataRole.UserRole, db_row[0])
                item.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
                self.mp3_table.setItem(i, j, item)
        conn.close()
        self.mp3_table.setColumnCount(2)
        self.mp3_table.setHorizontalHeaderLabels(["Start Number", "File Location"])
        self.mp3_table.resizeColumnsToContents()
        self.mp3_table.horizontalHeader().setStretchLastSection(True)
        self.mp3_table.blockSignals(False)

    def add_mp3_settings_item(self):
        dialog = FolderMappingDialog(self)
        dialog.setWindowTitle("Add Mp3 Mapping")
        if dialog.exec() == QDialog.DialogCode.Accepted:
            d = dialog.get_data()
            conn = sqlite3.connect(DB_FILE)
            conn.execute("INSERT INTO folder_mappings (start_no, full_link, mapping_type) VALUES (?,?, 'mp3')", (d["start_no"], d["link"]))
            conn.commit()
            conn.close()
            self.load_mp3_settings_data()

    def delete_mp3_settings_item(self):
        row = self.mp3_table.currentRow()
        if row < 0:
            return
        db_id = self.mp3_table.item(row, 0).data(Qt.ItemDataRole.UserRole)
        if QMessageBox.question(self, "Confirm", "Delete this Mp3 mapping?") == QMessageBox.StandardButton.Yes:
            conn = sqlite3.connect(DB_FILE)
            conn.execute("DELETE FROM folder_mappings WHERE start_no=? AND mapping_type='mp3'", (db_id,))
            conn.commit()
            conn.close()
            self.load_mp3_settings_data()

    def handle_mp3_settings_double_click(self, row, col):
        if self.mp3_table.is_locked:
            return
        if col == 1:
            folder = QFileDialog.getExistingDirectory(self, "Select Mp3 Folder", "/Volumes")
            if folder:
                self.mp3_table.item(row, col).setText(f"file://{folder}/")
                # Removed drive/name update logic
                self.mp3_table.item(row, 2).setText(parts[-1])
                self.save_mp3_settings_edit(self.mp3_table.item(row, col))

    def save_mp3_settings_edit(self, item):
        if self.mp3_table.is_locked:
            return
            
        db_id = item.data(Qt.ItemDataRole.UserRole)
        # If no DB ID, it might be a new row not yet saved or something went wrong
        if db_id is None:
            return

        # Prevent recursion loop when we update text/data programmatically
        self.mp3_table.blockSignals(True)
        try:
            r = item.row()
            s = self.mp3_table.item(r, 0).text().strip() # Start No
            d = self.mp3_table.item(r, 1).text().strip() # Drive
            f = self.mp3_table.item(r, 2).text().strip() # Folder
            l = self.mp3_table.item(r, 3).text().strip() # Link
            
            # Validate Start No is integer
            try:
                new_start_no = int(s)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Start Number must be an integer.")
                # Revert to old value (db_id) if it was the start_no column that changed
                if item.column() == 0:
                    item.setText(str(db_id))
                return

            conn = sqlite3.connect(DB_FILE)
            try:
                # Attempt Update
                # If db_id (original start_no) != new_start_no, we are renaming the key.
                # The WHERE clause must use the OLD ID (db_id).
                conn.execute("UPDATE folder_mappings SET start_no=?, drive_info=?, folder_name=?, full_link=? WHERE start_no=? AND mapping_type='mp3'", 
                             (new_start_no, d, f, l, db_id))
                conn.commit()
                
                # CRITICAL: If Start No changed, we MUST update the UserRole (DB ID) for ALL columns in this row
                # otherwise subsequent edits will try to update using the old ID and fail silently.
                if new_start_no != db_id:
                    for col in range(4):
                        it = self.mp3_table.item(r, col)
                        if it:
                            it.setData(Qt.ItemDataRole.UserRole, new_start_no)
                    self.log_message(f"Mp3 Mapping updated: {db_id} -> {new_start_no}")
                
            except sqlite3.IntegrityError:
                # likely duplicate start_no
                conn.rollback()
                QMessageBox.warning(self, "Duplicate ID", f"Start Number {new_start_no} already exists. Please choose a unique number.")
                # Revert text if it was the start_no col
                if item.column() == 0:
                    item.setText(str(db_id))
            except Exception as e:
                conn.rollback()
                QMessageBox.critical(self, "Save Error", f"Could not save Mp3 mapping: {e}")
                # Optional: Revert?
            finally:
                conn.close()
                
        finally:
            self.mp3_table.blockSignals(False)

    def import_mp3_folder_info(self):
        path = QFileDialog.getOpenFileName(self, "Select Info Sheet", "", "Excel/CSV (*.xlsx *.csv)")[0]
        if not path:
            return

        # WIPE CONFIRMATION
        msg = "Do you want to WIPE all existing Mp3 mappings before importing?\n\n" \
              "YES: Deletes all current mappings, then imports the file.\n" \
              "NO: Keeps existing mappings and merges/updates them."
        reply = QMessageBox.question(self, "Import Mode", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
        
        if reply == QMessageBox.StandardButton.Cancel:
            return
        
        should_wipe = (reply == QMessageBox.StandardButton.Yes)

        try:
            df = pd.read_csv(path) if path.endswith('.csv') else pd.read_excel(path)
            # Normalize column names: Number, Drive, Location 1, Location 2, Location 3, Link
            df.columns = [str(c).strip() for c in df.columns]
            
            conn = sqlite3.connect(DB_FILE)
            
            if should_wipe:
                conn.execute("DELETE FROM folder_mappings WHERE mapping_type='mp3'")
                self.log_message("Wiped mp3_links table before import.")
            
            count = 0
            for _, row in df.iterrows():
                # 1. Handle Number (safely convert, skip NaNs/bad str)
                raw_num = row.get('Number', row.get('Start_no'))
                try:
                    num = int(float(str(raw_num).strip()))
                except (ValueError, TypeError, AttributeError):
                    continue
                
                # 2. Get Drive
                drive = str(row.get('Drive', '')).strip()
                if drive == 'nan': drive = ""
                
                # 3. Get Folder Name (Priority: Loc 3 -> Loc 2 -> Loc 1)
                loc1 = str(row.get('Location 1', '')).strip()
                loc2 = str(row.get('Location 2', '')).strip()
                loc3 = str(row.get('Location 3', '')).strip()
                
                folder_name = ""
                if loc3 and loc3.lower() != 'nan':
                     folder_name = loc3
                elif loc2 and loc2.lower() != 'nan':
                     folder_name = loc2
                elif loc1 and loc1.lower() != 'nan':
                     folder_name = loc1
                     
                if folder_name == 'nan': folder_name = ""

                # 4. Get Link
                link = str(row.get('Link', '')).strip()
                if link == 'nan': link = ""
                
                if num:
                    conn.execute("INSERT INTO folder_mappings (start_no, drive_info, folder_name, full_link, mapping_type) VALUES (?,?,?,?, 'mp3')", (num, drive, folder_name, link))
                    count += 1
                    
            conn.commit()
            conn.close()
            self.load_mp3_settings_data()
            QMessageBox.information(self, "Import Success", f"Successfully imported {count} mappings.")
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import:\n{str(e)}")

    def load_tracks_sheet_data(self, table):
        """Loads data for the Tracks Sheet with Audio No join."""
        conn = sqlite3.connect(DB_FILE)
        try:
            # Query JOINing tracks and events
            query = """
                SELECT t.id, CAST(e.audio_no AS INTEGER) as audio_num, t.track_no, t.track_name 
                FROM tracks t 
                JOIN events e ON t.event_id = e.audio_no 
                ORDER BY audio_num ASC, CAST(t.track_no AS INTEGER) ASC
            """
            cursor = conn.execute(query)
            rows = cursor.fetchall()
            
            headers = ["ID", "Audio No", "Track No", "Tracks"]
            table.setColumnCount(len(headers))
            table.setHorizontalHeaderLabels(headers)
            table.setRowCount(len(rows))
            
            # Define Column Map for Saving: View Col Index -> DB Col Name
            # 0: ID (Read Only/Hidden), 1: Audio No (Events, Read Only), 2: Track No (tracks.track_no), 3: Track Name (tracks.track_name)
            col_map = {"2": "track_no", "3": "track_name"}
            table.setProperty("column_mapping", col_map)
            
            table.blockSignals(True)
            font_size = getattr(self, 'current_font_size', 12) # Use current setting or default
            
            for r, row_data in enumerate(rows):
                db_id = row_data[0] # t.id
                
                for c, val in enumerate(row_data):
                     display_val = str(val) if val is not None else ""
                     item = QTableWidgetItem(display_val)
                     item.setData(Qt.ItemDataRole.UserRole, db_id) 
                     
                     # Make ID and Audio No Read-Only
                     if c in [0, 1]:
                         item.setFlags(item.flags() ^ Qt.ItemFlag.ItemIsEditable)
                     else:
                         item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)

                     # Font for "Tracks" (Col 3)
                     if c == 3:
                         item.setFont(QFont(self.lisan_font_family, font_size))
                         # Ensure no extra newlines in display
                         item.setText(display_val.strip())
                     
                     table.setItem(r, c, item)
            
            table.blockSignals(False)
            
            # UI Formatting
            table.setColumnHidden(0, True) # Hide ID
            
            header = table.horizontalHeader()
            header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents) # Audio No
            header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents) # Track No
            header.setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch) # Tracks Title -> Stretch
            
            table.setWordWrap(True)
            # Force layout before resizing rows
            table.resizeColumnsToContents() 
            
            # Robust Resizing
            def resize_rows():
                table.resizeRowsToContents()
                
            QTimer.singleShot(100, resize_rows)
            QTimer.singleShot(500, resize_rows) # Retry later just in case
            
            # Connect column resize to row resize for dynamic adjustment
            # valid connection check
            try: header.sectionResized.disconnect() 
            except: pass
            header.sectionResized.connect(lambda: QTimer.singleShot(50, resize_rows))

        except Exception as e:
            print(f"Error loading tracks sheet: {e}")
            traceback.print_exc()
        finally:
            conn.close()

    def load_other_sheets(self):
        """Loads 'Other Sheets' from the database as tabs in the main view."""
        conn = sqlite3.connect(DB_FILE)
        try:
            sheets = conn.execute("SELECT id, name FROM old_sheets_meta ORDER BY id ASC").fetchall()
            
            # Get existing tab names to prevent duplicates
            existing_tabs = {self.sheets_tabs.tabText(i) for i in range(self.sheets_tabs.count())}
            
            for sheet_id, name in sheets:
                # Load ALL generic sheets, not just hardcoded ones
                if name not in existing_tabs:
                    self.create_other_sheet_tab(sheet_id, name)
        except Exception as e:
            print(f"Error loading other sheets: {e}")
        finally:
            conn.close()

    def create_other_sheet_tab(self, sheet_id, name):
        """Creates a single 'Other Sheet' tab."""
        # Table
        table = OtherSheetTable(self)
        table.setProperty("sheet_id", sheet_id)
        table.setAlternatingRowColors(self.alternating_rows)
        
        # Delegate
        delegate = EditorDelegate(self.lisan_font_family, "text", table)
        table.setItemDelegate(delegate)

        # Wrapping & Resizing
        table.setWordWrap(True)
        table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        # Debounced Resize (Reused logic)
        table.resize_timer = QTimer(table)
        table.resize_timer.setSingleShot(True)
        table.resize_timer.setInterval(500)
        table.resize_timer.timeout.connect(table.resizeRowsToContents)
        table.horizontalHeader().sectionResized.connect(lambda: table.resize_timer.start())
        
        # Lock State
        if self.master_locked:
            table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        else:
            table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed)
        
        # Load Headers
        conn = sqlite3.connect(DB_FILE)
        headers_res = conn.execute("SELECT col_index, header_name FROM old_sheet_headers WHERE sheet_id = ? ORDER BY col_index ASC", (sheet_id,)).fetchall()
        
        if headers_res:
             cols = len(headers_res)
             table.setColumnCount(cols)
             table.setHorizontalHeaderLabels([h[1] for h in headers_res])

             # Restore Saved Widths
             if hasattr(self, 'saved_ui_widths') and name in self.saved_ui_widths:
                 widths = self.saved_ui_widths[name]
                 if len(widths) == cols:
                     for c, w in enumerate(widths):
                         table.setColumnWidth(c, w)
        
        # Load Data
        try:
             data = conn.execute(f"SELECT rowid, * FROM old_sheet_data_{sheet_id}").fetchall()
             table.setRowCount(len(data))
             
             for r_idx, row_data in enumerate(data):
                 row_id = row_data[0]
                 vals = row_data[1:]
                 for c_idx, val in enumerate(vals):
                     val_str = str(val) if val is not None else ""
                     
                     # Get Header Name for this column
                     header_name = ""
                     if headers_res and c_idx < len(headers_res):
                         header_name = headers_res[c_idx][1]
                         
                     # Special Handling for "Year" Column -> Strict Integer
                     if "Year" in header_name or "year" in header_name:
                         # 1. Remove .0 (Float Artifact) FIRST
                         if val_str.endswith(".0"):
                             val_str = val_str[:-2]
                         # 2. Strict Integer: Remove anything that is not a digit
                         val_str = re.sub(r"\D", "", val_str)
                     else:
                         # Date Cleanup for other columns
                         # 1. YYYY__MM__DD or YYYY-MM-DD (e.g. 1960__06__26 00:00:00)
                         m1 = re.match(r"(\d{4})[_\W]+(\d{2})[_\W]+(\d{2})", val_str)
                         if m1:
                             val_str = f"{m1.group(3)}/{m1.group(2)}/{m1.group(1)}"
                         else:
                             # 2. Weird Format: 19/60/0629 -> 29/06/1960 (CC/YY/MMDD)
                             m2 = re.match(r"(\d{2})/(\d{2})/(\d{2})(\d{2})", val_str)
                             if m2:
                                  val_str = f"{m2.group(4)}/{m2.group(3)}/{m2.group(1)}{m2.group(2)}"
                             else:
                                  # 3. Handle standard YYYY-MM-DD if explicit
                                   m3 = re.match(r"(\d{4})-(\d{2})-(\d{2})", val_str)
                                   if m3:
                                       val_str = f"{m3.group(3)}/{m3.group(2)}/{m3.group(1)}"
                     
                     # Final Integer Cleanup for non-Year columns
                     if val_str.endswith(".0"):
                         val_str = val_str[:-2]
                         
                     it = QTableWidgetItem(val_str)
                     it.setData(Qt.ItemDataRole.UserRole, row_id)
                     
                     if hasattr(self, 'lisan_font_family'):
                        it.setFont(QFont(self.lisan_font_family, self.current_font_size))
                        
                     table.setItem(r_idx, c_idx, it)
             
             table.resizeColumnsToContents()
        except Exception as e:
             print(f"Error loading data for sheet {name}: {e}")
             
        conn.close()
        
        # Data Save
        table.itemChanged.connect(lambda item: self.save_other_sheet_cell(item)) # We need a save method
        
        self.sheets_tabs.addTab(table, name)

    def save_other_sheet_cell(self, item):
        """Saves a single cell edit for 'Other Sheets' or Generic Tables."""
        if self.master_locked: return
        
        table = item.tableWidget()
        if not table: return
        
        sheet_id = table.property("sheet_id")
        table_name_prop = table.property("table_name")
        col_map_prop = table.property("column_mapping") # Dict[int, str] mapping view col index to db col name
        
        # Determine Table Name
        if table_name_prop:
            target_table = table_name_prop
        elif sheet_id:
            target_table = f"old_sheet_data_{sheet_id}"
        else:
            return

        row = item.row()
        col = item.column()
        val = item.text()
        
        # Get ID (Stored in UserRole)
        row_id = item.data(Qt.ItemDataRole.UserRole)
        if row_id is None: return 
        
        conn = sqlite3.connect(DB_FILE)
        try:
            target_col = None
            
            # 1. Check Explicit Mapping
            if col_map_prop and str(col) in col_map_prop:
                 target_col = col_map_prop[str(col)]
            
            # 2. Fallback to Schema/Index Matching
            if not target_col:
                # Get Schema
                table_info = conn.execute(f"PRAGMA table_info({target_table})").fetchall()
                real_cols = [info[1] for info in table_info]
                
                if table_name_prop:
                     if col < len(real_cols):
                         target_col = real_cols[col]
                else:
                    headers = conn.execute("SELECT header_name FROM old_sheet_headers WHERE sheet_id = ? ORDER BY col_index ASC", (sheet_id,)).fetchall()
                    std_col = f"col_{col}"
                    if std_col in real_cols:
                        target_col = std_col
                    elif col < len(headers):
                         target_col = headers[col][0]
            
            if target_col:
                # Update Query
                # Check for PK
                table_info = conn.execute(f"PRAGMA table_info({target_table})").fetchall()
                real_cols = [info[1] for info in table_info]
                pk_col = next((info[1] for info in table_info if info[5] > 0), "id")
                
                if pk_col in real_cols:
                     where_clause = f"\"{pk_col}\" = ?"
                else:
                     where_clause = "rowid = ?"
                
                query = f"UPDATE {target_table} SET \"{target_col}\" = ? WHERE {where_clause}"
                conn.execute(query, (val, row_id))
                conn.commit()
                self.mark_unsaved()

        except Exception as e:
            print(f"Save Cell Error: {e}")
        finally:
            conn.close()

    def handle_add_rows(self):
        """Context-aware Add with dynamic count."""
        # Prompt for number of rows
        count, ok = QInputDialog.getInt(self, "Add Rows", "Number of rows to add:", 1, 1, 10000, 1)
        if not ok:
            return

        idx = self.sheets_tabs.currentIndex()
        if idx == 0:
            # Master Sheet
            self.add_rows_master(count)
        else:
            # Other Sheet
            self.add_rows_other(self.sheets_tabs.currentWidget(), count)

    def add_rows_other(self, table, count):
        if not table or not isinstance(table, OtherSheetTable): return
        
        sheet_id = table.property("sheet_id")
        if not sheet_id: return
        
        conn = sqlite3.connect(DB_FILE)
        try:
            cur = conn.cursor()
            row_ids = []
            
            # Batch insert into DB
            for _ in range(count):
                cur.execute(f"INSERT INTO old_sheet_data_{sheet_id} DEFAULT VALUES")
                row_ids.append(cur.lastrowid)
            conn.commit()
            
            # Update UI
            # Disable updates for performance
            table.setUpdatesEnabled(False)
            current_row = table.rowCount()
            
            for i, row_id in enumerate(row_ids):
                row_idx = current_row + i
                table.insertRow(row_idx)
                for c in range(table.columnCount()):
                    it = QTableWidgetItem("")
                    it.setData(Qt.ItemDataRole.UserRole, row_id)
                    if hasattr(self, 'lisan_font_family'):
                        it.setFont(QFont(self.lisan_font_family, self.current_font_size))
                    table.setItem(row_idx, c, it)
            
            table.setUpdatesEnabled(True)
            table.scrollToBottom()
            
            QMessageBox.information(self, "Success", f"Added {count} row(s).")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add rows: {e}")
        finally:
            conn.close()

    def search_active_sheet(self):
        """Context-aware Search."""
        idx = self.sheets_tabs.currentIndex()
        if idx == 0:
            self.search_entries_master()
        else:
            self.search_other_sheet_ui()

    def search_other_sheet_ui(self):
        """Simple UI-based filter for Other Sheets."""
        term = self.txt_search_master.text().strip().lower()
        table = self.sheets_tabs.currentWidget()
        if not table or not isinstance(table, QTableWidget): return
        
        for r in range(table.rowCount()):
            visible = False
            if not term:
                visible = True
            else:
                for c in range(table.columnCount()):
                    it = table.item(r, c)
                    if it and term in it.text().lower():
                        visible = True
                        break
            table.setRowHidden(r, not visible)

    def toggle_global_lock(self):
        """Toggles lock for ALL sheets."""
        # Toggle state
        self.master_locked = not self.master_locked
        
        # Update Button
        if self.master_locked:
            self.btn_master_lock.setText("🔒 MASTER LOCKED")
            self.btn_master_lock.setChecked(True)
            self.btn_master_lock.setStyleSheet("background-color: #742a2a; color: white; font-weight: bold; border: 2px solid #e53e3e; border-radius: 6px;")
        else:
            self.btn_master_lock.setText("🔓 UNLOCKED")
            self.btn_master_lock.setChecked(False)
            self.btn_master_lock.setStyleSheet("background-color: #38a169; color: white; font-weight: bold; border: 2px solid #2f855a; border-radius: 6px;")
            
        # Apply to ALL tabs
        for i in range(self.sheets_tabs.count()):
            w = self.sheets_tabs.widget(i)
            if isinstance(w, QTableWidget):
                if self.master_locked:
                    w.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
                else:
                    w.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed | QAbstractItemView.EditTrigger.AnyKeyPressed)

        # Update Add Button State
        self.btn_add.setEnabled(not self.master_locked)

    def export_active_sheet_selection(self):
        """Context-aware Export."""
        idx = self.sheets_tabs.currentIndex()
        if idx == 0:
            self.export_selection() # Existing method for Master
        else:
            self.export_simple_table_selection(self.sheets_tabs.currentWidget(), "Export")

    def refresh_active_sheet(self):
        """Context-aware Refresh."""
        idx = self.sheets_tabs.currentIndex()
        if idx == 0:
            # Refresh Master
            self.refresh_master_sheet_action()
        else:
            # Refresh Other Sheet (Reload Tab)
            msg = QMessageBox.question(self, "Refresh", "Reload all other sheets?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if msg == QMessageBox.StandardButton.Yes:
                # Remove tabs > 0
                while self.sheets_tabs.count() > 1:
                    self.sheets_tabs.removeTab(1)
                self.load_other_sheets()


# --- 5. SEARCH TABS CLASSES ---

class TrackSearchTab(QWidget):
    def __init__(self, parent_app):
        super().__init__()
        self.parent_app = parent_app
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)
        
        # --- TOP BAR ---
        top = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search tracks...")
        self.search_input.setFixedHeight(35)
        self.search_input.setFont(QFont(self.parent_app.lisan_font_family, 13))
        self.search_input.returnPressed.connect(self.run_track_search)
        
        self.btn_search = QPushButton("🔍 Search")
        self.btn_search.setFixedHeight(35)
        self.btn_search.setObjectName("PrimaryBtn")
        self.btn_search.clicked.connect(self.run_track_search)
        
        self.btn_c = QPushButton("🔄 Clear")
        self.btn_c.setFixedHeight(35)
        self.btn_c.clicked.connect(self.clear_all)
        self.btn_c.setObjectName("SecondaryBtn")
        
        self.btn_goto = QPushButton("📂 Go to Master")
        self.btn_goto.setFixedHeight(35)
        self.btn_goto.clicked.connect(self.handle_go_to_master)
        self.btn_goto.setObjectName("SecondaryBtn")

        # top.addWidget(QLabel("🔍 Track Text:")) # Removed label
        top.addWidget(self.search_input, 4)
        top.addWidget(self.btn_search, 1)
        top.addWidget(self.btn_c, 1)
        top.addWidget(self.btn_goto, 1)
        layout.addLayout(top)
        
        # --- FILTERS ---
        filters_l = QHBoxLayout()
        self.filters = {}
        # Same filters as EntrySearchTab
        for name in ["Person", "Occasion", "Category", "Place", "AV", "Year"]:
            v = QVBoxLayout()
            v.addWidget(QLabel(f"{name}:"))
            combo = FilterComboBox(name, self, self.parent_app.lisan_font_family)
            combo.currentIndexChanged.connect(self.run_track_search)
            self.filters[name] = combo
            v.addWidget(combo)
            filters_l.addLayout(v)
        layout.addLayout(filters_l)

        # --- RESULTS TABLE ---
        self.search_results = QTableWidget()
        self.search_results.setColumnCount(5)
        self.search_results.setHorizontalHeaderLabels(["Audio No", "Track No", "Tracks", "Person", "Date"])
        self.search_results.verticalHeader().setVisible(False)
        self.search_results.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.search_results.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        
        # User Request: Size to table, Wrap Text
        self.search_results.setWordWrap(True)
        self.search_results.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        self.search_results.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents) # Audio No
        self.search_results.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents) # Track No
        self.search_results.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch) # Tracks (Main)
        self.search_results.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents) # Person
        self.search_results.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents) # Date
        
        # Explicit resize Trigger
        self.search_results.horizontalHeader().sectionResized.connect(lambda *a: self.search_results.resizeRowsToContents())
        
        # Pass self.parent_app explicitly via a property or just ensure Delegate works
        # The EditorDelegate uses `parent_table.parent_app` if available
        self.search_results.parent_app = self.parent_app 
        
        self.search_results.setItemDelegate(EditorDelegate(self.parent_app.lisan_font_family, "text", self.search_results))
        layout.addWidget(self.search_results)
        
        self.update_dropdown_items()

    def update_dropdown_items(self):
        # Re-using logic from EntrySearchTab approximately
        for name, combo in self.filters.items():
            combo.blockSignals(True)
            combo.clear()
            items = []
            if name == "Person": items = self.parent_app.master_lists.get(2, [])
            elif name == "Occasion": items = self.parent_app.master_lists.get(3, [])
            elif name == "Category": items = self.parent_app.master_lists.get(4, [])
            elif name == "Place": items = self.parent_app.master_lists.get(5, [])
            elif name == "AV": items = self.parent_app.master_lists.get(13, [])
            elif name == "Year":
                conn = sqlite3.connect(DB_FILE)
                query_years = "SELECT DISTINCT year FROM events WHERE year != '' ORDER BY year DESC"
                items = [str(r[0]) for r in conn.execute(query_years).fetchall()]
                conn.close()
            combo.addItems([""] + items)
            combo.setCurrentIndex(0)
            combo.blockSignals(False)

    def clear_all(self):
        self.search_input.clear()
        for c in self.filters.values():
            c.blockSignals(True) # Prevent triggering search on every clear
            c.setCurrentIndex(0)
            c.blockSignals(False)
        self.run_track_search()

    def update_cascading_filters(self):
        """Dynamic filter updates for Track Search."""
        q = self.search_input.text().strip()
        
        # Snapshot current selections
        current_selections = {k: v.currentText().strip() for k, v in self.filters.items()}
        
        mapping = {"Person": "person", "Occasion": "occasion", "Category": "category", 
                   "Place": "place", "AV": "AV", "Year": "year"}
        
        conn = sqlite3.connect(DB_FILE)
        
        for target_name, target_combo in self.filters.items():
            target_col = mapping.get(target_name)
            if not target_col: continue
            
            # Build Query
            sql_parts = ["SELECT DISTINCT e." + target_col]
            sql_parts.append("FROM events e")
            
            # Join only if needed for text search on tracks
            if q:
                sql_parts.append("JOIN tracks t ON t.event_id = e.audio_no")
            
            where = ["1=1"]
            params = []
            
            if q:
                where.append("t.track_name LIKE ?")
                params.append(f"%{q}%")
            
            # Add other filters
            for other_name, other_val in current_selections.items():
                if other_name == target_name: continue # Skip self
                if other_val:
                    other_col = mapping.get(other_name)
                    # Filter for non-empty only
                    where.append(f"e.{other_col} = ?")
                    params.append(other_val)
            
            sql_parts.append("WHERE " + " AND ".join(where))
            sql_parts.append(f"ORDER BY e.{target_col} ASC")
            
            query = " ".join(sql_parts)
            
            try:
                rows = conn.execute(query, params).fetchall()
                # Clean items: valid strings only
                unique_vals = sorted(list(set(str(r[0]) for r in rows if r[0])), key=lambda x: x.lower())
                final_items = [""] + unique_vals
            except Exception as e:
                final_items = [""]
                
            # Update Combo
            target_combo.blockSignals(True)
            target_combo.clear()
            target_combo.addItems(final_items)
            
            # Restore selection if valid
            prev_val = current_selections[target_name]
            if prev_val in final_items:
                target_combo.setCurrentText(prev_val)
            else:
                target_combo.setCurrentIndex(0) # Reset if no longer valid
            
            target_combo.blockSignals(False)
            
        conn.close()

    def run_track_search(self):
        # 1. Update Filters
        self.update_cascading_filters()
        
        raw_q = self.search_input.text()
        q = raw_q.strip()
        
        # --- GLOBAL HIGHLIGHT UPDATE ---
        if hasattr(self.parent_app, 'current_search_term'):
            self.parent_app.current_search_term = q
            # Force repaint of this table
            self.search_results.viewport().update()
        # -------------------------------

        # Filter check
        filters_active = any(c.currentIndex() > 0 for c in self.filters.values())
        
        if not q and not filters_active:
            self.search_results.setRowCount(0)
            return

        conn = sqlite3.connect(DB_FILE)
        
        # Base Query
        sql = """SELECT t.track_no, t.track_name, e.audio_no, e.person, e.esavi_date 
                 FROM tracks t 
                 JOIN events e ON t.event_id = e.audio_no 
                 WHERE 1=1"""
        
        params = []
        
        # Term Search
        if q:
            sql += " AND t.track_name LIKE ?"
            params.append(f'%{q}%')
            
        # Filters
        mapping = {"Person": "e.person", "Occasion": "e.occasion", "Category": "e.category", 
                   "Place": "e.place", "AV": "e.AV", "Year": "e.year"}
        
        for name, col_db in mapping.items():
            val = self.filters[name].currentText().strip()
            if val:
                sql += f" AND {col_db} = ?"
                params.append(val)
                
        # Limit results to avoid freezing if query is too broad
        sql += " LIMIT 1000"

        try:
            df = pd.read_sql_query(sql, conn, params=params)
        except Exception as e:
            print(f"Search Error: {e}")
            df = pd.DataFrame()
            
        conn.close()
        
        self.search_results.setRowCount(0)
        for i, row in df.iterrows():
            self.search_results.insertRow(i)
            # Columns: Audio No, Track No, Tracks, Person, Date
            # DF Columns: track_no, track_name, audio_no, person, esavi_date
            
            # CHANGED: Reordered Values
            vals = [row['audio_no'], row['track_no'], row['track_name'], row['person'], row['esavi_date']]
            
            for j, val in enumerate(vals):
                # Clean value
                if pd.isna(val): val = ""
                val = str(val)
                item = LisanTableItem(val, self.parent_app.lisan_font_family, self.parent_app.current_font_size)
                
                # ALIGNMENT LOGIC
                align = Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
                if j == 2: # Tracks Column: Left Align
                    align = Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter
                
                # CRITICAL: Sync both roles to prevent jumping on click
                item.setTextAlignment(align)
                item.setData(Qt.ItemDataRole.TextAlignmentRole, align)
                
                self.search_results.setItem(i, j, item)
            
        # Final sizing tweaks
        self.search_results.resizeColumnToContents(0)
        self.search_results.resizeColumnToContents(1)
        self.search_results.resizeColumnToContents(3)
        self.search_results.resizeColumnToContents(4) # Date
        self.search_results.resizeRowsToContents()

    def handle_go_to_master(self):
        """Toolbar handler to jump to selected entry."""
        row = self.search_results.currentRow()
        if row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a row to go to.")
            return

        it = self.search_results.item(row, 0)
        if it:
            audio_no = it.text()
            if audio_no:
                self.parent_app.jump_to_audio_no(audio_no)



if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Global Font
    if os.path.exists(FONT_PATH):
        QFontDatabase.addApplicationFont(FONT_PATH)
    
    # 1. Show Splash with Fade-In
    splash = ModernSplash()
    splash.show()
    splash.raise_()
    splash.fade_in()
    
    # Wait for Fade-In to REALLY complete (No blocking here)
    while splash.fade_in_anim.state() == QPropertyAnimation.State.Running:
        QApplication.processEvents()
        time.sleep(0.01)
    
    # Ensure it's fully solid after animation
    splash.setWindowOpacity(1.0)
    QApplication.processEvents()
    
    # 2. Initialize App (Heavy lifting) - Background is now fully rendered
    # Pass splash for progress updates
    main_window = ArchiveApp(splash)
    
    # Extra pause to ensure user sees the splash (Min 1.5 seconds visible)
    # The progress bar is full now
    start_time = time.time()
    while time.time() - start_time < 0.5:
        QApplication.processEvents()
        time.sleep(0.05)
        
    # 3. Fade Out and then Show Main Window
    def reveal_app():
        main_window.showMaximized()
        main_window.raise_()
        main_window.activateWindow()
        
    splash.fade_out(reveal_app)
    
    sys.exit(app.exec())